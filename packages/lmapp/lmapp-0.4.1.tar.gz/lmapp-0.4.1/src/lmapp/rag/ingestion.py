"""
Document Ingestion for LMAPP RAG.
Supports Text, Markdown, PDF, and DOCX.
"""

import logging
from pathlib import Path
from typing import Optional
import mimetypes

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

from .models import Document

logger = logging.getLogger(__name__)


class DocumentIngestor:
    """Handles loading and parsing of various file formats."""

    def load_file(self, file_path: str) -> Optional[Document]:
        """Load a file and return a Document object."""
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        mime_type, _ = mimetypes.guess_type(file_path)
        content = ""

        try:
            if path.suffix.lower() == ".pdf":
                content = self._read_pdf(path)
            elif path.suffix.lower() in [".docx", ".doc"]:
                content = self._read_docx(path)
            else:
                # Default to text/markdown
                content = path.read_text(encoding="utf-8", errors="replace")

            if not content:
                return None

            return Document(
                doc_id=str(path.absolute()),
                title=path.name,
                content=content,
                file_path=str(path.absolute()),
                source_type=path.suffix.lower().lstrip(".") or "text",
            )

        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            return None

    def _read_pdf(self, path: Path) -> str:
        """Read PDF content."""
        if pypdf is None:
            logger.warning("pypdf not installed. Skipping PDF.")
            return ""

        text = []
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text())
        return "\n".join(text)

    def _read_docx(self, path: Path) -> str:
        """Read DOCX content."""
        if docx is None:
            logger.warning("python-docx not installed. Skipping DOCX.")
            return ""

        doc = docx.Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
