"""
Script chuyển đổi file Markdown (.md) thành file Word (.docx)

Cách sử dụng:
    python convert_md_to_word.py <đường_dẫn>                           # Chuyển file/thư mục
    python convert_md_to_word.py <đường_dẫn> --autofit content         # Table autofit to content
    python convert_md_to_word.py <đường_dẫn> --autofit window          # Table autofit to window (mặc định)

Ví dụ:
    python convert_md_to_word.py .
    python convert_md_to_word.py D:/Documents/README.md
    python convert_md_to_word.py D:/Documents/README.md --autofit content

Yêu cầu cài đặt:
    pip install python-docx markdown beautifulsoup4
"""

import sys
import os
import io
import argparse

# Fix encoding for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
from pathlib import Path
from typing import List

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import markdown
    from bs4 import BeautifulSoup
except ImportError as e:
    print(f"Lỗi: Thiếu thư viện. Vui lòng cài đặt bằng lệnh:")
    print("pip install python-docx markdown beautifulsoup4")
    sys.exit(1)


def convert_md_to_html(md_content: str) -> str:
    """Chuyển đổi Markdown sang HTML"""
    extensions = [
        'tables',
        'fenced_code',
        'codehilite',
        'toc',
        'nl2br'
    ]
    return markdown.markdown(md_content, extensions=extensions)


def add_formatted_text(paragraph, text: str, bold: bool = False,
                       italic: bool = False, code: bool = False):
    """Thêm text với định dạng vào paragraph"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


def process_html_element(doc: Document, element, current_list_level: int = 0, autofit_mode: str = 'window'):
    """Xử lý từng element HTML và thêm vào document Word"""

    if element.name is None:  # Text node
        return

    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        heading = doc.add_heading(element.get_text().strip(), level=level)

    elif element.name == 'p':
        para = doc.add_paragraph()
        process_inline_elements(para, element)

    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            para = doc.add_paragraph(style='List Bullet')
            process_inline_elements(para, li)
            # Xử lý nested lists
            nested_ul = li.find('ul')
            if nested_ul:
                for nested_li in nested_ul.find_all('li', recursive=False):
                    nested_para = doc.add_paragraph(style='List Bullet 2')
                    process_inline_elements(nested_para, nested_li)

    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            para = doc.add_paragraph(style='List Number')
            process_inline_elements(para, li)

    elif element.name == 'pre':
        # Code block
        code_text = element.get_text()
        para = doc.add_paragraph()
        run = para.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        para.paragraph_format.left_indent = Inches(0.3)

    elif element.name == 'code' and element.parent.name != 'pre':
        # Inline code - sẽ được xử lý trong process_inline_elements
        pass

    elif element.name == 'table':
        process_table(doc, element, autofit_mode)

    elif element.name == 'blockquote':
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        process_inline_elements(para, element)

    elif element.name == 'hr':
        para = doc.add_paragraph('─' * 50)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def process_inline_elements(paragraph, element):
    """Xử lý các inline elements (bold, italic, code, link)"""
    for child in element.children:
        if child.name is None:  # Text node
            text = str(child)
            if text.strip():
                paragraph.add_run(text)
        elif child.name == 'strong' or child.name == 'b':
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name == 'em' or child.name == 'i':
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif child.name == 'a':
            run = paragraph.add_run(child.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        elif child.name in ['span', 'p']:
            process_inline_elements(paragraph, child)
        else:
            # Đối với các element khác, lấy text
            text = child.get_text()
            if text.strip():
                paragraph.add_run(text)


def set_cell_shading(cell, color_hex: str):
    """Thiết lập màu nền cho cell"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, num_rows: int):
    """
    Thiết lập border cho table:
    - Border thường: 1/2 pt, Blue Accent 1 Lighter 60% (B4C6E7)
    - Border ngăn cách header với content: 3/2 pt (1.5pt)
    """
    # Blue Accent 1 Lighter 60% = B4C6E7
    BORDER_COLOR = "B4C6E7"
    # 1/2 pt = 4 (trong đơn vị eighths of a point)
    NORMAL_BORDER_SIZE = "4"
    # 3/2 pt = 1.5pt = 12 eighths
    HEADER_BORDER_SIZE = "12"

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Xóa border cũ nếu có
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    # Tạo table borders
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Thiết lập border bottom của header row = 3/2 pt
    if num_rows > 1:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # Xóa tcBorders cũ nếu có
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            # Thêm border bottom dày hơn cho header cells
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
                f'  <w:left w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
                f'  <w:bottom w:val="single" w:sz="{HEADER_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
                f'  <w:right w:val="single" w:sz="{NORMAL_BORDER_SIZE}" w:color="{BORDER_COLOR}"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)


def set_table_autofit(table, mode: str = 'window'):
    """
    Autofit table theo mode:
    - 'content': autofit to content (tự động theo nội dung)
    - 'window': autofit to window (phân bố 100% chiều rộng trang)
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Xóa tblW cũ nếu có
    for child in list(tblPr):
        if child.tag.endswith('tblW'):
            tblPr.remove(child)

    if mode == 'content':
        # Autofit to content: w:type="auto" w:w="0"
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>')
    else:
        # Autofit to window: w:type="pct" w:w="5000" (100%)
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.insert(0, tblW)

    # Set autofit behavior
    for child in list(tblPr):
        if child.tag.endswith('tblLayout'):
            tblPr.remove(child)

    # tblLayout với type="autofit" để tự động điều chỉnh theo nội dung
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>')
    tblPr.append(tblLayout)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def process_table(doc: Document, table_element, autofit_mode: str = 'window'):
    """Xử lý bảng HTML và tạo bảng Word với định dạng đẹp"""
    rows = table_element.find_all('tr')
    if not rows:
        return

    # Đếm số cột từ row đầu tiên
    first_row = rows[0]
    cols = first_row.find_all(['th', 'td'])
    num_cols = len(cols)

    if num_cols == 0:
        return

    # Tạo bảng Word
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Màu sắc cho bảng
    HEADER_BG_COLOR = "4472C4"  # Màu xanh dương đậm cho header row
    FIRST_COL_BG_COLOR = "D6DCE5"  # Màu xám nhạt cho first column
    ALT_ROW_COLOR = "F2F2F2"  # Màu xám rất nhạt cho alternate rows

    for row_idx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for col_idx, cell in enumerate(cells):
            if col_idx < num_cols:
                table_cell = table.rows[row_idx].cells[col_idx]
                table_cell.text = cell.get_text().strip()

                # Định dạng cho header row (row đầu tiên)
                if row_idx == 0 or cell.name == 'th':
                    set_cell_shading(table_cell, HEADER_BG_COLOR)
                    for para in table_cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)  # Chữ trắng
                # Định dạng cho first column (trừ header)
                elif col_idx == 0:
                    set_cell_shading(table_cell, FIRST_COL_BG_COLOR)
                    for para in table_cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                # Alternate row colors cho các cell còn lại
                elif row_idx % 2 == 0:
                    set_cell_shading(table_cell, ALT_ROW_COLOR)

    # Thiết lập border cho table
    set_table_borders(table, len(rows))

    # Autofit theo mode
    set_table_autofit(table, autofit_mode)

    # Thêm spacing sau bảng
    doc.add_paragraph()


def setup_page_format(doc: Document):
    """Thiết lập định dạng trang: A4, margins"""
    for section in doc.sections:
        # Khổ giấy A4: 21cm x 29.7cm
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        # Margins: top=2cm, bottom=2cm, left=3cm, right=2cm
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)


def convert_md_file_to_docx(md_path: Path, output_path: Path, autofit_mode: str = 'window'):
    """Chuyển đổi một file MD sang DOCX"""
    print(f"  Đang xử lý: {md_path.name}")

    # Đọc file MD
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Chuyển MD sang HTML
    html_content = convert_md_to_html(md_content)

    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Tạo document Word
    doc = Document()

    # Thiết lập định dạng trang (A4, margins)
    setup_page_format(doc)

    # Thiết lập styles
    style = doc.styles['Normal']
    # style.font.name = 'Times New Roman'
    style.font.name = 'Cambria'
    style.font.size = Pt(12)

    # Xử lý từng element
    for element in soup.children:
        if element.name:
            process_html_element(doc, element, autofit_mode=autofit_mode)

    # Lưu file
    doc.save(str(output_path))
    print(f"  ✓ Đã tạo: {output_path.name}")


def convert_single_file(file_path: str, autofit_mode: str = 'window'):
    """Chuyển đổi một file MD cụ thể sang DOCX"""
    md_path = Path(file_path)

    if not md_path.exists():
        print(f"Lỗi: File '{file_path}' không tồn tại!")
        sys.exit(1)

    if not md_path.is_file():
        print(f"Lỗi: '{file_path}' không phải là file!")
        sys.exit(1)

    if md_path.suffix.lower() != '.md':
        print(f"Lỗi: File '{file_path}' không phải là file Markdown (.md)!")
        sys.exit(1)

    print(f"\nChuyển đổi file: {md_path.name}")
    print(f"  Table autofit: {autofit_mode}")

    try:
        output_file = md_path.with_suffix('.docx')
        convert_md_file_to_docx(md_path, output_file, autofit_mode)
        print(f"\n{'='*50}")
        print(f"Hoàn thành!")
        print(f"  - File output: {output_file}")
    except Exception as e:
        print(f"  ✗ Lỗi khi xử lý: {str(e)}")
        sys.exit(1)


def convert_all_md_in_directory(directory: str, autofit_mode: str = 'window'):
    """Chuyển đổi tất cả file MD trong thư mục sang DOCX"""
    dir_path = Path(directory)

    if not dir_path.exists():
        print(f"Lỗi: Thư mục '{directory}' không tồn tại!")
        sys.exit(1)

    if not dir_path.is_dir():
        print(f"Lỗi: '{directory}' không phải là thư mục!")
        sys.exit(1)

    # Tìm tất cả file .md
    md_files: List[Path] = list(dir_path.glob('*.md'))

    if not md_files:
        print(f"Không tìm thấy file .md nào trong thư mục '{directory}'")
        return

    print(f"\nTìm thấy {len(md_files)} file Markdown:")
    for f in md_files:
        print(f"  - {f.name}")
    print(f"  Table autofit: {autofit_mode}")

    print(f"\nBắt đầu chuyển đổi...")

    success_count = 0
    error_count = 0

    for md_file in md_files:
        try:
            # Tạo tên file output (đổi .md thành .docx)
            output_file = md_file.with_suffix('.docx')
            convert_md_file_to_docx(md_file, output_file, autofit_mode)
            success_count += 1
        except Exception as e:
            print(f"  ✗ Lỗi khi xử lý {md_file.name}: {str(e)}")
            error_count += 1

    print(f"\n{'='*50}")
    print(f"Hoàn thành!")
    print(f"  - Thành công: {success_count} file")
    if error_count > 0:
        print(f"  - Lỗi: {error_count} file")


# =========================
# MAIN
# =========================
def md2w_main():
    global LOG_FILE

    parser = argparse.ArgumentParser()
    parser.add_argument("path")
    parser.add_argument("-a", "--autofit", choices=["window", "content"], default="window")
    parser.add_argument("-r", "--recursive", action="store_true")
    parser.add_argument("-f", "--force", action="store_true", help="Ghi đè file docx đã tồn tại")
    parser.add_argument("--filter", help="Chỉ convert file md có chứa chuỗi này trong tên")
    parser.add_argument("--log", help="Ghi log ra file")

    args = parser.parse_args()
    LOG_FILE = args.log

    if LOG_FILE:
        with open(LOG_FILE, "w", encoding="utf-8") as f:
            f.write(f"md2w log - {datetime.now()}\n")

    p = Path(args.path)
    print(f"Đường dẫn: {p.resolve()}")
    if p.is_file() and p.suffix == ".md":
        out = p.with_suffix(".docx")
        if out.exists() and not args.force:
            print("File đã tồn tại, bỏ qua")
        else:
            print(f"\nChuyển đổi file: {p}")
            # convert_single_file(input_path,   args.autofit)

    elif args.recursive:
        for D,_,F in os.walk(p):
            for f in F:
                if f.lower().endswith(".md"):
                    md_path = Path(D) / f
                    if args.filter and args.filter not in f:
                        continue
                    out_path = md_path.with_suffix(".docx")
                    if out_path.exists() and not args.force:
                        print(f"File {md_path} đã tồn tại docx, bỏ qua")
                    else:
                        print(f"\nChuyển đổi file: {md_path}")
                        convert_single_file(md_path, args.autofit)    
    else:
        print(f"\nChuyển đổi tất cả file .md trong thư mục: {p}")
        convert_all_md_in_directory(p,  args.autofit)


if __name__ == "__main__":
    md2w_main()

