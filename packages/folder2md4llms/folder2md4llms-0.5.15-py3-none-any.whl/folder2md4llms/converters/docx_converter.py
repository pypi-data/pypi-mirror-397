"""DOCX document converter."""

import logging
from pathlib import Path
from typing import Any

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base import BaseConverter, ConversionError

logger = logging.getLogger(__name__)


class DOCXConverter(BaseConverter):
    """Converts DOCX files to text."""

    def __init__(self, config: dict[str, Any] | None = None):
        super().__init__(config)
        self.extract_images = False  # Image extraction disabled for simplicity

    def can_convert(self, file_path: Path) -> bool:
        """Check if this converter can handle the given file."""
        return (
            DOCX_AVAILABLE
            and file_path.suffix.lower() == ".docx"
            and file_path.exists()
        )

    def convert(self, file_path: Path) -> str | None:
        """Convert DOCX to text."""
        if not DOCX_AVAILABLE:
            return "DOCX conversion not available. Install python-docx: pip install python-docx"

        try:
            doc = Document(str(file_path))

            text_parts = []
            text_parts.append(f"Word Document: {file_path.name}")
            text_parts.append("=" * 50)
            text_parts.append("")

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                    text_parts.append("")

            # Extract tables
            if doc.tables:
                text_parts.append("=" * 30 + " TABLES " + "=" * 30)
                text_parts.append("")

                for i, table in enumerate(doc.tables):
                    text_parts.append(f"Table {i + 1}:")
                    text_parts.append("")

                    # Convert table to markdown format
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        text_parts.append(f"| {row_text} |")

                    text_parts.append("")

            result = "\n".join(text_parts)
            # Validate output to ensure no binary content leaked through
            return self._validate_text_output(result, file_path)

        except Exception as e:
            logger.error(f"Error converting DOCX {file_path}: {e}")
            raise ConversionError(f"Failed to convert DOCX: {str(e)}") from e

    def get_supported_extensions(self) -> set:
        """Get the file extensions this converter supports."""
        return {".docx"}

    def get_document_info(self, file_path: Path) -> dict[str, Any]:
        """Get DOCX-specific information."""
        info = self.get_file_info(file_path)

        if not DOCX_AVAILABLE:
            info["error"] = "DOCX library not available"
            return info

        try:
            doc = Document(str(file_path))

            # Count content elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # Calculate text length
            total_text = ""
            for paragraph in doc.paragraphs:
                total_text += paragraph.text

            info.update(
                {
                    "paragraphs": paragraph_count,
                    "tables": table_count,
                    "text_length": len(total_text),
                    "words": len(total_text.split()) if total_text else 0,
                }
            )

            # Try to get document properties
            if hasattr(doc, "core_properties"):
                props = doc.core_properties
                info.update(
                    {
                        "title": props.title or "",
                        "author": props.author or "",
                        "subject": props.subject or "",
                        "created": props.created.isoformat() if props.created else "",
                        "modified": props.modified.isoformat()
                        if props.modified
                        else "",
                    }
                )

        except Exception as e:
            info["error"] = str(e)

        return info
