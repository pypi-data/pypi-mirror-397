import copy
import traceback
import io

import base64
from typing import List
import warnings

import docx
from docx.shared import Inches, Pt, RGBColor

import tempfile
import os

from pathlib import Path
import zipfile, os, sys
from io import BytesIO

import markdown

try:
    from pydocmaker.backend.baseformatter import BaseFormatter
except Exception as err:
    from .baseformatter import BaseFormatter

can_run_pandoc = lambda : False

from enum import Enum

try:
    from pydocmaker.backend.pandoc_api import can_run_pandoc, pandoc_convert, pandoc_convert_file
except Exception as err:
    from .pandoc_api import can_run_pandoc, pandoc_convert, pandoc_convert_file

try:
    from pydocmaker.backend.ex_html import convert as convert_html
except Exception as err:
    from .ex_html import convert as convert_html



from docx import Document

gwin32 = None
gcomposer = None
gmailmerge = None

def _make_output(bts, output_path_or_buffer):
    
    if isinstance(output_path_or_buffer, (str, Path)):
        os.makedirs(os.path.dirname(output_path_or_buffer), exist_ok=True)
        with open(output_path_or_buffer, 'wb') as f:
            f.write(bts)
        return os.path.exists(output_path_or_buffer)
    elif hasattr(output_path_or_buffer, 'write'):
        return output_path_or_buffer.write(bts)
    else:
        return bts
    

def _get_bytes_file_or_buffer(file_path_or_buffer):
    if hasattr(file_path_or_buffer, "read"):
        bts_data = file_path_or_buffer.read()
    elif isinstance(file_path_or_buffer, bytes):
        bts_data = file_path_or_buffer
    else:
        # Read the original DOCX file
        with open(file_path_or_buffer, 'rb') as f:
            bts_data = f.read()
    return bts_data

def _test_docxw32_installed(verb=0, force_reload=False):
        """tests if win32com and Microsoft Word is available

        Returns:
            int: 0 if both are available, 1 if win32com is not available 2 if win32com is available and word is not available.
        """
        if _test_docxw32_installed.cache is None or force_reload:
            global gwin32

            try:
                
                import win32com.client
                if verb: print("win32com is available.")
                
                gwin32 = win32com.client

                # Attempt to create an instance of the Word application
                word = win32com.client.Dispatch("Word.Application")
                word.Quit()
                if verb: print("Microsoft Word is available.")
                _test_docxw32_installed.cache = 0
            except ImportError:
                if verb: print("win32com is not available.")
                _test_docxw32_installed.cache = 1
            except Exception as e:
                if verb: print(f"An error occurred: {e}")
                _test_docxw32_installed.cache = 2
        
        return _test_docxw32_installed.cache

_test_docxw32_installed.cache = None

class msoPictureCompress(Enum):
    """Enumeration of picture compression types for MS Office."""
    Default = 0
    HQPrint = 1
    Print = 2
    Email = 3
    Screen = 4
    Photo = 16

class DocxFileW32:
    """A context manager for working with DOCX files using Win32COM automation.
    
    This class provides methods to manipulate DOCX documents via Microsoft Word's COM interface,
    including image compression, field updates, and exporting to PDF format. It ensures proper
    cleanup by closing the Word application and document upon exiting the context manager.

    WARNING! This class needs the win32com library and word installed to work properly!

    """

    @staticmethod
    def is_installed(verb=0, force_reload=False, ret_int=False):
        r = _test_docxw32_installed(verb=verb, force_reload=force_reload)
        return r == 0 if not ret_int else r


    def __init__(self, docx_path, outpath=None):
        """Initialize the DocxFileW32 instance.
        
        WARNING! This class needs the win32com library and word installed to work properly!

        Args:
            docx_path (str): The path to the input DOCX file.
            outpath (str, optional): The output path for saving changes, None defaults to the input path. Defaults to None.
        """
        self.docx_path = docx_path
        self.outpath = outpath
        self.word = None
        self.worddoc = None
        self._created_word_app = False

    def __enter__(self):
        """Enter the runtime context for the DocxFileW32 instance. if "win32com.client" is not available it will be 
        imported. 

        Returns:
            DocxFileW32: The instance itself for use in a 'with' statement.
        """
        global gwin32
        if gwin32 is None:
            import win32com.client as _gwin32
            gwin32 = _gwin32

        try:
            # Try to get an existing Word application
            self.word = gwin32.GetActiveObject("Word.Application")
        except Exception:
            # If no Word application is running, create a new one
            self.word = gwin32.Dispatch("Word.Application")
            self.word.Visible = False  # Keep Word hidden
            self._created_word_app = True

        # Open the document
        self.worddoc = self.word.Documents.Open(self.docx_path)
        return self

    def compress_images(self, compressionQuality=msoPictureCompress.Screen):
        """Compress all images within the document based on specified quality settings.
        
        Args:
            compressionQuality (msoPictureCompress or int): The compression level to apply to images.
                Defaults to msoPictureCompress.Screen (quality 4).

        Returns:
            DocxFileW32: The instance itself to allow method chaining.
        """
        quality = compressionQuality.value if hasattr(compressionQuality, "value") else (int(compressionQuality))

        # Iterate through all shapes in the document
        for shape in self.worddoc.InlineShapes:
            if shape.Type == 13:  # 13 corresponds to a picture
                # Compress the image
                shape.PictureFormat.CompressionType = quality  # msoPictureCompressPhoto
        return self
    
    def update_fields(self):
        """Update all fields in the document, including those in headers and footers.

        Returns:
            DocxFileW32: The instance itself to allow method chaining.
        """
        self.worddoc.Fields.Update()

        for story_range in self.worddoc.StoryRanges:
            story_range.Fields.Update()
            # Some story ranges have linked ranges (e.g., next header/footer)
            while story_range.NextStoryRange is not None:
                story_range = story_range.NextStoryRange
                story_range.Fields.Update()
        return self
    
    def export(self, pdf_path=None, optimize_for_screen=True):
        """Export the document to PDF format with specified options.

        Args:
            pdf_path (str, optional): The path where the PDF will be saved. If None, uses outpath.
            optimize_for_screen (bool): True for wdExportOptimizeForOnScreen (=1) else wdExportOptimizeForPrint (=0).

        Returns:
            DocxFileW32: The instance itself to allow method chaining.
        
        Raises:
            AssertionError: If pdf_path is None and outpath is also None.
        """
        if pdf_path is None:
            assert self.outpath, 'If pdf_path is None, the objects "outpath" must be valid!'
            pdf_path = self.outpath

        pdf_path = pdf_path.replace('.docx', '.pdf')
        # Define the PDF save options
        pdf_options = {
            'OutputFileName': pdf_path,
            'ExportFormat': 17,  # FileFormat=17 is for PDF
            'OptimizeFor': 1 if optimize_for_screen else 0
        }

        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        self.worddoc.ExportAsFixedFormat(**pdf_options)
        
        return self
    
    def save(self):
        # Save the document
        if self.worddoc:
            if self.outpath is None:
                self.worddoc.Save()
            else:
                os.makedirs(os.path.dirname(self.outpath), exist_ok=True)
                self.worddoc.SaveAs(self.outpath)
        return self
    
    def close(self):
        # Close the document and quit Word
        if self.worddoc:
            self.worddoc.Close()
        # Quit the Word application if it was created by this script
        if self.word and self._created_word_app:
            self.word.Quit()
        return self
    
    def __exit__(self, exc_type, exc_value, traceback):
        """Exit the runtime context for the DocxFileW32 instance, saving and closing resources.

        Args:
            exc_type: Exception type (if any) that occurred during execution.
            exc_value: Exception value (if any) that occurred during execution.
            traceback: Traceback object (if any) that occurred during execution.
        """
        self.save()
        self.close()


# def docx_replace_color(file_path_or_buffer, output_path_or_buffer=None, color_to_replace=(0, 0, 0), new_color=(0, 112, 192), do_replace_default_color=True):

#     docx_data = _get_bytes_file_or_buffer(file_path_or_buffer)
        
#     # Create a temporary zip file from the DOCX data
#     doc = Document(BytesIO(docx_data))

#     if isinstance(color_to_replace, (tuple, list)):
#         color_to_replace = RGBColor(*color_to_replace)
    
#     if isinstance(new_color, (tuple, list)):
#         new_color = RGBColor(*new_color)
    

#     # Iterate through all paragraphs and runs
#     for para in doc.paragraphs:
#         for run in para.runs:
#             # Check if the run has a font color set to black
#             if (run.font.color is None and do_replace_default_color) or run.font.color.rgb == color_to_replace:
#                 run.font.color.rgb = new_color  # set to navyblue

#     with BytesIO() as fp:
#         doc.save(fp)
#         bts = fp.getvalue()

#     return _make_output(bts, output_path_or_buffer)




class DocxFile:
    """
    A class to handle operations on DOCX files such as appending multiple documents,
    replacing fields, and saving the modified document.
    """

    def __init__(self, file_path_or_buffer):
        """
        Initialize the DocxFile with a file path or buffer containing DOCX data.

        :param file_path_or_buffer: bytes, or a path to the DOCX file or a buffer containing DOCX data.
        """
                
        self.file_path_or_buffer = file_path_or_buffer
        
        self.docx_data = self._get_bytes_file_or_buffer(file_path_or_buffer)
        self.inp_data = copy.deepcopy(self.docx_data)
        
    def _get_bytes_file_or_buffer(self, file_path_or_buffer):
        """
        Retrieve bytes from a file path or buffer.

        :param file_path_or_buffer: Path to the DOCX file or a buffer containing DOCX data.
        :return: Bytes of the DOCX file.
        """
        # Existing implementation to get bytes from file path or buffer
        if isinstance(file_path_or_buffer, (io.BytesIO, bytes)):
            return file_path_or_buffer if isinstance(file_path_or_buffer, bytes) else file_path_or_buffer.getvalue()
        else:
            with open(file_path_or_buffer, 'rb') as f:
                return f.read()


    def append(self, *files, verb=0) -> bytes:
        """
        Append multiple Word (.docx) documents into a single document.

        :param files: Variable length argument list of either (bytes, or file paths, or buffers) to be appended.
        :param verb: Verbosity level (0 for silent, 1 for verbose).
        :return: The current instance of DocxFile with the appended DOCX data.
        """
        if not files:
            raise ValueError(f'Must supply files to merge, but given was {files=}!')
        
        global gcomposer

        if gcomposer is None:
            from docxcompose.composer import Composer
            gcomposer = Composer


        if isinstance(files[0], (tuple, list)):
            files = files[0]
        
        i = 0
        if verb: print(f'{i}/{len(files)} loading current document into composer ...')
        composer = gcomposer(Document(io.BytesIO(self.docx_data)))

        for i, file in enumerate(files):
            bts = self._get_bytes_file_or_buffer(file)
            if verb:
                print(f'{i}/{len(files)} appending doc')
                print("load report...")
                print(f'doc={file}')
            doc_b = Document(io.BytesIO(bts))
            if verb: print("adding report to template...")
            composer.append(doc_b)
        if verb: print("saving merged as bytes...")
            
        with io.BytesIO() as fp:
            composer.save(fp)
            fp.seek(0)
            bts = fp.getvalue()

        if verb: print("returning...")
        self.docx_data = bts
        return self

    def replace_fields(self, replace_dict):
        """
        Replace all MergeFields of a DOCX file with given text in form of a dict.

        :param replace_dict: Dictionary where keys are field names and values are replacement texts.
        :return: The current instance of DocxFile with the replaced fields.
        """
        global gmailmerge
        if gmailmerge is None:
            from mailmerge import MailMerge
            gmailmerge = MailMerge

        outbuf = BytesIO()

        with gmailmerge(BytesIO(self.docx_data)) as document:
            document.merge(**replace_dict)
            document.write(outbuf)

            
        self.docx_data = outbuf.getvalue()
        return self
    

    def get_fields(self):
        """
        Retrieve the merge fields from the DOCX document.

        Returns:
            list: A list of merge fields present in the DOCX document.
        """
        global gmailmerge
        if gmailmerge is None:
            from mailmerge import MailMerge
            gmailmerge = MailMerge

        with gmailmerge(BytesIO(self.docx_data)) as document:
            return list(document.get_merge_fields())
            

            
    

    def replace_keywords(self, replace_dict):
        """
        Edit raw XML content of a DOCX file by replacing specified strings using python-docx.

        :param replace_dict: Dictionary where keys are strings to be replaced and values are replacement strings.
        :return: The current instance of DocxFile with the replaced keywords.
        """
        doc = Document(BytesIO(self.docx_data))

        for p in doc.paragraphs:
            for key_to_replace, new_value in replace_dict.items():
                p.text = p.text.replace(str(key_to_replace), str(new_value))
    
        with BytesIO() as fp:
            doc.save(fp)
            bts = fp.getvalue()
        self.docx_data = bts

        return self

    def replace_keywords_raw(self, replace_dict):
        """
        Edit raw XML content of a DOCX file by replacing specified strings in all XML files within the document.

        :param replace_dict: Dictionary where keys are strings to be replaced and values are replacement strings.
        :return: The current instance of DocxFile with the replaced keywords in raw XML.
        """
        zip_buffer = BytesIO(self.docx_data)
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            file_list = zip_file.namelist()

            output_buffer = BytesIO()
            with zipfile.ZipFile(output_buffer, 'w', zipfile.ZIP_DEFLATED) as new_zip:
                for filename in file_list:
                    content = zip_file.read(filename)
                    
                    for key_to_replace, new_value in replace_dict.items():
                        content = content.replace(key_to_replace.encode('utf-8'), new_value.encode('utf-8'))
                    
                    new_zip.writestr(filename, content)
            bts = output_buffer.getvalue()
        self.docx_data = bts

        return self
    

    def save(self, output_path_or_buffer=None):
        """
        Save the modified DOCX data to a file or buffer.

        :param output_path_or_buffer: File path or buffer to save the DOCX data.
        :return: The current instance of DocxFile.
        """

        return _make_output(self.docx_data, output_path_or_buffer)




def blue(run):
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)

def red(run):
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

def convert_pandoc(doc:List[dict]) -> bytes:

    with tempfile.TemporaryDirectory() as temp_dir:
        html_file_path = os.path.join(temp_dir, 'temp.html')
        docx_file_path = os.path.join(temp_dir, 'temp.docx')

        with open(html_file_path, 'w', encoding='utf-8') as fp:
            fp.write(convert_html(doc))
        
        pandoc_convert_file(html_file_path, docx_file_path)
        with open(docx_file_path, 'rb') as fp:
            return fp.read()
        


def convert(doc:List[dict], template = None, template_params=None, use_w32=False, as_pdf=False, compress_images=False, filename=None, **kwargs) -> bytes:
    """
    Convert a list of document sections into a DOCX or PDF (via docx) file using a specified template.

    Parameters:
    - doc (List[dict]): The document content structured as a list of dictionaries.
    - template (str, optional): Path to a DOCX template file. Defaults to None.
    - template_params (dict, optional): Parameters to replace fields in the template. Defaults to None.
    - use_w32 (bool, optional): Whether to use win32com for document field updating and any of the following arguments, THIS OPTION NEEDS win32com and word installed. Defaults to False.
    - as_pdf (bool, optional): Whether to output the document as a PDF (via docx and win32com). Defaults to False.
    - compress_images (bool, optional): Whether to compress images in the document using win32com. Defaults to False.
    - filename (str, optional): The optional filename to give the document in case saving it as a tempfile is necessary. Default will try to get from metadata and if not found use tempfile.docx.
    - **kwargs: only used to check if invalid keyword arguments were passed.

    Returns:
    - bytes: The byte content of the generated DOCX or PDF file.

    Raises:
    - ValueError: If attempting to export to PDF without win32com and Word.Application installed and use_w32 set to True.
    """
    if not template_params:
        template_params = {}

    unknown_params = kwargs
    if unknown_params:
        warnings.warn(f'Unknown parameters passed: {unknown_params=}')
        
    _pandoc = can_run_pandoc()
    if _pandoc and not template:
        return convert_pandoc(doc)
    if _pandoc and template:
        bts = None
        bts_sub = convert_pandoc(doc)
        tmplt = DocxFile(template)
        tmplt.append(bts_sub)
        tmplt.replace_fields(template_params)
        if DocxFileW32.is_installed() and use_w32:
            with tempfile.TemporaryDirectory() as td:
                if not filename: # try to get filename from metadata or parameters
                    # get metadata from fields
                    metadata = next((k for k in doc if isinstance(k, dict) and k.get('typ') == 'meta'), {}).get('data', {})
                    metadata.update(template_params)
                    filename = metadata.get('filename', metadata.get('FILENAME', metadata.get('Filename', 'tempfile')))

                filename, ext = os.path.splitext(os.path.basename(filename))
                filepath = os.path.join(td, f'{filename}.docx')
                
                tmplt.save(filepath)
                with DocxFileW32(filepath) as docxw32:
                    docxw32.update_fields()
                    if compress_images:
                        docxw32.compress_images()
                    if as_pdf:
                        out = os.path.join(td, f'{filename}.pdf')
                        docxw32.export(out, optimize_for_screen=compress_images)
                        with open(out, 'rb') as fp:
                            bts = fp.read()

                if not as_pdf: # read back in
                    with open(filepath, 'rb') as fp:
                        bts = fp.read()

        elif use_w32:
            raise ValueError(f'{use_w32=} but either win32com and Word.Application is not installed! But was {use_w32=} and {DocxFileW32.is_installed(ret_int=True)=}')
        
        elif as_pdf:
            raise ValueError(f'can only export via docx to pdf if win32com and Word.Application is installed and the input parameter "use_w32" is True. But was {use_w32=} and {DocxFileW32.is_installed()=}')
        else:
            bts = tmplt.save()

        return bts

                    
    else:
        renderer = docx_renderer(None, make_blue=False)
        renderer.digest(doc)
        return renderer.doc_to_bytes()

class docx_renderer(BaseFormatter):
    def __init__(self, template_path:str=None, make_blue=False) -> None:
        self.d = docx.Document(template_path)
        self.make_blue = make_blue

    def add_paragraph(self, newtext, *args, **kwargs):
        new_paragraph = self.d.add_paragraph(newtext, *args, **kwargs)
        if self.make_blue:
            for r in new_paragraph.runs:
                blue(r)
        return new_paragraph
    
    def add_run(self, text, *args, **kwargs):
        if not self.d.paragraphs:
            self.add_paragraph('')

        last_paragraph = self.d.paragraphs[-1]
        
        if not last_paragraph.runs:
            last_run = last_paragraph.add_run(text)
        else:
            last_run = last_paragraph.runs[-1]
            last_run.add_text(text)
            
        if self.make_blue:
            blue(last_run)
        return last_run
        
    def digest_text(self, children, *args, **kwargs):
        return self.add_paragraph(children)
    

    def digest_str(self, children, *args, **kwargs):
        return self.add_run(children)

    def digest_line(self, children, *args, **kwargs):
        return self.add_run(children + '\n')
    
    def digest_markdown(self, children, *args, **kwargs):
        return self.add_paragraph(children, style='Normal')
        
    def digest_verbatim(self, children, *args, **kwargs):
        new_run = self.add_run(children)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        return new_run

    def digest_latex(self, children, *args, **kwargs):
        new_run = self.add_run(children)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        return new_run


    def handle_error(self, err, el=None) -> list:
        if isinstance(err, BaseException):
            traceback.print_exc(limit=5)
            err = '\n'.join(traceback.format_exception(type(err), value=err, tb=err.__traceback__, limit=5))

        new_run = self.add_run(err)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        red(new_run)
        return new_run


    def digest_iterator(self, children, *args, **kwargs):
        if children:
            return [self.digest(val, *args, **kwargs) for val in children]
        return []

    def digest_table(self, children=None, **kwargs) -> str:
        self.handle_error(NotImplementedError(f'exporter of type {type(self)} can not handle tables'))
    
    def digest_image(self, children, *args, **kwargs):

        image_width = Inches(max(1, kwargs.get('width', 0.8)*5))
        image_caption = kwargs.get('caption', '')
        image_blob = kwargs.get('imageblob', '')

        assert image_blob, 'no image data given!'

        btsb64 = image_blob.split(',')[-1]

        # Decode the base64 image
        img_bytes = base64.b64decode(btsb64)

        # Create an image stream from the bytes
        image_stream = io.BytesIO(img_bytes)
        
        picture = self.d.add_picture(image_stream, width=image_width)
        # picture.width = image_width  # Ensure fixed width
        # picture.height = None  # Adjust height automatically
        picture.alignment = 1

        run = self.add_paragraph(image_caption)
        # run.style = 'Caption'  # Apply the 'Caption' style for formatting

        return run

    def format(self, *args, **kwargs):
        raise NotImplementedError('Can not format a docx document directly')
    

    def doc_to_bytes(self):
        with io.BytesIO() as fp:
            self.d.save(fp)
            fp.seek(0)
            return fp.read()

    def save(self, filepath):
        self.d.save(filepath)
