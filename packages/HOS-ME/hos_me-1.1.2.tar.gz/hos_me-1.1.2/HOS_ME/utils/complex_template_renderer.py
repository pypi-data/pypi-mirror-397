from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import json
from datetime import datetime

class ComplexTemplateRenderer:
    """
    复杂模板渲染器，用于根据模板渲染包含表格、图片等复杂元素的文档
    """
    
    def __init__(self):
        pass
    
    def render_document(self, template, data):
        """
        根据模板和数据渲染文档
        
        Args:
            template: 模板信息，包含content和structure
            data: 要填充的数据
            
        Returns:
            Document: 渲染后的docx.Document对象
        """
        try:
            # 创建新文档
            doc = Document()
            
            # 获取模板结构
            structure = template.get('structure', {})
            content = template.get('content', '')
            
            # 渲染段落
            self._render_paragraphs(doc, structure.get('paragraphs', []), data)
            
            # 渲染表格
            self._render_tables(doc, structure.get('tables', []), data)
            
            # 渲染图片
            self._render_images(doc, structure.get('images', []), data)
            
            return {
                'success': True,
                'document': doc,
                'message': '文档渲染成功'
            }
        except Exception as e:
            return {
                'success': False,
                'message': f'渲染文档失败: {str(e)}'
            }
    
    def _render_paragraphs(self, doc, paragraphs, data):
        """渲染段落"""
        for para_data in paragraphs:
            text = para_data.get('text', '')
            format = para_data.get('format', {})
            
            # 替换占位符
            processed_text = self._replace_placeholders(text, data)
            
            # 添加段落
            paragraph = doc.add_paragraph(processed_text)
            
            # 设置段落格式
            self._set_paragraph_format(paragraph, format)
    
    def _render_tables(self, doc, tables, data):
        """渲染表格，支持跨行跨列"""
        for table_data in tables:
            rows = table_data.get('rows', 0)
            cols = table_data.get('cols', 0)
            cells = table_data.get('cells', [])
            
            if rows <= 0 or cols <= 0:
                continue
            
            # 创建表格
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'  # 使用默认表格样式
            
            # 填充表格内容
            for cell_data in cells:
                row = cell_data.get('row', 0)
                col = cell_data.get('col', 0)
                text = cell_data.get('text', '')
                format = cell_data.get('format', {})
                colspan = cell_data.get('colspan', 1)
                rowspan = cell_data.get('rowspan', 1)
                
                # 替换占位符
                processed_text = self._replace_placeholders(text, data)
                
                # 获取单元格
                if row < rows and col < cols:
                    cell = table.cell(row, col)
                    cell.text = processed_text
                    
                    # 处理跨列合并
                    if colspan > 1 and (col + colspan - 1) < cols:
                        # 向右合并单元格
                        cell.merge(table.cell(row, col + colspan - 1))
                    
                    # 处理跨行合并
                    if rowspan > 1 and (row + rowspan - 1) < rows:
                        # 向下合并单元格
                        cell.merge(table.cell(row + rowspan - 1, col))
                    
                    # 设置单元格格式
                    if cell.paragraphs:
                        self._set_paragraph_format(cell.paragraphs[0], format)
    
    def _render_images(self, doc, images, data):
        """渲染图片"""
        # 简化实现，实际需要根据图片占位符替换为真实图片
        # 这里仅添加图片占位符标记
        for image_data in images:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.text = f"[图片位置: {image_data.get('image_id', 'unknown')}]"
            run.bold = True
    
    def _replace_placeholders(self, text, data):
        """替换文本中的占位符"""
        import re
        
        # 匹配{{placeholder}}格式的占位符
        def replace_match(match):
            placeholder = match.group(1)
            return data.get(placeholder, match.group(0))
        
        return re.sub(r'{{(\w+)}}', replace_match, text)
    
    def _set_paragraph_format(self, paragraph, format):
        """设置段落格式"""
        # 设置对齐方式
        alignment = format.get('alignment', 'None')
        if alignment == 'CENTER':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'RIGHT':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'LEFT':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'JUSTIFY':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 设置行间距
        line_spacing = format.get('line_spacing', 1.0)
        paragraph.paragraph_format.line_spacing = line_spacing
        
        # 设置段前间距
        space_before = format.get('space_before', 0)
        paragraph.paragraph_format.space_before = space_before
        
        # 设置段后间距
        space_after = format.get('space_after', 0)
        paragraph.paragraph_format.space_after = space_after
    
    def save_document(self, doc, file_path):
        """
        保存文档到文件
        
        Args:
            doc: docx.Document对象
            file_path: 保存路径
            
        Returns:
            bool: 保存成功返回True，否则返回False
        """
        try:
            doc.save(file_path)
            return {
                'success': True,
                'message': '文档保存成功'
            }
        except Exception as e:
            return {
                'success': False,
                'message': f'保存文档失败: {str(e)}'
            }
    
    def render_and_save(self, template, data, file_path):
        """
        渲染模板并保存为DOCX文件
        
        Args:
            template: 模板信息
            data: 要填充的数据
            file_path: 保存路径
            
        Returns:
            dict: 渲染和保存结果
        """
        # 渲染文档
        render_result = self.render_document(template, data)
        
        if not render_result['success']:
            return render_result
        
        # 保存文档
        save_result = self.save_document(render_result['document'], file_path)
        
        if save_result['success']:
            return {
                'success': True,
                'file_path': file_path,
                'message': '文档渲染并保存成功'
            }
        else:
            return save_result
