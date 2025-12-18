from docx import Document
from docx.shared import RGBColor
import os
import json
from datetime import datetime

class DocxTemplateParser:
    """
    DOCX模板解析器，用于解析DOCX文件结构并生成模板配置
    """
    
    def __init__(self):
        pass
    
    def parse_docx(self, file_path):
        """
        解析DOCX文件，提取结构信息
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            dict: 包含模板结构信息的字典
        """
        try:
            # 加载DOCX文档
            doc = Document(file_path)
            
            # 解析文档结构
            structure = {
                'tables': self._parse_tables(doc),
                'images': self._parse_images(doc),
                'paragraphs': self._parse_paragraphs(doc),
                'placeholders': self._extract_placeholders(doc)
            }
            
            # 提取纯文本内容
            text_content = self._extract_text_content(doc)
            
            return {
                'success': True,
                'content': text_content,
                'structure': structure,
                'message': 'DOCX文件解析成功'
            }
        except Exception as e:
            return {
                'success': False,
                'message': f'解析DOCX文件失败: {str(e)}'
            }
    
    def _parse_tables(self, doc):
        """解析表格结构"""
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'table_id': f'table_{table_idx}',
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': [],
                'style': self._get_table_style(table)
            }
            
            # 遍历表格单元格
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # 检查是否合并单元格
                    is_merged = False
                    if row_idx > 0 or col_idx > 0:
                        # 简单合并检测（实际实现可能需要更复杂的逻辑）
                        if cell._tc.top != row_idx or cell._tc.left != col_idx:
                            is_merged = True
                    
                    cell_data = {
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell.text.strip(),
                        'is_merged': is_merged,
                        'format': self._get_cell_format(cell)
                    }
                    table_data['cells'].append(cell_data)
            
            tables.append(table_data)
        
        return tables
    
    def _parse_images(self, doc):
        """解析图片位置和格式"""
        images = []
        image_idx = 0
        
        for paragraph_idx, paragraph in enumerate(doc.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                for inline_shape in run._inline_shapes:
                    if inline_shape.type == 3:  # 图片类型
                        image_data = {
                            'image_id': f'image_{image_idx}',
                            'paragraph_idx': paragraph_idx,
                            'run_idx': run_idx,
                            'width': inline_shape.width.inches,
                            'height': inline_shape.height.inches,
                            'format': self._get_image_format(inline_shape)
                        }
                        images.append(image_data)
                        image_idx += 1
        
        return images
    
    def _parse_paragraphs(self, doc):
        """解析段落样式"""
        paragraphs = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # 检查段落类型
            is_title = self._is_title(paragraph)
            is_list = self._is_list(paragraph)
            
            para_data = {
                'para_id': f'para_{para_idx}',
                'text': paragraph.text.strip(),
                'is_title': is_title,
                'is_list': is_list,
                'heading_level': self._get_heading_level(paragraph),
                'format': self._get_paragraph_format(paragraph),
                'runs': self._parse_runs(paragraph)
            }
            paragraphs.append(para_data)
        
        return paragraphs
    
    def _parse_runs(self, paragraph):
        """解析段落中的文本运行（runs）"""
        runs = []
        
        for run_idx, run in enumerate(paragraph.runs):
            run_data = {
                'run_id': f'run_{run_idx}',
                'text': run.text,
                'format': self._get_run_format(run)
            }
            runs.append(run_data)
        
        return runs
    
    def _extract_placeholders(self, doc):
        """提取模板中的占位符，如{{content}}"""
        placeholders = []
        placeholder_set = set()
        
        # 从段落中提取
        for paragraph in doc.paragraphs:
            text = paragraph.text
            self._find_placeholders_in_text(text, placeholder_set)
        
        # 从表格中提取
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    self._find_placeholders_in_text(text, placeholder_set)
        
        # 转换为列表
        for placeholder in placeholder_set:
            placeholders.append({
                'name': placeholder,
                'placeholder': f'{{{{{placeholder}}}}}'
            })
        
        return placeholders
    
    def _find_placeholders_in_text(self, text, placeholder_set):
        """在文本中查找占位符"""
        import re
        # 匹配{{placeholder}}格式的占位符
        matches = re.findall(r'{{(\w+)}}', text)
        for match in matches:
            placeholder_set.add(match.strip())
    
    def _extract_text_content(self, doc):
        """提取文档的纯文本内容"""
        text_content = []
        
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        return '\n'.join(text_content)
    
    # 格式获取辅助方法
    
    def _get_table_style(self, table):
        """获取表格样式"""
        return {
            'style_name': table.style.name if table.style else 'Default'
        }
    
    def _get_cell_format(self, cell):
        """获取单元格格式"""
        # 简化实现，仅获取第一个段落的格式
        if cell.paragraphs:
            para_format = self._get_paragraph_format(cell.paragraphs[0])
            return para_format
        return {}
    
    def _get_paragraph_format(self, paragraph):
        """获取段落格式"""
        format = {
            'alignment': str(paragraph.alignment),
            'space_before': paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else 0,
            'space_after': paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0,
            'line_spacing': paragraph.paragraph_format.line_spacing
        }
        
        # 获取首行缩进
        if paragraph.paragraph_format.first_line_indent:
            format['first_line_indent'] = paragraph.paragraph_format.first_line_indent.pt
        
        return format
    
    def _get_run_format(self, run):
        """获取文本运行格式"""
        font = run.font
        format = {
            'bold': font.bold,
            'italic': font.italic,
            'underline': font.underline,
            'font_name': font.name,
            'font_size': font.size.pt if font.size else None,
            'color': self._get_color(font.color)
        }
        
        return format
    
    def _get_color(self, color):
        """获取颜色信息"""
        if color.type == 1:  # RGB颜色
            return {
                'type': 'rgb',
                'value': f'#{color.rgb[0]:02x}{color.rgb[1]:02x}{color.rgb[2]:02x}'
            }
        return {
            'type': 'default'
        }
    
    def _get_image_format(self, inline_shape):
        """获取图片格式"""
        return {
            'type': 'image',
            'shape_type': inline_shape.type
        }
    
    def _is_title(self, paragraph):
        """检查是否为标题"""
        style = paragraph.style.name
        return style.startswith('Heading') or style.startswith('标题')
    
    def _is_list(self, paragraph):
        """检查是否为列表"""
        # 简化实现，通过段落样式判断
        style = paragraph.style.name
        return 'List' in style or '列表' in style
    
    def _get_heading_level(self, paragraph):
        """获取标题级别"""
        style = paragraph.style.name
        if 'Heading' in style or '标题' in style:
            # 提取级别数字
            import re
            match = re.search(r'\d+', style)
            if match:
                return int(match.group())
        return 0
    
    def generate_template_config(self, file_path, template_name, template_type='weekly_report'):
        """
        生成完整的模板配置
        
        Args:
            file_path: DOCX文件路径
            template_name: 模板名称
            template_type: 模板类型
            
        Returns:
            dict: 完整的模板配置
        """
        parse_result = self.parse_docx(file_path)
        
        if not parse_result['success']:
            return parse_result
        
        # 生成模板配置
        template_config = {
            'id': f'template_{int(datetime.now().timestamp())}',
            'name': template_name,
            'description': f'从DOCX导入的{template_type}模板',
            'type': template_type,
            'output_format': 'docx',
            'prompt': f'请根据以下模板和提示词生成{template_type}文档，内容要详细、具体、符合实际工作情况。',
            'format_settings': {
                'font_name': '微软雅黑',
                'font_size': 12,
                'line_spacing': 1.5,
                'margin': {
                    'top': 2.54,
                    'right': 2.54,
                    'bottom': 2.54,
                    'left': 2.54
                }
            },
            'batch_settings': {
                'enabled': True,
                'delimiter': '\n',
                'max_batch_size': 10
            },
            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'is_default': False,
            'content': parse_result['content'],
            'structure': parse_result['structure']
        }
        
        return {
            'success': True,
            'template': template_config,
            'message': '模板配置生成成功'
        }
    
    def save_template_config(self, template_config, save_dir):
        """
        保存模板配置到文件
        
        Args:
            template_config: 模板配置字典
            save_dir: 保存目录
            
        Returns:
            str: 保存的文件路径
        """
        # 确保保存目录存在
        os.makedirs(save_dir, exist_ok=True)
        
        # 保存模板配置
        template_id = template_config['id']
        config_path = os.path.join(save_dir, f'{template_id}_config.json')
        
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(template_config, f, ensure_ascii=False, indent=2)
        
        return config_path
