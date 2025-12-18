import os
import json
from datetime import datetime
import time

class Config:
    def __init__(self):
        self.api_key = self._read_api_key()
        self.template = ""  # 不再在初始化时读取模板，改为动态获取
        self.reports_dir = os.path.join(os.getcwd(), "reports")
        self.config_file = os.path.join(os.getcwd(), "hos_config.json")
        os.makedirs(self.reports_dir, exist_ok=True)
        # API来源配置
        self.api_sources = {
            "deepseek": {
                "name": "DeepSeek API",
                "base_url": "https://api.deepseek.com/v1/chat/completions"
            },
            "ollama": {
                "name": "本地Ollama",
                "base_url": "http://localhost:11434/v1/chat/completions"
            }
        }
        # 加载系统设置
        self.system_settings = self._load_system_settings()
        # 加载HOS平台配置（用于自定义模块和工作流）
        self.hos_config = self._load_hos_config()
    
    def _read_api_key(self):
        """读取DeepSeek API密钥"""
        try:
            with open("key.txt", "r", encoding="utf-8") as f:
                return f.read().strip()
        except FileNotFoundError:
            # 如果key.txt文件不存在，返回空字符串
            return ""
        except Exception as e:
            raise Exception(f"读取API密钥失败: {str(e)}")
    
    def _load_system_settings(self):
        """加载系统设置"""
        system_settings_file = os.path.join(os.getcwd(), "system_settings.json")
        if os.path.exists(system_settings_file):
            with open(system_settings_file, "r", encoding="utf-8") as f:
                return json.load(f)
        return {
            "template_settings": {
                "output_before": {
                    "format": "txt",
                    "encoding": "utf-8"
                },
                "output_during": {
                    "docx": {
                        "font_name": "微软雅黑",
                        "font_size": 12,
                        "margin": {"top": 2.54, "right": 2.54, "bottom": 2.54, "left": 2.54},
                        "line_spacing": 1.5
                    },
                    "pdf": {
                        "page_size": "A4",
                        "orientation": "portrait"
                    },
                    "excel": {
                        "sheet_name": "周报"
                    }
                },
                "output_after": {
                    "default_save_location": "reports",
                    "naming_rule": "{date}_{user}_{type}.{format}",
                    "auto_save": True
                }
            },
            "api_settings": {
                "default_api_source": "deepseek",
                "request_timeout": 30,
                "max_retries": 3
            },
            "system_settings": {
                "app_name": "HOS可扩展式办公平台",
                "version": "1.0.0",
                "debug": False,
                "max_file_size": 10485760,
                "supported_file_types": ["txt", "docx", "pdf", "xls", "xlsx", "csv"]
            },
            "rag_settings": {
                "default_library_id": "",
                "auto_generate_embeddings": True,
                "default_embedding_model": "all-MiniLM-L6-v2",
                "similarity_threshold": 0.7,
                "max_results": 5,
                "enable_rag": True,
                "default_visibility": "private",
                "supported_embedding_models": ["all-MiniLM-L6-v2", "all-mpnet-base-v2", "paraphrase-MiniLM-L6-v2"],
                "summary_settings": {
                    "enabled": True,
                    "max_chars": 100,
                    "include_filename": True,
                    "include_content_preview": True
                }
            }
        }
    
    def _load_hos_config(self):
        """加载HOS平台配置"""
        default_config = {
            "workflows": [],
            "workflow_templates": self._load_workflow_templates(),
            "custom_modules": []
        }
        
        if os.path.exists(self.config_file):
            with open(self.config_file, "r", encoding="utf-8") as f:
                config = json.load(f)
                # 合并配置，确保所有默认值存在
                return self._merge_configs(default_config, config)
        return default_config
    
    def _load_workflow_templates(self):
        """加载工作流模板"""
        workflow_templates_file = os.path.join(os.getcwd(), "workflow_templates.json")
        if os.path.exists(workflow_templates_file):
            with open(workflow_templates_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("workflow_templates", [])
        return []
    
    def _load_workflow_prompt_templates(self):
        """加载工作流提示词模板"""
        workflow_prompt_templates_file = os.path.join(os.getcwd(), "workflow_prompt_templates.json")
        if os.path.exists(workflow_prompt_templates_file):
            with open(workflow_prompt_templates_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("workflow_prompt_templates", [])
        return []
    
    def _merge_configs(self, default, custom):
        """合并默认配置和自定义配置"""
        for key, value in custom.items():
            if key in default and isinstance(default[key], dict) and isinstance(value, dict):
                default[key] = self._merge_configs(default[key], value)
            else:
                default[key] = value
        return default
    
    def save_hos_config(self):
        """保存HOS平台配置"""
        with open(self.config_file, "w", encoding="utf-8") as f:
            json.dump(self.hos_config, f, ensure_ascii=False, indent=2)
    
    def update_template_settings(self, settings):
        """更新模板配置"""
        self.system_settings["template_settings"].update(settings)
        self.save_system_settings()
    
    def get_template_settings(self):
        """获取模板配置"""
        return self.system_settings["template_settings"]
    
    def save_system_settings(self):
        """保存系统设置"""
        system_settings_file = os.path.join(os.getcwd(), "system_settings.json")
        with open(system_settings_file, "w", encoding="utf-8") as f:
            json.dump(self.system_settings, f, ensure_ascii=False, indent=2)
    
    def get_workflows(self):
        """获取工作流列表"""
        return self.hos_config["workflows"]
    
    def add_workflow(self, workflow):
        """添加工作流"""
        workflow["id"] = f"workflow_{int(time.time())}"
        workflow["created_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.hos_config["workflows"].append(workflow)
        self.save_hos_config()
        return workflow
    
    def update_workflow(self, workflow_id, workflow):
        """更新工作流"""
        for i, wf in enumerate(self.hos_config["workflows"]):
            if wf["id"] == workflow_id:
                workflow["id"] = workflow_id
                workflow["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                self.hos_config["workflows"][i] = workflow
                self.save_hos_config()
                return workflow
        return None
    
    def delete_workflow(self, workflow_id):
        """删除工作流"""
        self.hos_config["workflows"] = [wf for wf in self.hos_config["workflows"] if wf["id"] != workflow_id]
        self.save_hos_config()
    
    def get_custom_modules(self):
        """获取自定义模块列表"""
        return self.hos_config["custom_modules"]
    
    def add_custom_module(self, module):
        """添加自定义模块"""
        module["id"] = f"module_{int(time.time())}"
        module["created_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.hos_config["custom_modules"].append(module)
        self.save_hos_config()
        return module
    
    def update_custom_module(self, module_id, module):
        """更新自定义模块"""
        for i, mod in enumerate(self.hos_config["custom_modules"]):
            if mod["id"] == module_id:
                module["id"] = module_id
                module["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                self.hos_config["custom_modules"][i] = module
                self.save_hos_config()
                return module
        return None
    
    def delete_custom_module(self, module_id):
        """删除自定义模块"""
        self.hos_config["custom_modules"] = [mod for mod in self.hos_config["custom_modules"] if mod["id"] != module_id]
        self.save_hos_config()
    
    def reorder_custom_modules(self, module_ids):
        """重新排序自定义模块"""
        # 创建ID到模块的映射
        module_map = {mod["id"]: mod for mod in self.hos_config["custom_modules"]}
        # 按照新的顺序重新排列模块
        reordered_modules = []
        for module_id in module_ids:
            if module_id in module_map:
                reordered_modules.append(module_map[module_id])
        # 添加未在列表中的模块
        for module in self.hos_config["custom_modules"]:
            if module["id"] not in module_ids:
                reordered_modules.append(module)
        # 更新配置
        self.hos_config["custom_modules"] = reordered_modules
        self.save_hos_config()
    
    def generate_workflow(self, prompt):
        """
        基于提示词生成工作流
        
        Args:
            prompt: 工作流需求描述
            
        Returns:
            dict: 生成的工作流信息
        """
        # 这里应该调用AI API生成工作流，但由于是演示，我们返回一个示例工作流
        # 在实际实现中，应该调用像DeepSeek这样的AI API来生成工作流
        
        # 示例工作流生成逻辑
        workflow = {
            "name": "AI生成工作流",
            "description": f"基于提示词生成的工作流: {prompt}",
            "steps": [
                {
                    "id": "step_1",
                    "name": "数据收集",
                    "action": "collect_data",
                    "params": {
                        "sources": ["系统1", "系统2"],
                        "data_types": ["结构化数据", "非结构化数据"]
                    },
                    "next_step": "step_2"
                },
                {
                    "id": "step_2",
                    "name": "数据处理",
                    "action": "process_data",
                    "params": {
                        "operations": ["清洗", "转换", "整合"]
                    },
                    "next_step": "step_3"
                },
                {
                    "id": "step_3",
                    "name": "AI分析",
                    "action": "ai_analysis",
                    "params": {
                        "model": "default",
                        "prompt": prompt
                    },
                    "next_step": "step_4"
                },
                {
                    "id": "step_4",
                    "name": "生成报告",
                    "action": "generate_report",
                    "params": {
                        "format": "docx",
                        "template": "default"
                    },
                    "next_step": "step_5"
                },
                {
                    "id": "step_5",
                    "name": "报告分发",
                    "action": "distribute_report",
                    "params": {
                        "recipients": ["经理", "团队成员"],
                        "channels": ["邮件", "系统通知"]
                    }
                }
            ]
        }
        
        return workflow
    
    def convert_workflow_to_module(self, workflow_id):
        """
        将工作流转换为自定义模块
        
        Args:
            workflow_id: 工作流ID
            
        Returns:
            dict: 转换后的模块信息
        """
        # 获取工作流
        workflow = next((wf for wf in self.hos_config["workflows"] if wf["id"] == workflow_id), None)
        if not workflow:
            return None
        
        # 创建自定义模块
        module = {
            "name": workflow["name"],
            "description": workflow.get("description", "从工作流转换的模块"),
            "type": "workflow_module",
            "workflow_id": workflow_id,
            "workflow_config": {
                "steps": workflow.get("steps", []),
                "settings": workflow.get("settings", {})
            }
        }
        
        # 添加到自定义模块列表
        return self.add_custom_module(module)

class ReportGenerator:
    def __init__(self, config, api_client):
        self.config = config
        self.api_client = api_client
        # 不适当内容过滤列表
        self.inappropriate_content = [
            '色情', '暴力', '赌博', '毒品', '恐怖', '极端',
            '政治敏感', '宗教敏感', '民族敏感', '地域敏感',
            '广告', '推销', '诈骗', '虚假信息', '谣言',
            '侮辱', '诽谤', '歧视', '攻击', '威胁'
        ]
        # 批次存储文件路径
        self.batches_file = os.path.join(os.getcwd(), "batches.json")
        # 加载现有批次数据
        self.batches = self._load_batches()
    
    def _load_batches(self):
        """加载批次数据"""
        try:
            if os.path.exists(self.batches_file):
                with open(self.batches_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            return []
        except Exception as e:
            print(f"加载批次数据失败: {str(e)}")
            return []
    
    def _save_batches(self):
        """保存批次数据"""
        try:
            with open(self.batches_file, "w", encoding="utf-8") as f:
                json.dump(self.batches, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存批次数据失败: {str(e)}")
    
    def create_batch(self, batch_name, files):
        """创建新批次"""
        batch = {
            "id": f"batch_{int(time.time())}",
            "name": batch_name,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "files": files,
            "total": len(files)
        }
        self.batches.append(batch)
        self._save_batches()
        return batch
    
    def get_batches(self):
        """获取所有批次"""
        return self.batches
    
    def get_batch(self, batch_id):
        """获取特定批次"""
        for batch in self.batches:
            if batch["id"] == batch_id:
                return batch
        return None
    
    def add_file_to_batch(self, batch_id, filename):
        """添加文件到批次"""
        batch = self.get_batch(batch_id)
        if batch:
            if filename not in batch["files"]:
                batch["files"].append(filename)
                batch["total"] = len(batch["files"])
                self._save_batches()
            return True
        return False
    
    def remove_file_from_batch(self, batch_id, filename):
        """从批次中移除文件"""
        batch = self.get_batch(batch_id)
        if batch:
            if filename in batch["files"]:
                batch["files"].remove(filename)
                batch["total"] = len(batch["files"])
                self._save_batches()
            return True
        return False
    
    def delete_batch(self, batch_id):
        """删除批次"""
        self.batches = [batch for batch in self.batches if batch["id"] != batch_id]
        self._save_batches()
    
    def get_files_by_batch(self, batch_id):
        """获取批次中的文件列表"""
        batch = self.get_batch(batch_id)
        if batch:
            return batch["files"]
        return []
    
    def filter_content(self, content):
        """过滤不适当内容"""
        filtered_content = content
        
        # 检查是否包含不适当内容
        for word in self.inappropriate_content:
            if word in filtered_content:
                # 替换不适当内容
                filtered_content = filtered_content.replace(word, '*' * len(word))
        
        return filtered_content
    
    def generate_single_report(self, prompt, template_id=None, progress_callback=None, rag_library_id=None):
        """生成单个周报，支持进度回调和RAG库"""
        if progress_callback:
            progress_callback(0, 100, 0, '准备生成周报...')
        
        # 导入模板管理器
        from HOS_ME.utils.template_manager import TemplateManager
        template_manager = TemplateManager()
        
        if progress_callback:
            progress_callback(20, 100, 20, '加载模板...')
        
        # 获取模板内容
        if template_id:
            # 使用指定模板
            template_info = template_manager.get_template(template_id)
            if template_info:
                template_content = template_info['content']
                template_prompt = template_info.get('prompt', '请根据以下模板和提示词生成专业的文档，内容要详细、具体、符合实际工作情况。')
            else:
                # 模板不存在，使用默认模板
                default_template = template_manager.get_current_template()
                template_content = default_template['content']
                template_prompt = '请根据以下模板和提示词生成周报：'
        else:
            # 使用默认模板
            default_template = template_manager.get_current_template()
            template_content = default_template['content']
            template_prompt = '请根据以下模板和提示词生成周报：'
        
        # RAG增强：如果提供了RAG库ID，使用RAG库增强提示词
        enhanced_prompt = prompt
        if rag_library_id:
            if progress_callback:
                progress_callback(30, 100, 30, '正在使用RAG库增强内容...')
            
            # 导入知识库
            from HOS_ME.utils.knowledge_base import KnowledgeBase
            knowledge_base = KnowledgeBase(self.config)  # 传递config参数
            
            # 使用RAG库查询相关内容
            rag_results = knowledge_base.rag_query(prompt, rag_library_id, top_k=3)
            
            if rag_results:
                # 构建RAG增强的提示词，使用文档的总结
                rag_context = "\n".join([f"相关参考 {i+1}: {result['document']['summary']}" for i, result in enumerate(rag_results)])
                enhanced_prompt = f"{prompt}\n\n相关参考资料：\n{rag_context}\n\n请根据以上参考资料和您的知识生成内容。"
        
        if progress_callback:
            progress_callback(40, 100, 40, '准备API请求...')
        
        full_prompt = f"{template_prompt}\n\n模板：\n{template_content}\n\n提示词：\n{enhanced_prompt}\n\n请生成符合模板格式的文档，内容要详细、专业，符合实际工作情况。"
        
        if progress_callback:
            progress_callback(60, 100, 60, '调用API生成周报...')
        
        report_content = self.api_client.generate_report(full_prompt)
        
        # 过滤不适当内容
        report_content = self.filter_content(report_content)
        
        if progress_callback:
            progress_callback(100, 100, 100, '生成完成')
        
        return report_content
    
    def generate_batch_reports(self, prompts, template_id=None, progress_callback=None, file_format="txt", rag_library_id=None):
        """批量生成周报，支持进度回调和RAG库"""
        reports = []
        saved_files = []
        total = len(prompts)
        
        # 过滤空提示词
        valid_prompts = [(i+1, prompt) for i, prompt in enumerate(prompts) if prompt.strip()]
        total = len(valid_prompts)
        
        if total == 0:
            if progress_callback:
                progress_callback(0, 0, 100, "没有可生成的提示词")
            return reports
        
        for i, (original_index, prompt) in enumerate(valid_prompts):
            try:
                # 发送进度更新
                if progress_callback:
                    progress = (i + 1) / total * 100
                    progress_callback(i + 1, total, progress, f"正在生成第 {i + 1} 份周报...")
                
                # 使用generate_single_report方法，支持RAG库
                report = self.generate_single_report(prompt, template_id, None, rag_library_id)
                
                # 保存报告到文件系统，使用指定的文件格式
                filename = f"batch_{i+1}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}_周报.{file_format}"
                self.save_report(report, filename, file_format)
                saved_files.append(filename)
                
                reports.append((original_index, prompt, report, "成功", filename))
            except Exception as e:
                reports.append((original_index, prompt, "", f"失败: {str(e)}", None))
        
        # 创建批次记录
        if saved_files:
            batch_name = f"批次_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"
            self.create_batch(batch_name, saved_files)
        
        # 发送完成通知
        if progress_callback:
            progress_callback(total, total, 100, f"批量生成完成，共生成 {len(saved_files)} 个文件")
        
        return reports
    
    def save_report(self, report_content, filename=None, file_format="txt"):
        """保存周报，支持多种格式"""
        if not filename:
            filename = f"{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}_周报.{file_format}"
        
        filepath = os.path.join(self.config.reports_dir, filename)
        
        if file_format == "txt":
            # 保存为文本文件
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(report_content)
        elif file_format == "docx":
            # 保存为docx文件，使用HOS_M2F包处理Markdown
            try:
                # 尝试使用HOS_M2F包处理
                from hos_m2f import MarkdownToDocx
                
                # 创建MarkdownToDocx实例
                m2d = MarkdownToDocx()
                
                # 转换并保存
                m2d.convert(report_content, filepath)
            except ImportError:
                # 如果HOS_M2F包不可用，使用默认处理方式
                from docx import Document
                from docx.shared import Pt
                from docx.oxml.ns import qn
                
                # 创建文档
                doc = Document()
                
                # 设置默认字体为微软雅黑
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 添加标题
                title = doc.add_heading('周报', 0)
                title.alignment = 1  # 居中对齐
                for run in title.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(16)
                
                # 处理内容，按换行符分割段落
                paragraphs = report_content.split('\n')
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue
                    
                    # 处理标题
                    if para.startswith('# '):
                        heading_level = 1
                        content = para[2:]
                    elif para.startswith('## '):
                        heading_level = 2
                        content = para[3:]
                    elif para.startswith('### '):
                        heading_level = 3
                        content = para[4:]
                    else:
                        # 普通段落
                        p = doc.add_paragraph()
                        run = p.add_run(para)
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        run.font.size = Pt(12)
                        continue
                    
                    # 添加标题
                    heading = doc.add_heading(content, heading_level)
                    for run in heading.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        if heading_level == 1:
                            run.font.size = Pt(14)
                        elif heading_level == 2:
                            run.font.size = Pt(13)
                        else:
                            run.font.size = Pt(12)
                
                # 保存文档
                doc.save(filepath)
        else:
            # 其他格式暂不支持
            raise Exception(f"不支持的文件格式: {file_format}")
        
        return filename
    
    def load_reports(self):
        """加载历史周报列表，支持所有格式"""
        reports = []
        if not os.path.exists(self.config.reports_dir):
            return reports
        
        supported_formats = [".txt", ".docx"]
        for filename in os.listdir(self.config.reports_dir):
            # 检查文件扩展名是否支持
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in supported_formats:
                filepath = os.path.join(self.config.reports_dir, filename)
                reports.append({
                    "filename": filename,
                    "filepath": filepath,
                    "date": datetime.fromtimestamp(os.path.getctime(filepath)).strftime("%Y-%m-%d %H:%M:%S")
                })
        
        # 按创建时间倒序排列
        reports.sort(key=lambda x: x["date"], reverse=True)
        return reports
    
    def read_report(self, filename):
        """读取周报内容，支持多种格式"""
        filepath = os.path.join(self.config.reports_dir, filename)
        if not os.path.exists(filepath):
            raise Exception("周报不存在")
        
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == ".txt":
            # 读取txt文件
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        elif file_ext == ".docx":
            # 读取docx文件
            try:
                from docx import Document
                doc = Document(filepath)
                content = []
                for para in doc.paragraphs:
                    content.append(para.text)
                return "\n".join(content)
            except Exception as e:
                return f"无法读取docx文件: {str(e)}"
        else:
            # 不支持的格式
            return f"不支持读取{file_ext}格式文件"

    
    def delete_report(self, filename):
        """删除单个周报，并从所有批次中移除该文件"""
        filepath = os.path.join(self.config.reports_dir, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
            # 从所有批次中移除该文件
            for batch in self.batches:
                if filename in batch["files"]:
                    batch["files"].remove(filename)
                    batch["total"] = len(batch["files"])
            # 保存批次更新
            self._save_batches()
            return True
        return False
    
    def batch_delete_reports(self, filenames, progress_callback=None):
        """批量删除周报，支持进度回调"""
        success_count = 0
        failed_count = 0
        failed_files = []
        total = len(filenames)
        
        if progress_callback:
            progress_callback(0, total, 0, f'准备删除 {total} 个文件...')
        
        for index, filename in enumerate(filenames):
            try:
                if progress_callback:
                    current_progress = (index / total) * 100
                    progress_callback(index + 1, total, current_progress, f'正在删除第 {index + 1} 个文件...')
                
                if self.delete_report(filename):
                    success_count += 1
                else:
                    failed_count += 1
                    failed_files.append(filename)
            except Exception as e:
                failed_count += 1
                failed_files.append(filename)
                continue
        
        if progress_callback:
            progress_callback(total, total, 100, 
                            f'批量删除完成，成功删除 {success_count} 个文件，失败 {failed_count} 个文件')

        return {
            "success": success_count,
            "failed": failed_count,
            "failed_files": failed_files
        }

    def import_excel(self, file, progress_callback=None):
        """从Excel文件导入数据并生成文档，支持进度回调"""
        generated = 0
        failed = 0
        failed_records = []

        try:
            # 导入pandas库
            import pandas as pd
            import io

            if progress_callback:
                progress_callback(0, 100, 0, '正在读取Excel文件...')

            # 读取Excel文件
            df = pd.read_excel(io.BytesIO(file.read()))

            # 检查必要的列
            if 'prompt' not in df.columns:
                raise Exception('Excel文件必须包含prompt列')

            total_rows = len(df)

            if progress_callback:
                progress_callback(20, 100, 20, 
                                f'开始生成文档，共 {total_rows} 行数据...')

            # 遍历数据行
            for index, row in df.iterrows():
                try:
                    prompt = row['prompt']
                    if pd.isna(prompt) or not str(prompt).strip():
                        failed += 1
                        failed_records.append(f'第{index+2}行：提示词为空')
                        if progress_callback:
                            current_progress = 20 + ((index + 1) / total_rows) * 70
                            progress_callback(index + 1, total_rows, current_progress, 
                                            f'跳过空提示词，已处理 {index + 1}/{total_rows} 行...')
                        continue

                    # 获取其他可选列
                    template_id = row.get('template_id')
                    file_format = row.get('file_format', 'txt')
                    filename = row.get('filename')

                    if progress_callback:
                        current_progress = 20 + (index / total_rows) * 70
                        progress_callback(index + 1, total_rows, current_progress, 
                                        f'正在生成第 {index + 1} 个文档...')

                    # 生成文档，添加进度回调
                    report_content = self.generate_single_report(
                        str(prompt), template_id, progress_callback=None
                    )  # 内部已处理进度

                    # 保存文档
                    self.save_report(report_content, filename, file_format)
                    generated += 1

                    if progress_callback:
                        current_progress = 20 + ((index + 1) / total_rows) * 70
                        progress_callback(index + 1, total_rows, current_progress, 
                                        f'已完成 {index + 1}/{total_rows} 个文档...')
                except Exception as e:
                    failed += 1
                    failed_records.append(f'第{index+2}行：{str(e)}')
                    if progress_callback:
                        current_progress = 20 + ((index + 1) / total_rows) * 70
                        progress_callback(index + 1, total_rows, current_progress, 
                                        f'生成失败，已处理 {index + 1}/{total_rows} 行...')
                    continue

            if progress_callback:
                progress_callback(total_rows, total_rows, 100, 
                                f'Excel导入完成，共生成 {generated} 个文档，失败 {failed} 个')
        except Exception as e:
            if progress_callback:
                progress_callback(0, 100, 0, f'Excel导入失败: {str(e)}')
            raise Exception(f'Excel导入失败: {str(e)}')

        return {
            "generated": generated,
            "failed": failed,
            "failed_records": failed_records
        }
