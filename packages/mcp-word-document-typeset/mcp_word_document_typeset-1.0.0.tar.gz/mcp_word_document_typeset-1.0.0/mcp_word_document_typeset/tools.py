"""
Word文档排版工具函数

仅保留一个核心功能：为指定级别标题（Heading 1-9）统一设置样式。
"""

import os
import re
from typing import Optional
from .utils import (
    ensure_docx_extension,
    check_file_writeable,
)


async def style_heading(
    filename: str,
    level: int,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    color: Optional[str] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    align: Optional[str] = None,
    spacing_before: Optional[int] = None,
    spacing_after: Optional[int] = None,
    add_numbering: Optional[bool] = None,
    numbering_separator: Optional[str] = None,
    numbering_suffix: Optional[str] = None,
    remove_numbering: Optional[bool] = None,
    remove_all_levels: Optional[bool] = None,
) -> str:
    """为指定级别标题样式（Heading 1-9）应用格式。

    说明：此操作会直接修改文档内对应的样式对象（如 "Heading 1"），从而使所有使用该样式的段落统一更新。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    filename = ensure_docx_extension(filename)

    # 校验参数类型
    try:
        level = int(level)
        if font_size is not None:
            font_size = int(font_size)
        if spacing_before is not None:
            spacing_before = int(spacing_before)
        if spacing_after is not None:
            spacing_after = int(spacing_after)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: level/font_size/spacing_before/spacing_after must be integers")

    if level < 1 or level > 9:
        raise ValueError("Heading level must be between 1 and 9")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # 可写性检查
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    # 颜色映射与解析
    def parse_color(value: str):
        color_map = {
            'red': RGBColor(255, 0, 0),
            'blue': RGBColor(0, 0, 255),
            'green': RGBColor(0, 128, 0),
            'black': RGBColor(0, 0, 0),
            'white': RGBColor(255, 255, 255),
            'yellow': RGBColor(255, 255, 0),
            'orange': RGBColor(255, 165, 0),
            'purple': RGBColor(128, 0, 128),
            'gray': RGBColor(128, 128, 128),
            'grey': RGBColor(128, 128, 128),
        }
        if not value:
            return None
        try:
            v = value.strip()
            if v.lower() in color_map:
                return color_map[v.lower()]
            # hex like #RRGGBB or RRGGBB
            if v.startswith('#'):
                v = v[1:]
            if len(v) == 6:
                r = int(v[0:2], 16)
                g = int(v[2:4], 16)
                b = int(v[4:6], 16)
                return RGBColor(r, g, b)
        except Exception:
            pass
        return None

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        style_name = f"Heading {level}"
        heading_style = None
        
        # 按名称查找样式（更可靠，因为有些文档的 style_id 不标准）
        try:
            heading_style = doc.styles[style_name]
        except KeyError:
            # 尝试中文名称
            try:
                heading_style = doc.styles[f"标题 {level}"]
            except KeyError:
                # 最后尝试按标准 ID 查找
                target_style_id = f"Heading{level}"
                for s in doc.styles:
                    if getattr(s, 'style_id', '') == target_style_id:
                        heading_style = s
                        break
                
                if heading_style is None:
                    raise ValueError(f"Style '{style_name}' (or '标题 {level}') not found in document.")

        # 获取实际使用的 style_name (用于日志或错误提示) 和 style_id (用于匹配段落)
        actual_style_name = heading_style.name
        actual_style_id = heading_style.style_id
        
        font = heading_style.font
        pf = heading_style.paragraph_format

        if bold is not None:
            font.bold = bold
        if italic is not None:
            font.italic = italic
        if underline is not None:
            font.underline = underline
        if font_size is not None and font_size > 0:
            font.size = Pt(font_size)
        if font_name:
            font.name = font_name
            # 针对中文字体，必须设置 eastAsia 属性
            try:
                if heading_style.element.rPr is not None:
                    if heading_style.element.rPr.rFonts is not None:
                        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            except Exception:
                pass
        if color:
            rgb = parse_color(color)
            if rgb is not None:
                font.color.rgb = rgb

        if align:
            al = align.strip().lower()
            if al == 'left':
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif al == 'center':
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif al == 'right':
                pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif al == 'justify':
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                raise ValueError("Invalid align value. Use left/center/right/justify")

        if spacing_before is not None and spacing_before >= 0:
            pf.space_before = Pt(spacing_before)
        if spacing_after is not None and spacing_after >= 0:
            pf.space_after = Pt(spacing_after)

        # 遍历文档段落，清除直接格式（direct formatting）以确保样式生效
        # 仅针对用户本次明确请求设置的属性进行清除，避免误删用户其他手动格式
        modified_count = 0
        for para in doc.paragraphs:
            # 检查段落样式是否匹配（注意：para.style 可能为 None 或其他类型，需安全访问）
            # 改为匹配 style_id，以支持中文 "标题 1" 等情况
            try:
                if not para.style:
                    continue
                # 优先按名称匹配（更可靠），因为有些文档的 style_id 不标准
                if para.style.name != actual_style_name:
                    # 如果名称不匹配，再尝试 ID 匹配（兼容性）
                    p_style_id = getattr(para.style, 'style_id', None)
                    if not (p_style_id and actual_style_id and p_style_id == actual_style_id):
                        continue
            except Exception:
                continue
            
            modified_count += 1
            for run in para.runs:
                # 1. 颜色
                if color:
                    # 尝试1：清除 run 级别的 RGB
                    try:
                        run.font.color.rgb = None
                    except Exception:
                        pass
                    
                    # 尝试2：从 XML 中彻底移除所有颜色相关节点
                    try:
                        if run.element.rPr is not None:
                            # 移除所有可能影响颜色的元素
                            # textFill 是关键！它的优先级高于 w:color
                            for tag_name in ['w:color', 'w:shd', 'w:highlight', 'w14:textFill', 'w14:textOutline']:
                                color_element = run.element.rPr.find(qn(tag_name))
                                if color_element is not None:
                                    run.element.rPr.remove(color_element)
                    except Exception:
                        pass
                    
                    # 尝试3：清除字符样式 (Character Style) 引用，防止样式覆盖
                    # 注意：如果原本有加粗/斜体等其他字符样式，这会将其移除。
                    # 但为了解决“不变色”的痛点，这是必要的妥协。
                    try:
                         if run.style and run.style.name != "Default Paragraph Font":
                             run.style = "Default Paragraph Font"
                    except Exception:
                        pass

                    # 尝试4：如果以上清除操作后，理论上应该继承段落样式。
                    # 但为了双重保险，如果清理后还没变色（或为了确保绝对生效），
                    # 我们可以选择强制在 Run 上应用颜色。
                    # 策略：如果用户指定了颜色，强制应用到 Run 上，确保视觉效果。
                    # 弊端：以后改样式不会自动更新。但解决当前“无效”问题优先。
                    try:
                        rgb_val = parse_color(color)
                        if rgb_val is not None:
                            from docx.oxml import OxmlElement
                            rPr = run.element.rPr
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                run.element.insert(0, rPr)
                            
                            # 创建新的颜色元素并添加
                            color_elem = OxmlElement('w:color')
                            color_elem.set(qn('w:val'), '%02X%02X%02X' % (rgb_val.r, rgb_val.g, rgb_val.b))
                            rPr.append(color_elem)
                    except Exception:
                        # 回退方案
                        try:
                            rgb_val = parse_color(color)
                            if rgb_val is not None:
                                run.font.color.rgb = rgb_val
                        except:
                            pass
                
                # 2. 字号
                if font_size is not None:
                    run.font.size = None
                    # 强制设置（双保险）
                    try:
                        run.font.size = Pt(font_size)
                    except Exception:
                        pass
                
                # 3. 字体名称
                if font_name:
                    run.font.name = None
                    try:
                        run.font.name = font_name
                        # 针对中文字体，必须设置 eastAsia 属性
                        if run.element.rPr is not None:
                             if run.element.rPr.rFonts is not None:
                                 run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    except Exception:
                        pass
                
                # 4. 粗体
                if bold is not None:
                    run.font.bold = None
                    try:
                        run.font.bold = bold
                    except Exception:
                        pass
                
                # 5. 斜体
                if italic is not None:
                    run.font.italic = None
                    try:
                        run.font.italic = italic
                    except Exception:
                        pass
                
                # 6. 下划线
                if underline is not None:
                    run.font.underline = None
                    try:
                        run.font.underline = underline
                    except Exception:
                        pass

        # 仅移除编号（优先级最高）
        if remove_numbering:
            # 支持常见模式："1 ", "1.", "1. ", "1)", "1- ", 多级如 "1.2.3 " 等
            try:
                _sep = numbering_separator if (isinstance(numbering_separator, str) and numbering_separator) else "."
                import re as _re
                # 开头可能有空白；数字+（.或-或)）可重复；可带自定义分隔符；最后跟一个可选的 .|)|- 再跟空格
                _pat = _re.compile(r"^\s*\d+(?:" + _re.escape(_sep) + r"\d+)*(?:[\.|\)|\-])?\s+")
            except Exception:
                _pat = None

            for para in doc.paragraphs:
                try:
                    style_n = para.style.name if para.style else ""
                except Exception:
                    style_n = ""
                if style_n.startswith("Heading "):
                    try:
                        lvl = int(style_n.split()[-1])
                    except Exception:
                        continue
                    # 若 remove_all_levels 为真，则所有级别都移除；否则仅移除目标 level
                    if remove_all_levels or lvl == level:
                        if _pat is not None:
                            para.text = _pat.sub("", para.text or "")
            # 若仅移除，不再添加编号
        # 自动编号（多级）
        elif add_numbering:
            sep = numbering_separator if (isinstance(numbering_separator, str) and numbering_separator) else "."
            suffix = numbering_suffix if (isinstance(numbering_suffix, str)) else " "
            counters = [0] * 9
            # 识别并去除已有编号前缀，避免重复编号
            num_pattern = None
            try:
                import re as _re
                num_pattern = _re.compile(r"^\s*\d+(?:" + _re.escape(sep) + r"\d+)*(?:[\.|\)|\-])?\s+")
            except Exception:
                num_pattern = None

            for para in doc.paragraphs:
                try:
                    style_n = para.style.name if para.style else ""
                except Exception:
                    style_n = ""
                if style_n.startswith("Heading "):
                    try:
                        lvl = int(style_n.split()[-1])
                    except Exception:
                        continue
                    if 1 <= lvl <= 9:
                        # 更新计数器以保持层级关系（即便不改该级文本）
                        counters[lvl - 1] += 1
                        for i in range(lvl, 9):
                            counters[i] = 0
                        # 仅当当前段落级别等于目标 level 时，才写入编号前缀
                        if lvl == level:
                            parts = [str(counters[i]) for i in range(lvl) if counters[i] > 0]
                            num_str = sep.join(parts)
                            base_text = para.text or ""
                            if num_pattern is not None:
                                base_text = num_pattern.sub("", base_text)
                            new_text = f"{num_str}{suffix}{base_text}".strip()
                            para.text = new_text

        doc.save(filename)
        return f"Styled '{actual_style_name}' successfully. (Updated {modified_count} paragraphs)"
    except Exception as e:
        raise RuntimeError(f"Failed to style heading: {str(e)}")
