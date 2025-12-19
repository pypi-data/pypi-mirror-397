from __future__ import annotations

import io


def _heading_prefix(style_name: str) -> str | None:
    s = (style_name or "").strip().lower()
    if not s.startswith("heading"):
        return None
    parts = s.split()
    if len(parts) >= 2 and parts[-1].isdigit():
        level = max(1, min(6, int(parts[-1])))
        return "#" * level + " "
    return "# "


def convert_docx_bytes(data: bytes) -> dict:
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        raise RuntimeError("DOCX conversion requires `python-docx`. Install with `pip install 'kytchen[converters]'`.") from e

    doc = Document(io.BytesIO(data))

    out_lines: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").rstrip()
        if not text:
            continue
        prefix = _heading_prefix(getattr(p.style, "name", "") or "")
        if prefix:
            out_lines.append(prefix + text)
        else:
            out_lines.append(text)

    table_count = 0
    for t in doc.tables:
        table_count += 1
        out_lines.append("")
        out_lines.append(f"[Table {table_count}]")
        for row in t.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            out_lines.append("\t".join(cells))

    text = "\n".join(out_lines).strip()

    return {
        "type": "docx",
        "paragraphs": len(doc.paragraphs),
        "tables": table_count,
        "text": text,
    }
