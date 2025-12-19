"""
DOCX Generator
Converts Markdown files to properly formatted Word documents.
"""
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Document = None

from seokit.config import OUTPUTS_DIR


def setup_document_styles(doc):
    """Configure document styles for consistent formatting."""
    # Normal text style
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(24)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def parse_inline_formatting(paragraph, text: str):
    """Parse and apply inline markdown formatting (bold, italic, links)."""
    # Simple approach: handle bold first, then regular text
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Handle links - extract just the text
            link_pattern = r'\[([^\]]+)\]\([^)]+\)'
            link_free = re.sub(link_pattern, r'\1', part)
            paragraph.add_run(link_free)


def md_to_docx(md_path: str, output_path: str = None) -> str:
    """
    Convert markdown file to DOCX format.

    Args:
        md_path: Path to markdown file
        output_path: Optional custom output path

    Returns:
        Path to generated DOCX file
    """
    if Document is None:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    md_path = Path(md_path)

    if not md_path.exists():
        # Try in outputs directory
        md_path = OUTPUTS_DIR / md_path.name
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {md_path}")

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    # Create document
    doc = Document()
    setup_document_styles(doc)

    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block - add as formatted paragraph
                if code_content:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip metadata blocks (> lines at start)
        if stripped.startswith('>'):
            i += 1
            continue

        # Skip horizontal rules
        if stripped in ['---', '***', '___']:
            i += 1
            continue

        # H1 - Title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading = doc.add_heading(stripped[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # H2 - Main sections
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        # H3 - Subsections
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        # H4 - Sub-subsections
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)
        i += 1

    # Determine output path
    if output_path:
        docx_path = Path(output_path)
    else:
        docx_path = md_path.with_suffix('.docx')

    # Ensure output directory exists
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    # Save document
    doc.save(str(docx_path))

    return str(docx_path)


def count_words_in_docx(docx_path: str) -> int:
    """Count words in a DOCX file."""
    if Document is None:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = Document(docx_path)
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    return word_count


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python -m seokit.core.docx_generator <markdown_file> [output_file]")
        print("Example: python -m seokit.core.docx_generator outputs/article-seo-tips.md")
        sys.exit(1)

    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        result = md_to_docx(md_file, output_file)
        word_count = count_words_in_docx(result)
        print(f"DOCX generated successfully: {result}")
        print(f"Word count: {word_count}")
    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        sys.exit(1)
