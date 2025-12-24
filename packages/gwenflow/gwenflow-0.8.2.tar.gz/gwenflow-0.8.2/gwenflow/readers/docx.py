import io
from typing import List, Union, ClassVar
from pathlib import Path

from gwenflow.logger import logger
from gwenflow.types import Document
from gwenflow.readers.base import Reader



class DocxReader(Reader):
    trans: ClassVar[dict[int, int | None]] = {
        0x00A0: 0x20,
        0x202F: 0x20,
        0x2007: 0x20,
        0x200B: None,
        0x200C: None,
        0x200D: None,
        0xFEFF: None,
    }

    def getText(self, file_obj) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is not installed. Please install it with `pip install python-docx`")
        doc = docx.Document(file_obj)
        return "\n".join((p.text.translate(self.trans) if p.text else "") for p in doc.paragraphs)

    def getTables(self, file_obj):
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is not installed. Please install it with `pip install python-docx`")
        doc = docx.Document(file_obj)

        tables = []
        for t in doc.tables:
            rows = []
            for r in t.rows:
                cells = []
                for c in r.cells:
                    txt = "\n".join(p.text for p in c.paragraphs) if c.paragraphs else ""
                    txt = txt.translate(self.trans) if txt else ""
                    cells.append(txt)
                rows.append(cells)
            tables.append(rows)
        return tables

    def read(self, file: Union[Path, io.BytesIO]) -> List[Document]:
        try:
            filename = self.get_file_name(file)
            content = self.get_file_content(file)

            data = content.getvalue() if isinstance(content, io.BytesIO) else content

            text = self.getText(io.BytesIO(data))
            tables = self.getTables(io.BytesIO(data))

            doc = Document(
                id=self.key(f"{filename}"),
                content=text,
                metadata={"filename": filename, "tables": tables},
            )
            return [doc]

        except Exception as e:
            logger.exception(f"Error reading file: {e}")
            return []