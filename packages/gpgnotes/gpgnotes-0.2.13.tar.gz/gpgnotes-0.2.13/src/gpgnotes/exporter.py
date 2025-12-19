"""File exporter for GPGNotes - exports notes to various formats."""

import re
from pathlib import Path

from .note import Note


class ExportError(Exception):
    """Raised when file export fails."""

    pass


class MissingDependencyError(ExportError):
    """Raised when required dependency is not installed."""

    pass


def _check_dependency(module_name: str, package_name: str):
    """Check if a dependency is available, raise helpful error if not."""
    try:
        __import__(module_name)
    except ModuleNotFoundError:
        raise MissingDependencyError(
            f"The '{package_name}' package is required to export to this format.\n"
            f"Install it with: pip install gpgnotes[import]\n"
            f"Or: pip install {package_name}"
        )


def export_markdown(note: Note) -> str:
    """
    Export note to markdown format.

    Args:
        note: Note to export

    Returns:
        Markdown string
    """
    content = f"# {note.title}\n\n"
    content += f"**Modified:** {note.modified.strftime('%Y-%m-%d %H:%M')}\n"
    content += f"**Tags:** {', '.join(note.tags) if note.tags else 'none'}\n\n"
    content += f"{note.content}\n"
    return content


def export_text(note: Note) -> str:
    """
    Export note to plain text format.

    Args:
        note: Note to export

    Returns:
        Plain text string
    """
    content = f"{note.title}\n"
    content += "=" * len(note.title) + "\n\n"
    content += f"Modified: {note.modified.strftime('%Y-%m-%d %H:%M')}\n"
    content += f"Tags: {', '.join(note.tags) if note.tags else 'none'}\n\n"
    content += f"{note.content}\n"
    return content


def export_html(note: Note) -> str:
    """
    Export note to HTML format.

    Args:
        note: Note to export

    Returns:
        HTML string
    """
    content = "<!DOCTYPE html>\n<html>\n<head>\n"
    content += "<meta charset='utf-8'>\n"
    content += f"<title>{note.title}</title>\n"
    content += "<style>body { font-family: Arial, sans-serif; max-width: 800px; "
    content += "margin: 40px auto; padding: 20px; }"
    content += "h1 { color: #333; } .meta { color: #666; font-size: 0.9em; } "
    content += ".tags { color: #0066cc; } pre { background: #f5f5f5; padding: 15px; "
    content += "overflow-x: auto; }</style>\n"
    content += "</head>\n<body>\n"
    content += f"<h1>{note.title}</h1>\n"
    content += f"<div class='meta'>Modified: {note.modified.strftime('%Y-%m-%d %H:%M')}</div>\n"
    content += f"<div class='tags'>Tags: {', '.join(note.tags) if note.tags else 'none'}</div>\n"
    content += f"<pre>{note.content}</pre>\n"
    content += "</body>\n</html>"
    return content


def export_json(note: Note) -> str:
    """
    Export note to JSON format.

    Args:
        note: Note to export

    Returns:
        JSON string
    """
    import json

    data = {
        "title": note.title,
        "content": note.content,
        "tags": note.tags,
        "created": note.created.isoformat() if hasattr(note, "created") else None,
        "modified": note.modified.isoformat(),
    }
    return json.dumps(data, indent=2, ensure_ascii=False)


def export_rtf(note: Note) -> str:
    """
    Export note to RTF format.

    Args:
        note: Note to export

    Returns:
        RTF string
    """
    # RTF header
    rtf = r"{\rtf1\ansi\deff0"
    rtf += r"{\fonttbl{\f0 Arial;}}"
    rtf += r"{\colortbl;\red0\green0\blue0;\red100\green100\blue100;}"
    rtf += "\n"

    # Title (bold, larger)
    rtf += r"{\pard\b\fs32 " + _escape_rtf(note.title) + r"\b0\par}"
    rtf += "\n"

    # Metadata
    rtf += r"{\pard\cf2\fs20 Modified: " + note.modified.strftime("%Y-%m-%d %H:%M") + r"\par}"
    tags_str = ", ".join(note.tags) if note.tags else "none"
    rtf += r"{\pard\cf2\fs20 Tags: " + _escape_rtf(tags_str) + r"\par}"
    rtf += r"{\pard\par}"  # Empty line
    rtf += "\n"

    # Content
    for line in note.content.split("\n"):
        rtf += r"{\pard " + _escape_rtf(line) + r"\par}"
        rtf += "\n"

    rtf += "}"
    return rtf


def _escape_rtf(text: str) -> str:
    """Escape special RTF characters."""
    text = text.replace("\\", "\\\\")
    text = text.replace("{", "\\{")
    text = text.replace("}", "\\}")
    # Handle unicode characters
    result = []
    for char in text:
        if ord(char) > 127:
            result.append(f"\\u{ord(char)}?")
        else:
            result.append(char)
    return "".join(result)


def export_pdf(note: Note, output_path: Path) -> None:
    """
    Export note to PDF format.

    Args:
        note: Note to export
        output_path: Path to save the PDF

    Note:
        This function writes directly to file instead of returning content.
    """
    _check_dependency("pypdf", "pypdf")

    # Use reportlab if available for better PDF generation
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=12,
        )
        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            fontSize=10,
            textColor="gray",
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=11,
            leading=14,
        )

        story = []

        # Title
        story.append(Paragraph(note.title, title_style))

        # Metadata
        story.append(Paragraph(f"Modified: {note.modified.strftime('%Y-%m-%d %H:%M')}", meta_style))
        tags_str = ", ".join(note.tags) if note.tags else "none"
        story.append(Paragraph(f"Tags: {tags_str}", meta_style))
        story.append(Spacer(1, 0.25 * inch))

        # Content - split into paragraphs
        for para in note.content.split("\n\n"):
            if para.strip():
                # Escape HTML entities
                para = para.replace("&", "&amp;")
                para = para.replace("<", "&lt;")
                para = para.replace(">", "&gt;")
                para = para.replace("\n", "<br/>")
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.1 * inch))

        doc.build(story)

    except ImportError:
        # pypdf doesn't support creating PDFs from scratch easily
        # So we'll raise an error suggesting reportlab
        raise MissingDependencyError(
            "PDF export requires 'reportlab' for proper formatting.\n"
            "Install it with: pip install reportlab"
        )


def export_docx(note: Note, output_path: Path) -> None:
    """
    Export note to DOCX format.

    Args:
        note: Note to export
        output_path: Path to save the DOCX

    Note:
        This function writes directly to file instead of returning content.
    """
    _check_dependency("docx", "python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    doc.add_heading(note.title, level=1)

    # Metadata
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Modified: {note.modified.strftime('%Y-%m-%d %H:%M')}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    tags_para = doc.add_paragraph()
    tags_run = tags_para.add_run(f"Tags: {', '.join(note.tags) if note.tags else 'none'}")
    tags_run.font.size = Pt(10)
    tags_run.font.color.rgb = RGBColor(128, 128, 128)

    # Add spacing
    doc.add_paragraph()

    # Content - process markdown-like formatting
    for para_text in note.content.split("\n\n"):
        if not para_text.strip():
            continue

        # Check for headings
        if para_text.startswith("# "):
            doc.add_heading(para_text[2:].strip(), level=1)
        elif para_text.startswith("## "):
            doc.add_heading(para_text[3:].strip(), level=2)
        elif para_text.startswith("### "):
            doc.add_heading(para_text[4:].strip(), level=3)
        elif para_text.startswith("- ") or para_text.startswith("* "):
            # Bullet list
            for line in para_text.split("\n"):
                if line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.strip():
                    doc.add_paragraph(line)
        else:
            # Regular paragraph
            para = doc.add_paragraph()
            _add_formatted_text(para, para_text)

    doc.save(str(output_path))


def _add_formatted_text(paragraph, text: str):
    """Add text with basic markdown formatting to a docx paragraph."""
    # Simple bold/italic handling
    # This is a basic implementation - could be enhanced

    # Process **bold** and *italic*
    parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)", text)

    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def get_supported_export_formats() -> list[str]:
    """Return list of supported export formats."""
    return ["markdown", "text", "html", "json", "rtf", "pdf", "docx"]
