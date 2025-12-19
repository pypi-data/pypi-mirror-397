"""File importer for GPGNotes - converts various formats to markdown notes."""

import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple


class ImportError(Exception):
    """Raised when file import fails."""

    pass


class MissingDependencyError(ImportError):
    """Raised when required dependency is not installed."""

    pass


def _check_dependency(module_name: str, package_name: str):
    """Check if a dependency is available, raise helpful error if not."""
    try:
        __import__(module_name)
    except ModuleNotFoundError:
        raise MissingDependencyError(
            f"The '{package_name}' package is required to import this file type.\n"
            f"Install it with: pip install gpgnotes[import]\n"
            f"Or: pip install {package_name}"
        )


def import_markdown(file_path: Path) -> tuple[str, str]:
    """
    Import a markdown file.

    Args:
        file_path: Path to the markdown file

    Returns:
        Tuple of (title, content)
    """
    content = file_path.read_text(encoding="utf-8")

    # Try to extract title from first heading
    title = file_path.stem
    lines = content.split("\n")
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    return title, content


def import_text(file_path: Path) -> tuple[str, str]:
    """
    Import a plain text file.

    Args:
        file_path: Path to the text file

    Returns:
        Tuple of (title, content)
    """
    content = file_path.read_text(encoding="utf-8")
    title = file_path.stem

    # Use first non-empty line as title if it looks like a title
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if line and len(line) < 100:
            title = line
            break

    return title, content


def import_rtf(file_path: Path) -> tuple[str, str]:
    """
    Import an RTF file, converting to plain text.

    Args:
        file_path: Path to the RTF file

    Returns:
        Tuple of (title, content)
    """
    _check_dependency("striprtf", "striprtf")
    from striprtf.striprtf import rtf_to_text

    rtf_content = file_path.read_text(encoding="utf-8", errors="ignore")
    content = rtf_to_text(rtf_content)

    # Clean up the content
    content = content.strip()

    # Use filename as title, or first line if short enough
    title = file_path.stem
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if line and len(line) < 100:
            title = line
            break

    return title, content


def import_pdf(file_path: Path) -> tuple[str, str]:
    """
    Import a PDF file, extracting text content.

    Args:
        file_path: Path to the PDF file

    Returns:
        Tuple of (title, content)
    """
    _check_dependency("pypdf", "pypdf")
    from pypdf import PdfReader

    reader = PdfReader(file_path)

    # Extract text from all pages
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    content = "\n\n".join(text_parts)

    # Try to get title from PDF metadata
    title = file_path.stem
    if reader.metadata and reader.metadata.title:
        title = reader.metadata.title

    # Clean up content - normalize whitespace
    content = re.sub(r"\n{3,}", "\n\n", content)
    content = content.strip()

    return title, content


def import_docx(file_path: Path) -> tuple[str, str]:
    """
    Import a Word document (.docx), converting to markdown.

    Args:
        file_path: Path to the docx file

    Returns:
        Tuple of (title, content)
    """
    _check_dependency("docx", "python-docx")
    from docx import Document

    doc = Document(file_path)

    # Extract content with basic formatting preservation
    content_parts = []
    title = file_path.stem

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            content_parts.append("")
            continue

        # Check for headings
        if para.style.name.startswith("Heading"):
            level = 1
            try:
                level = int(para.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            prefix = "#" * level
            content_parts.append(f"{prefix} {text}")

            # Use first heading as title
            if i == 0 or (title == file_path.stem and level == 1):
                title = text
        else:
            # Handle basic inline formatting
            formatted_text = _format_docx_paragraph(para)
            content_parts.append(formatted_text)

    # Handle tables
    for table in doc.tables:
        content_parts.append("")
        content_parts.append(_convert_docx_table(table))

    content = "\n\n".join(content_parts)

    # Clean up multiple blank lines
    content = re.sub(r"\n{3,}", "\n\n", content)
    content = content.strip()

    return title, content


def _format_docx_paragraph(para) -> str:
    """Format a docx paragraph with basic markdown formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue

        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts)


def _convert_docx_table(table) -> str:
    """Convert a docx table to markdown format."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

        # Add header separator after first row
        if i == 0:
            separator = "| " + " | ".join(["---"] * len(cells)) + " |"
            rows.append(separator)

    return "\n".join(rows)


def import_file(file_path: Path, title: Optional[str] = None) -> tuple[str, str]:
    """
    Import a file based on its extension.

    Args:
        file_path: Path to the file to import
        title: Optional custom title (overrides auto-detected title)

    Returns:
        Tuple of (title, content)

    Raises:
        ImportError: If file type is not supported
        FileNotFoundError: If file does not exist
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    importers = {
        ".md": import_markdown,
        ".markdown": import_markdown,
        ".txt": import_text,
        ".text": import_text,
        ".rtf": import_rtf,
        ".pdf": import_pdf,
        ".docx": import_docx,
    }

    if suffix not in importers:
        supported = ", ".join(sorted(importers.keys()))
        raise ImportError(f"Unsupported file type: {suffix}\nSupported formats: {supported}")

    detected_title, content = importers[suffix](file_path)

    # Use custom title if provided
    final_title = title if title else detected_title

    return final_title, content


def get_supported_extensions() -> list[str]:
    """Return list of supported file extensions."""
    return [".md", ".markdown", ".txt", ".text", ".rtf", ".pdf", ".docx"]


def import_url(url: str, title: Optional[str] = None) -> Tuple[str, str]:
    """
    Import content from a URL (web clipping).

    Args:
        url: URL to fetch
        title: Optional custom title

    Returns:
        Tuple of (title, content_markdown)
    """
    try:
        import urllib.request
        from html.parser import HTMLParser
    except ImportError as e:
        raise ImportError(f"Failed to import required modules: {e}")

    # HTML to Markdown converter with content extraction
    class HTMLToMarkdown(HTMLParser):
        # Tags to completely ignore (skip content)
        SKIP_TAGS = {
            "nav",
            "header",
            "footer",
            "aside",
            "script",
            "style",
            "noscript",
            "iframe",
            "form",
            "button",
        }

        # Content container tags (prefer these)
        CONTENT_TAGS = {"article", "main"}

        def __init__(self):
            super().__init__()
            self.markdown = []
            self.current_tag = None
            self.list_level = 0
            self.in_pre = False
            self.title = None
            self.skip_depth = 0  # Track depth of skipped tags
            self.in_content = False  # Track if we're in article/main
            self.content_depth = 0  # Track depth of content tags

        def handle_starttag(self, tag, attrs):
            # Skip ignored tags and their content
            if tag in self.SKIP_TAGS:
                self.skip_depth += 1
                return

            # Track content container tags
            if tag in self.CONTENT_TAGS:
                self.in_content = True
                self.content_depth += 1
                return

            # Skip processing if we're in a skipped tag
            if self.skip_depth > 0:
                return

            # Process content tags
            if tag == "h1":
                self.current_tag = "h1"
            elif tag == "h2":
                self.current_tag = "h2"
            elif tag == "h3":
                self.current_tag = "h3"
            elif tag == "h4":
                self.current_tag = "h4"
            elif tag == "p":
                self.current_tag = "p"
            elif tag == "strong" or tag == "b":
                self.markdown.append("**")
            elif tag == "em" or tag == "i":
                self.markdown.append("*")
            elif tag == "code":
                self.markdown.append("`")
            elif tag == "pre":
                self.in_pre = True
                self.markdown.append("\n```\n")
            elif tag == "a":
                href = dict(attrs).get("href", "")
                self.markdown.append("[")
                self.current_tag = ("a", href)
            elif tag == "ul":
                self.list_level += 1
            elif tag == "ol":
                self.list_level += 1
            elif tag == "li":
                self.markdown.append(f"\n{'  ' * (self.list_level - 1)}- ")
            elif tag == "br":
                self.markdown.append("\n")

        def handle_endtag(self, tag):
            # Handle skipped tags
            if tag in self.SKIP_TAGS:
                self.skip_depth = max(0, self.skip_depth - 1)
                return

            # Handle content container tags
            if tag in self.CONTENT_TAGS:
                self.content_depth = max(0, self.content_depth - 1)
                if self.content_depth == 0:
                    self.in_content = False
                return

            # Skip processing if we're in a skipped tag
            if self.skip_depth > 0:
                return

            # Process end tags
            if tag in ["h1", "h2", "h3", "h4", "p"]:
                self.markdown.append("\n\n")
                self.current_tag = None
            elif tag == "strong" or tag == "b":
                self.markdown.append("**")
            elif tag == "em" or tag == "i":
                self.markdown.append("*")
            elif tag == "code":
                self.markdown.append("`")
            elif tag == "pre":
                self.in_pre = False
                self.markdown.append("\n```\n")
            elif tag == "a":
                if isinstance(self.current_tag, tuple) and self.current_tag[0] == "a":
                    href = self.current_tag[1]
                    self.markdown.append(f"]({href})")
                    self.current_tag = None
            elif tag == "ul" or tag == "ol":
                self.list_level -= 1
                if self.list_level == 0:
                    self.markdown.append("\n")

        def handle_data(self, data):
            # Skip if we're in a skipped tag
            if self.skip_depth > 0:
                return

            data = data.strip()
            if data:
                if self.current_tag == "h1":
                    self.markdown.append(f"# {data}")
                    if not self.title:
                        self.title = data
                elif self.current_tag == "h2":
                    self.markdown.append(f"## {data}")
                elif self.current_tag == "h3":
                    self.markdown.append(f"### {data}")
                elif self.current_tag == "h4":
                    self.markdown.append(f"#### {data}")
                elif self.in_pre:
                    self.markdown.append(data)
                else:
                    # Regular text
                    if self.markdown and not self.markdown[-1].endswith(
                        (" ", "\n", "**", "*", "`", "[")
                    ):
                        self.markdown.append(" ")
                    self.markdown.append(data)

        def get_markdown(self):
            return "".join(self.markdown).strip()

    # Fetch the URL
    try:
        req = urllib.request.Request(
            url, headers={"User-Agent": "Mozilla/5.0 (GPGNotes Web Clipper)"}
        )
        with urllib.request.urlopen(req, timeout=30) as response:
            html_content = response.read().decode("utf-8", errors="ignore")
    except Exception as e:
        raise ImportError(f"Failed to fetch URL: {e}")

    # Parse HTML and convert to markdown
    parser = HTMLToMarkdown()
    parser.feed(html_content)

    markdown_content = parser.get_markdown()
    detected_title = parser.title or "Web Clip"

    # Add metadata header
    clipped_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    metadata_header = f"""---
source_url: {url}
clipped_at: {clipped_at}
---

*Clipped from [{url}]({url})*

"""

    final_content = metadata_header + markdown_content
    final_title = title if title else detected_title

    return final_title, final_content
