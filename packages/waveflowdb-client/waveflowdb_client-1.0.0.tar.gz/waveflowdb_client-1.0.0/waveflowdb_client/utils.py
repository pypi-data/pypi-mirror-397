import os
import csv
import logging
from pathlib import Path
from datetime import datetime
import json
import shutil
from typing import List,Tuple

from .exceptions import FileProcessingError

# logger = logging.getLogger(__name__)
# logging.basicConfig(level=logging.INFO)

import logging

# Configure the logging settings
logging.basicConfig(
    level=logging.DEBUG,  # Set to INFO to see only high-level actions, or DEBUG for detail
    format='%(asctime)s - %(levelname)s - %(name)s - %(message)s',
    filename='batch_creation.log',  # Log output to a file
    filemode='w'  # Overwrite the log file each run
)

# If you also want output to the console:
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO) 
formatter = logging.Formatter('%(levelname)s: %(message)s')
console_handler.setFormatter(formatter)
logging.getLogger().addHandler(console_handler)

# Get the logger instance for your class (if needed, though basicConfig covers the root)
logger = logging.getLogger(__name__)

# -------------------------------------------------------------------------
#                UNIVERSAL HELPERS FOR FILE CONVERSION & CHUNKING
# -------------------------------------------------------------------------

def paragraph_chunk(text: str, max_chars: int) -> list:
    """
    Strict paragraph-based chunking.
    No paragraph is ever split.
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    current = ""

    for p in paragraphs:
        if len(current) + len(p) + 2 > max_chars:
            if current:
                chunks.append(current.strip())
            current = p + "\n\n"
        else:
            current += p + "\n\n"

    if current.strip():
        chunks.append(current.strip())

    return chunks


def json_to_csv(json_path: str) -> str:
    out_path = json_path.replace(".json", ".csv")

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, list):
        raise FileProcessingError("JSON must be an array of objects.")

    # Extract all keys across JSON objects
    headers = sorted({k for row in data for k in row.keys()})

    with open(out_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=headers)
        writer.writeheader()
        writer.writerows(data)

    return out_path


def convert_to_txt(file_path: str) -> str:
    base, _ = os.path.splitext(file_path)
    out_path = base + ".txt"

    if file_path.endswith(".py"):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

    elif file_path.endswith(".ipynb"):
        import json
        with open(file_path, "r", encoding="utf-8") as f:
            nb = json.load(f)
        content = "\n\n".join(
            cell["source"] if isinstance(cell["source"], str) 
            else "".join(cell["source"])
            for cell in nb.get("cells", []) if cell.get("cell_type") == "code"
        )

    else:
        raise FileProcessingError("Unsupported file for text conversion")

    with open(out_path, "w", encoding="utf-8") as out:
        out.write(content)

    return out_path


def extract_text(file_path: str) -> str:
    """
    Extract text from PDF, DOCX, or TXT.
    """
    ext = file_path.lower().split(".")[-1]

    if ext == "pdf":
        try:
            import PyPDF2
        except Exception as e:
            raise FileProcessingError(f"PyPDF2 not available: {e}")
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for p in reader.pages:
                text += (p.extract_text() or "") + "\n\n"
        return text

    if ext == "docx":
        try:
            from docx import Document
        except Exception as e:
            raise FileProcessingError(f"python-docx not available: {e}")
        doc = Document(file_path)
        return "\n\n".join(p.text for p in doc.paragraphs)

    if ext == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    raise FileProcessingError("Unsupported extraction type.")


def split_csv_if_needed(file_path: str, max_bytes: int) -> list:
    """
    Split CSV into smaller parts using 20% safety buffer.
    No partial rows.
    """
    SAFE_LIMIT = int(max_bytes * 0.8)
    size = os.path.getsize(file_path)

    if size <= SAFE_LIMIT:
        return [file_path]

    base, ext = os.path.splitext(file_path)
    output_files = []

    with open(file_path, "r", newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        header = next(reader)
        header_bytes = len(",".join(header).encode("utf-8"))

        part = 1
        rows = []
        current_size = header_bytes

        for row in reader:
            row_bytes = len(",".join(row).encode("utf-8"))

            if row_bytes > SAFE_LIMIT:
                raise FileProcessingError(f"A CSV row exceeds the safe max size.")

            if current_size + row_bytes > SAFE_LIMIT:
                out_file = f"{base}_part_{part}{ext}"
                with open(out_file, "w", newline="", encoding="utf-8") as out:
                    writer = csv.writer(out)
                    writer.writerow(header)
                    writer.writerows(rows)
                output_files.append(out_file)

                part += 1
                rows = []
                current_size = header_bytes

            rows.append(row)
            current_size += row_bytes

        # last chunk
        if rows:
            out_file = f"{base}_part_{part}{ext}"
            with open(out_file, "w", newline="", encoding="utf-8") as out:
                writer = csv.writer(out)
                writer.writerow(header)
                writer.writerows(rows)
            output_files.append(out_file)

    return output_files


# -------------------------------------------------------------------------
#                            FILE PROCESSOR
# -------------------------------------------------------------------------

class FileProcessor:
    SUPPORTED_EXTENSIONS = ['txt', 'csv', 'json', 'py', 'docx', 'pdf', 'ipynb']

    @staticmethod
    def read_file_content(filepath: str) -> str:
        ext = filepath.lower().split('.')[-1]
        if ext in ['txt', 'csv', 'py', 'json']:
            with open(filepath, encoding='utf-8') as f:
                return f.read()
        if ext == 'docx':
            from docx import Document
            doc = Document(filepath)
            return '\n'.join(p.text for p in doc.paragraphs)
        if ext == 'pdf':
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return '\n'.join(p.extract_text() or '' for p in reader.pages)
        raise FileProcessingError(f'Unsupported extension: {ext}')

    @staticmethod
    def get_file_size(filepath: str) -> int:
        return os.path.getsize(filepath)

    @staticmethod
    def is_supported_file(filename: str) -> bool:
        ext = filename.lower().split('.')[-1]
        return ext in FileProcessor.SUPPORTED_EXTENSIONS


# -------------------------------------------------------------------------
#                               LOGGER
# -------------------------------------------------------------------------

class Logger:
    def __init__(self, log_dir: str):
        Path(log_dir).mkdir(parents=True, exist_ok=True)
        self.skipped_log = Path(log_dir) / 'skipped_files_log.csv'
        self.api_error_log = Path(log_dir) / 'api_error_log.csv'
        self.performance_log = Path(log_dir) / 'api_performance_log.csv'

    def _write_csv_log(self, path: Path, header: List[str], row: List):
        exists = path.exists()
        with open(path, 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not exists:
                writer.writerow(header)
            writer.writerow(row)

    def log_skipped_file(self, filename: str, reason: str):
        self._write_csv_log(self.skipped_log,
                            ['ts', 'filename', 'reason'],
                            [datetime.utcnow().isoformat(), filename, reason])

    def log_api_error(self, operation: str, batch_num: int, error_message: str):
        self._write_csv_log(self.api_error_log,
                            ['ts', 'operation', 'batch_num', 'err'],
                            [datetime.utcnow().isoformat(), operation, batch_num, error_message])

    def log_performance(self, **kwargs):
        """
        SAFE CSV logger for performance metrics.
        Accepts ANY keyword args and serializes them safely.
        """
        logfile = self.performance_log

        # Start row with timestamp
        row = {"timestamp": datetime.utcnow().isoformat()}

        # Safely encode all user-provided fields
        for k, v in kwargs.items():
            try:
                if isinstance(v, (list, tuple, set)):
                    row[k] = ",".join(str(i) for i in v)
                else:
                    row[k] = str(v) if v is not None else ""
            except Exception:
                row[k] = "UNSERIALIZABLE"

        # Write row
        file_exists = logfile.exists()
        with open(logfile, "a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=row.keys())
            if not file_exists:
                writer.writeheader()
            writer.writerow(row)


# -------------------------------------------------------------------------
#                         BATCH MANAGER (UPDATED)
# -------------------------------------------------------------------------
class BatchManager:
    def __init__(self, max_files: int, max_size_mb: int):
        self.max_files = max_files
        self.max_bytes = max_size_mb * 1024 * 1024  # convert MB to bytes


    def create_batches(self, files: List[str], base_path: str) -> Tuple[List[List[str]], str]:
        """
        Process files into batches based on max files & max bytes, with detailed logging.
        """
        
        # Log initial parameters
        logger.info(f"Starting batch creation. Max files: {self.max_files}, Max bytes: {self.max_bytes}")
        
        SAFE_LIMIT = int(self.max_bytes * 0.8)
        logger.debug(f"Calculated SAFE_LIMIT (80% of max_bytes): {SAFE_LIMIT} bytes")
        
        chunks_dir = os.path.join(base_path, "chunks")
        os.makedirs(chunks_dir, exist_ok=True)
        logger.info(f"Chunks directory created: {chunks_dir}")

        file_info = []

        # -----------------------------
        # Step 1: Preprocess files
        # -----------------------------
        logger.info("--- Step 1: Preprocessing and Chunking Files ---")
        
        for f in files:
            filepath = os.path.join(base_path, f)
            
            if not os.path.isfile(filepath):
                logger.warning(f"File not found, skipping: {filepath}")
                continue

            ext = f.lower().split(".")[-1]
            logger.debug(f"Processing file: {f}, detected extension: {ext}")

            # JSON → CSV
            if ext == "json":
                original_path = filepath
                filepath = json_to_csv(filepath)
                ext = "csv"
                logger.info(f"Converted JSON file {original_path} to CSV at {filepath}")

            # PY/IPYNB → TXT
            elif ext in ("py", "ipynb"):
                original_path = filepath
                filepath = convert_to_txt(filepath)
                ext = "txt"
                logger.info(f"Converted code file {original_path} to TXT at {filepath}")

            # PDF/DOCX/TXT → paragraph chunks if needed
            if ext in ("pdf", "docx", "txt") and ext != "csv":
                text = extract_text(filepath)
                chunks = paragraph_chunk(text, max_chars=SAFE_LIMIT)
                base_fname, _ = os.path.splitext(os.path.basename(filepath))
                part_files = []

                for i, block in enumerate(chunks, 1):
                    out_file = os.path.join(chunks_dir, f"{base_fname}_part_{i}.txt")
                    with open(out_file, "w", encoding="utf-8") as out:
                        out.write(block)
                    size = os.path.getsize(out_file)
                    part_files.append((out_file, size))
                    logger.debug(f"Chunked file {f} into part {i}: {os.path.basename(out_file)}, Size: {size} bytes")

                file_info.extend(part_files)
                logger.info(f"Successfully chunked {f} into {len(chunks)} parts.")
                continue

            # CSV → safe split
            if ext == "csv":
                parts = self._split_csv_safe(filepath, SAFE_LIMIT, chunks_dir)
                sizes = [os.path.getsize(p) for p in parts]
                file_info.extend(list(zip(parts, sizes)))
                logger.info(f"Successfully split CSV file {f} into {len(parts)} safe chunks.")
                continue

            # Other small files → move to chunks folder or chunk if large
            size = os.path.getsize(filepath)
            if size > SAFE_LIMIT:
                logger.warning(f"File {f} ({size} bytes) exceeds SAFE_LIMIT ({SAFE_LIMIT} bytes). Attempting generic TXT chunking.")
                try:
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f_in:
                        content = f_in.read()
                    
                    chunks = paragraph_chunk(content, max_chars=SAFE_LIMIT)
                    base_fname, _ = os.path.splitext(os.path.basename(filepath))
                    for i, block in enumerate(chunks, 1):
                        out_file = os.path.join(chunks_dir, f"{base_fname}_part_{i}.txt")
                        with open(out_file, "w", encoding="utf-8") as out:
                            out.write(block)
                        file_info.append((out_file, os.path.getsize(out_file)))
                        logger.debug(f"Generic chunked file {f} into part {i}: {os.path.basename(out_file)}")
                
                except Exception as e:
                    logger.error(f"Failed to chunk large file {filepath} as TXT. Error: {e}")
                    raise FileProcessingError(f"File {filepath} exceeds batch limit and cannot be chunked")
            
            else:
                # Copy small files to chunks directory
                new_path = os.path.join(chunks_dir, os.path.basename(filepath))
                if not os.path.exists(new_path):
                    shutil.copy2(filepath, new_path)
                file_info.append((new_path, size))
                logger.debug(f"Copied small file {f} to chunks_dir. Size: {size} bytes.")

        logger.info(f"Preprocessing complete. Total items ready for batching: {len(file_info)}")

        # -----------------------------
        # Step 2: Create dynamic batches
        # -----------------------------
        logger.info("--- Step 2: Creating Dynamic Batches ---")
        batches = []
        cur_batch = []
        cur_size = 0

        for filepath, size in file_info:
            filename = os.path.basename(filepath)
            logger.debug(f"Considering file: {filename}, Size: {size} bytes. Current batch size: {cur_size} bytes, files: {len(cur_batch)}")
            
            # Absolute safety check (shouldn't happen if Step 1 worked)
            if size > self.max_bytes:
                logger.critical(f"Chunked file {filename} ({size} bytes) still exceeds max batch size ({self.max_bytes} bytes).")
                raise FileProcessingError(f"File {filepath} exceeds max batch size ({self.max_bytes} bytes)")

            # Flush current batch if limits exceeded
            if cur_batch and (len(cur_batch) >= self.max_files or (cur_size + size) > self.max_bytes):
                logger.info(f"Batch limit hit. Flushing batch with {len(cur_batch)} files and {cur_size} bytes.")
                batches.append(cur_batch)
                cur_batch = []
                cur_size = 0
                
            cur_batch.append(filename)
            cur_size += size
            logger.debug(f"Added file {filename}. New batch size: {cur_size} bytes, files: {len(cur_batch)}")

        if cur_batch:
            logger.info(f"Flushing final batch with {len(cur_batch)} files and {cur_size} bytes.")
            batches.append(cur_batch)
        
        logger.info(f"Batch creation complete. Total batches created: {len(batches)}")

        return batches, chunks_dir

    def _split_csv_safe(self, file_path: str, safe_limit: int, chunks_dir: str, preferred_encoding: str = "utf-8") -> List[str]:
        """
        Split CSV safely into chunks respecting row boundaries.
        Tries multiple common encodings if the preferred one fails, ensuring robustness
        across different operating systems and file origins (e.g., handling cp1252/Latin-1).
        Includes an aggressive fallback to 'cp1252' if 'utf-8' is suspected of failing deep within the file.
        Uses 'errors='replace'' in the final read step for maximum stability against corrupted bytes.

        Returns a list of chunk file paths inside chunks_dir.
        """
        input_file = Path(file_path)
        output_dir = Path(chunks_dir)
        output_files = []

        # 1. Ensure the output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)

        base_fname = input_file.stem
        ext = input_file.suffix

        # 2. Robust Encoding Handling: Determine the correct encoding
        encodings_to_try = [preferred_encoding, "cp1252", "latin-1"]
        effective_encoding = None

        for encoding in encodings_to_try:
            try:
                # Attempt to open and read only the header to test the encoding.
                # We do not use 'errors' here to ensure we only accept encodings that cleanly read the header.
                with open(input_file, "r", newline="", encoding=encoding) as f_test:
                    test_reader = csv.reader(f_test)
                    # Attempt to read the first row (header)
                    header_test = next(test_reader)
                    
                    effective_encoding = encoding
                    break
            except UnicodeDecodeError:
                print(f"Failed to decode with {encoding}. Trying next encoding...")
                continue
            except StopIteration:
                raise FileProcessingError(f"CSV file {file_path} is empty.")
            except Exception as e:
                raise FileProcessingError(f"Error opening or reading file {file_path}: {e}")

        # AGGRESSIVE FALLBACK LOGIC:
        # If the file appears to be UTF-8 based on the header check, but is running on a Windows environment ('nt'),
        # we preemptively switch to cp1252 to handle deep file corruption caused by non-UTF-8 characters.
        if effective_encoding == preferred_encoding and preferred_encoding == "utf-8" and os.name == 'nt':
            effective_encoding = "cp1252"
            print(f"Warning: Aggressively switching encoding from {preferred_encoding} to {effective_encoding} to prevent deep file UnicodeDecodeError common in Windows CSVs.")
            
        if not effective_encoding:
            raise FileProcessingError(f"Could not decode CSV file {file_path} using any of the tested encodings: {', '.join(encodings_to_try)}")

        # 3. Main Splitting Logic: Open the file again using the confirmed or aggressively set encoding
        # CRITICAL FIX: Add errors='replace' to the main open call. This guarantees 
        # that if any remaining problematic bytes are found deep in the file, Python
        # replaces them with a safe character instead of crashing the process.
        with open(input_file, "r", newline="", encoding=effective_encoding, errors='replace') as f:
            reader = csv.reader(f)
            try:
                header = next(reader)
            except StopIteration:
                # Should not happen if Step 2 passed, but defensive coding is included.
                return output_files

            # Calculate header size (using UTF-8 as the target output encoding for consistency)
            header_bytes = len(",".join(header).encode("utf-8"))

            part = 1
            rows = []
            current_size = header_bytes
            
            # The row iteration now runs with the errors='replace' safety net from 'f'
            for row in reader:
                try:
                    # Calculate row size based on the output encoding (UTF-8)
                    row_bytes = len(",".join(row).encode("utf-8"))
                except Exception as e:
                    print(f"Warning: Skipping problematic row during size calculation: {e}")
                    continue

                if row_bytes > safe_limit:
                    raise FileProcessingError(f"A CSV row exceeds the safe max size ({safe_limit} bytes) at part {part}")

                if current_size + row_bytes > safe_limit:
                    # Write the current chunk
                    out_path = output_dir / f"{base_fname}_part_{part}{ext}"
                    with open(out_path, "w", newline="", encoding="utf-8") as out:
                        writer = csv.writer(out)
                        writer.writerow(header)
                        writer.writerows(rows)
                    
                    output_files.append(str(out_path))
                    
                    # Reset for the next chunk
                    part += 1
                    rows = []
                    current_size = header_bytes

                # Add the row to the current chunk
                rows.append(row)
                current_size += row_bytes

            # Write the final chunk
            if rows:
                out_path = output_dir / f"{base_fname}_part_{part}{ext}"
                with open(out_path, "w", newline="", encoding="utf-8") as out:
                    writer = csv.writer(out)
                    writer.writerow(header)
                    writer.writerows(rows)
                output_files.append(str(out_path))

        return output_files