from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QTimer, Qt, QDate
from PyQt5.QtWidgets import QMainWindow, QWidget, QDialog, QInputDialog, QApplication, QLineEdit
from pyqtgraph import PlotWidget, plot
from random import random
import pyqtgraph as pg
import sqlite3
import traceback
import os
import time
import threading
import sys
import math
import shutil
from docx.shared import Pt, Cm
from docx import Document
from sys import platform
from Exchange_data_str_bu7 import *
from Exchange_data_str_itm import *
# os.environ["QT_IM_MODULE"] = "qtvirtualkeyboard"

class MainWindow(QMainWindow):

    def __init__(self):
        super().__init__()
        self.setWindowIcon(QtGui.QIcon('Logo_etalon.ico'))
        self.setupUi(self)

    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1024, 600)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(MainWindow.sizePolicy().hasHeightForWidth())
        MainWindow.setSizePolicy(sizePolicy)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setMaximumSize(QtCore.QSize(1024, 600))
        self.centralwidget.setObjectName("centralwidget")
        
        self.stackedWidget_main = QtWidgets.QStackedWidget(self.centralwidget)
        self.stackedWidget_main.setGeometry(QtCore.QRect(9, 9, 981, 581))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.stackedWidget_main.sizePolicy().hasHeightForWidth())
        self.stackedWidget_main.setSizePolicy(sizePolicy)
        self.stackedWidget_main.setObjectName("stackedWidget_main")
         
#   -----------------------------------------------------------------
#   вкладка Лого
#   -----------------------------------------------------------------
       
        self.Logo = QtWidgets.QWidget()
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.Logo.sizePolicy().hasHeightForWidth())
        self.Logo.setSizePolicy(sizePolicy)
        self.Logo.setObjectName("Logo")
        self.label_logo = QtWidgets.QLabel(self.Logo)
        self.label_logo.setEnabled(True)
        self.label_logo.setGeometry(QtCore.QRect(0, 0, 1024, 600))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Maximum, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_logo.sizePolicy().hasHeightForWidth())
        self.label_logo.setSizePolicy(sizePolicy)
        self.label_logo.setStyleSheet("background-color: rgb(243, 242, 242);")
        self.label_logo.setText("")
        self.label_logo.setPixmap(QtGui.QPixmap("Logo_etalon.png"))
        self.label_logo.setObjectName("label_logo")
        self.tab_logo = self.stackedWidget_main.addWidget(self.Logo)
        
#   -----------------------------------------------------------------
#   вкладка Настройки
#   -----------------------------------------------------------------

        self.Settings = QtWidgets.QWidget()
        self.Settings.setObjectName("Settings")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.Settings)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.stackedWidget_settings_tabs = QtWidgets.QStackedWidget(self.Settings)
        self.stackedWidget_settings_tabs.setObjectName("stackedWidget_settings_tabs")
        self.settings_basic = QtWidgets.QWidget()
        self.settings_basic.setObjectName("settings_basic")
        self.verticalLayout_settings = QtWidgets.QVBoxLayout(self.settings_basic)
        self.verticalLayout_settings.setObjectName("verticalLayout_settings")
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_device_nubmer = QtWidgets.QLabel(self.settings_basic)
        self.label_device_nubmer.setFont(font)
        self.label_device_nubmer.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_device_nubmer.setObjectName("label_device_nubmer")
        self.verticalLayout_settings.addWidget(self.label_device_nubmer)
        self.lineEdit_device_number = QtWidgets.QLineEdit(self.settings_basic)
        self.lineEdit_device_number.setFont(font)
        self.lineEdit_device_number.setObjectName("lineEdit_device_number")
        self.lineEdit_device_number.textEdited.connect(self.device_number_changed)
        self.verticalLayout_settings.addWidget(self.lineEdit_device_number)
        self.label_device_name = QtWidgets.QLabel(self.settings_basic)
        self.label_device_name.setFont(font)
        self.label_device_name.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_device_name.setObjectName("label_device_name")
        self.verticalLayout_settings.addWidget(self.label_device_name)
        self.lineEdit_device_name = QtWidgets.QLineEdit(self.settings_basic)
        self.lineEdit_device_name.setFont(font)
        self.lineEdit_device_name.setObjectName("lineEdit_device_name")
        self.lineEdit_device_name.textEdited.connect(self.device_name_changed)
        self.verticalLayout_settings.addWidget(self.lineEdit_device_name)
        self.label_device_next_veryfing_date = QtWidgets.QLabel(self.settings_basic)
        self.label_device_next_veryfing_date.setFont(font)
        self.label_device_next_veryfing_date.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_device_next_veryfing_date.setObjectName("label_device_next_veryfing_date")
        self.verticalLayout_settings.addWidget(self.label_device_next_veryfing_date)
        self.lineEdit_device_next_veryfing_date = QtWidgets.QLineEdit(self.settings_basic)
        self.lineEdit_device_next_veryfing_date.setFont(font)
        self.lineEdit_device_next_veryfing_date.setObjectName("lineEdit_device_next_veryfing_date")
        self.lineEdit_device_next_veryfing_date.textEdited.connect(self.device_next_veryfing_date_changed)
        self.verticalLayout_settings.addWidget(self.lineEdit_device_next_veryfing_date)
        self.label_devic_veryfing_date = QtWidgets.QLabel(self.settings_basic)
        self.label_devic_veryfing_date.setFont(font)
        self.label_devic_veryfing_date.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_devic_veryfing_date.setObjectName("label_devic_veryfing_date")
        self.verticalLayout_settings.addWidget(self.label_devic_veryfing_date)
        self.lineEdit_device_produce_date = QtWidgets.QLineEdit(self.settings_basic)
        self.lineEdit_device_produce_date.setFont(font)
        self.lineEdit_device_produce_date.setObjectName("lineEdit_device_produce_date")
        self.lineEdit_device_produce_date.textEdited.connect(self.device_produce_date_changed)
        self.verticalLayout_settings.addWidget(self.lineEdit_device_produce_date)
        self.label_device_produce_date = QtWidgets.QLabel(self.settings_basic)
        self.label_device_produce_date.setFont(font)
        self.label_device_produce_date.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_device_produce_date.setObjectName("label_device_produce_date")
        self.verticalLayout_settings.addWidget(self.label_device_produce_date)
        self.lineEdit_device_veryfing_date = QtWidgets.QLineEdit(self.settings_basic)
        self.lineEdit_device_veryfing_date.setFont(font)
        self.lineEdit_device_veryfing_date.setObjectName("lineEdit_device_veryfing_date")
        self.lineEdit_device_veryfing_date.textEdited.connect(self.device_veryfing_date_changed)
        self.verticalLayout_settings.addWidget(self.lineEdit_device_veryfing_date)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout_settings.addItem(spacerItem)
        self.settings_basic_tab = self.stackedWidget_settings_tabs.addWidget(self.settings_basic)
        self.horizontalLayout.addWidget(self.stackedWidget_settings_tabs)

        self.settings_connections = QtWidgets.QWidget()
        self.settings_connections.setObjectName("settings_connections")
        self.settings_connections_tab = self.stackedWidget_settings_tabs.addWidget(self.settings_connections)
        self.settings_verifying = QtWidgets.QWidget()
        self.settings_verifying.setObjectName("settings_verifying")
        self.settings_verifying_tab = self.stackedWidget_settings_tabs.addWidget(self.settings_verifying)
        self.settings_device = QtWidgets.QWidget()
        self.settings_device.setObjectName("settings_device")
        self.verticalLayout_settings_device = QtWidgets.QVBoxLayout(self.settings_device)
        self.verticalLayout_settings_device.setObjectName("verticalLayout_settings_device")
        self.label_comport_bu7 = QtWidgets.QLabel(self.settings_device)
        self.label_comport_bu7.setFont(font)
        self.label_comport_bu7.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_comport_bu7.setObjectName("label_comport_bu7")
        self.verticalLayout_settings_device.addWidget(self.label_comport_bu7)
        self.lineEdit_comport_bu7 = QtWidgets.QLineEdit(self.settings_device)
        self.lineEdit_comport_bu7.setFont(font)
        self.lineEdit_comport_bu7.setObjectName("lineEdit_comport_bu7")
        self.lineEdit_comport_bu7.textEdited.connect(self.comport_bu7_changed)
        self.verticalLayout_settings_device.addWidget(self.lineEdit_comport_bu7)
        self.label_comport_itm = QtWidgets.QLabel(self.settings_device)
        self.label_comport_itm.setFont(font)
        self.label_comport_itm.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_comport_itm.setObjectName("label_comport_itm")
        self.verticalLayout_settings_device.addWidget(self.label_comport_itm)
        self.lineEdit_comport_itm = QtWidgets.QLineEdit(self.settings_device)
        self.lineEdit_comport_itm.setFont(font)
        self.lineEdit_comport_itm.setObjectName("lineEdit_comport_itm")
        self.lineEdit_comport_itm.textEdited.connect(self.comport_itm_changed)
        self.verticalLayout_settings_device.addWidget(self.lineEdit_comport_itm)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout_settings_device.addItem(spacerItem)
        self.settings_device_tab = self.stackedWidget_settings_tabs.addWidget(self.settings_device)
        self.settings_system_update = QtWidgets.QWidget()
        self.settings_system_update.setObjectName("settings_system_update")
        self.settings_system_update_tab = self.stackedWidget_settings_tabs.addWidget(self.settings_system_update)

        self.verticalLayout_buttons = QtWidgets.QVBoxLayout()
        self.verticalLayout_buttons.setSizeConstraint(QtWidgets.QLayout.SizeConstraint.SetDefaultConstraint)
        self.verticalLayout_buttons.setSpacing(20)
        self.verticalLayout_buttons.setObjectName("verticalLayout_buttons")
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        self.pushButton_settings_basic = QtWidgets.QPushButton(self.Settings)
        sizePolicy.setHeightForWidth(self.pushButton_settings_basic.sizePolicy().hasHeightForWidth())
        self.pushButton_settings_basic.setSizePolicy(sizePolicy)
        self.pushButton_settings_basic.setFont(font)
        self.pushButton_settings_basic.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_settings_basic.setObjectName("pushButton_settings_date")
        self.pushButton_settings_basic.setCheckable(True)
        self.pushButton_settings_basic.setChecked(True)
        self.pushButton_settings_basic.clicked.connect(self.set_settings_basic_tab)
        self.verticalLayout_buttons.addWidget(self.pushButton_settings_basic)
        self.pushButton_settings_connections = QtWidgets.QPushButton(self.Settings)
        sizePolicy.setHeightForWidth(self.pushButton_settings_connections.sizePolicy().hasHeightForWidth())
        self.pushButton_settings_connections.setSizePolicy(sizePolicy)
        self.pushButton_settings_connections.setFont(font)
        self.pushButton_settings_connections.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_settings_connections.setObjectName("pushButton_settings_connections")
        self.pushButton_settings_connections.setCheckable(True)
        self.pushButton_settings_connections.clicked.connect(self.set_settings_connections_tab)
        self.verticalLayout_buttons.addWidget(self.pushButton_settings_connections)
        self.pushButton_settings_verifying = QtWidgets.QPushButton(self.Settings)
        sizePolicy.setHeightForWidth(self.pushButton_settings_verifying.sizePolicy().hasHeightForWidth())
        self.pushButton_settings_verifying.setSizePolicy(sizePolicy)
        self.pushButton_settings_verifying.setFont(font)
        self.pushButton_settings_verifying.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_settings_verifying.setObjectName("pushButton_settings_verifying")
        self.pushButton_settings_verifying.setCheckable(True)
        self.pushButton_settings_verifying.clicked.connect(self.set_settings_verifying_tab)
        self.verticalLayout_buttons.addWidget(self.pushButton_settings_verifying)
        self.pushButton_settings_device = QtWidgets.QPushButton(self.Settings)
        self.pushButton_settings_device.setEnabled(True)
        sizePolicy.setHeightForWidth(self.pushButton_settings_device.sizePolicy().hasHeightForWidth())
        self.pushButton_settings_device.setSizePolicy(sizePolicy)
        self.pushButton_settings_device.setFont(font)
        self.pushButton_settings_device.setMouseTracking(False)
        self.pushButton_settings_device.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_settings_device.setObjectName("pushButton_settings_device")
        self.pushButton_settings_device.setCheckable(True)
        self.pushButton_settings_device.clicked.connect(self.set_settings_device_tab)
        self.verticalLayout_buttons.addWidget(self.pushButton_settings_device)
        self.pushButton_settings_system_update = QtWidgets.QPushButton(self.Settings)
        self.pushButton_settings_system_update.setEnabled(True)
        sizePolicy.setHeightForWidth(self.pushButton_settings_system_update.sizePolicy().hasHeightForWidth())
        self.pushButton_settings_system_update.setSizePolicy(sizePolicy)
        self.pushButton_settings_system_update.setFont(font)
        self.pushButton_settings_system_update.setMouseTracking(False)
        self.pushButton_settings_system_update.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_settings_system_update.setObjectName("pushButton_settings_system_update")
        self.pushButton_settings_system_update.setCheckable(True)
        self.pushButton_settings_system_update.clicked.connect(self.set_settings_system_update_tab)
        self.verticalLayout_buttons.addWidget(self.pushButton_settings_system_update)
        self.horizontalLayout.addLayout(self.verticalLayout_buttons)
        self.tab_settings = self.stackedWidget_main.addWidget(self.Settings)

        #   -----------------------------------------------------------------
        #   вкладка Настройка ИСХ
        #   -----------------------------------------------------------------
        self.Settings_ISH = QtWidgets.QWidget()
        self.Settings_ISH.setObjectName("Settings_ISH")

        #   Таблица эталонных датчиков
   
        self.tableWidget_settings_ISH = QtWidgets.QTableWidget(self.Settings_ISH)
        self.tableWidget_settings_ISH.setGeometry(QtCore.QRect(0, 0, 971, 541))
        self.tableWidget_settings_ISH.setObjectName("tableWidget_settings_ISH")
        columnWidths = [50, 80, 321]
        self.tableWidget_settings_ISH.setColumnCount(len(columnWidths))
        rowCount = self.db_query('SELECT COUNT(*) FROM ish_data')[0][0]
        self.tableWidget_settings_ISH.setRowCount(rowCount)
        for column in range(len(columnWidths)):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_settings_ISH.setHorizontalHeaderItem(column, item)
            self.tableWidget_settings_ISH.setColumnWidth(column, columnWidths[column])
        self.tableWidget_settings_ISH.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        sensor_ish_data=self.db_query('SELECT sensor_id FROM ish_data')
        for row in range(rowCount):
            font.setPointSize(16)
            item = QtWidgets.QTableWidgetItem() # Создаём строки таблицы ИСХ
            item.setFont(font)
            item.setText(str(row+1))
            self.tableWidget_settings_ISH.setVerticalHeaderItem(row, item)
            self.tableWidget_settings_ISH.setRowHeight (row, 41)
            font.setPointSize(14)
            sensor=self.db_query('SELECT * FROM sensors WHERE id='+str(sensor_ish_data[row][0]))
            item = QtWidgets.QTableWidgetItem(f'{str(sensor[0][1])}, sn: {sensor[0][2]}')
            item.setFont(font)
            item.setFlags(Qt.ItemFlag.ItemIsSelectable)
            self.tableWidget_settings_ISH.setItem(row, 0, item) # Записываем в таблцу ИСХ название датчика
            item = QtWidgets.QTableWidgetItem(sensor[0][4])
            item.setFont(font)
            item.setFlags(Qt.ItemFlag.ItemIsSelectable)
            self.tableWidget_settings_ISH.setItem(row, 1, item) # Записываем в таблицу ИСХ тип датчика
            match str(sensor[0][4]):
                case "ЭТС":
                    sensor_type="ets"
                case "ППО":
                    sensor_type="ppo"
                case "ПРО":
                    sensor_type="pro"
                case _:
                    sensor_type="error"
            coef_set = self.db_query('SELECT value FROM options WHERE option="coef_set_ish_'+sensor_type+'"')
            if sensor_type == 'ets':
                coef_set[0] = list(coef_set[0])
                coef_set[0][0] += ',δ'
            group_name = "self.groupBox_coef_set_ish_"+sensor_type+"_"+str(row)
            exec(group_name+" = QtWidgets.QGroupBox(self.Settings_ISH)")
            exec(group_name+".setGeometry(QtCore.QRect(0, 0, "+str((coef_set[0][0].count(',')+1)*40)+", 40))")
            exec(group_name+".setTitle('')")
            exec(group_name+".setObjectName('"+group_name[5:]+"')")
            item = QtWidgets.QTableWidgetItem()
            item.setFlags(Qt.ItemFlag.ItemIsSelectable)
            self.tableWidget_settings_ISH.setItem(row, 2, item)
            self.tableWidget_settings_ISH.setCellWidget(row, 2, eval(group_name))
            i = 0
            for coef in coef_set[0][0].split(","):
                button_name = "self.pushButton_"+sensor[0][4]+"_"+coef+"_"+str(sensor_ish_data[row][0])
                exec(button_name+" = QtWidgets.QPushButton("+group_name+")")
                exec(button_name+".setFont(font)")
                exec(button_name+".setGeometry(QtCore.QRect("+str(i*40)+", 0, 40, 40))")
                i+=1
                exec(button_name+".setObjectName('"+button_name[5:]+"')")
                exec(button_name+".setText('"+coef+"')")
                if coef != 'δ':
                    coef_val = self.db_query(f"SELECT {coef} FROM ish_data WHERE sensor_id={str(sensor_ish_data[row][0])}")
                    exec(button_name+".clicked.connect(self.coef_edit)")
                else:
                    exec(button_name+".clicked.connect(self.confidence_error)")
                if (coef_val[0][0] is not None and coef_val[0][0] != 0):
                    exec(button_name+".setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(38, 0, 255);')")
 
            # brush = QtGui.QBrush(QtGui.QColor(0, 0, 255))
            # brush.setStyle(QtCore.Qt.BrushStyle.NoBrush)
            # brush = QtGui.QBrush(QtGui.QColor(0, 0, 0))
            # brush.setStyle(QtCore.Qt.BrushStyle.NoBrush)
            # item.setBackground(brush)

        # item = QtWidgets.QTableWidgetItem()
        # self.tableWidget_settings_ISH.setItem(0, 1, item)
        # item = QtWidgets.QTableWidgetItem()
        # self.tableWidget_settings_ISH.setItem(1, 1, item)
        # item = QtWidgets.QTableWidgetItem()
        # brush = QtGui.QBrush(QtGui.QColor(0, 0, 0))
        # brush.setStyle(QtCore.Qt.BrushStyle.NoBrush)
        # item.setBackground(brush)
        # brush = QtGui.QBrush(QtGui.QColor(0, 0, 0))
        # brush.setStyle(QtCore.Qt.BrushStyle.NoBrush)
        # item.setForeground(brush)
        # self.tableWidget_settings_ISH.setItem(3, 1, item)
        # self.tableWidget_settings_ISH.horizontalHeader().setVisible(False)
        # self.tableWidget_settings_ISH.horizontalHeader().setStretchLastSection(False)
        
        self.pushButton_ish_back = QtWidgets.QPushButton(self.Settings_ISH)
        self.pushButton_ish_back.setGeometry(QtCore.QRect(414, 550, 141, 31))
        self.pushButton_ish_back.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_ish_back.setObjectName("pushButton_back")
        self.pushButton_ish_back.clicked.connect(self.back_tab)
        self.tab_settings_ish = self.stackedWidget_main.addWidget(self.Settings_ISH)


        #   -----------------------------------------------------------------
        #   вкладка Датчики
        #   -----------------------------------------------------------------
        self.Sensors = QtWidgets.QWidget()
        self.Sensors.setObjectName("Sensors")

        #   Таблица датчиков
   
        self.tableWidget_sensors = QtWidgets.QTableWidget(self.Sensors)
        self.tableWidget_sensors.setGeometry(QtCore.QRect(0, 0, 971, 541))
        self.tableWidget_sensors.setObjectName("tableWidget_sensors")
        self.tableWidget_sensors.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows);
        self.tableWidget_sensors.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection);
        columnWidths = [50, 178, 78, 98, 100, 130]
        columnCount = len(columnWidths)
        self.tableWidget_sensors.setColumnCount(columnCount)
        rowCount = self.db_query('SELECT COUNT(*) FROM sensors')[0][0]
        self.tableWidget_sensors.setRowCount(rowCount)
        for column in range(columnCount):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_sensors.setHorizontalHeaderItem(column, item)
            self.tableWidget_sensors.setColumnWidth(column, columnWidths[column])
        self.tableWidget_sensors.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        sensor=self.db_query('SELECT * FROM sensors')
        for row in range(rowCount):
            font.setPointSize(16)
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            item.setText(str(row+1))
            self.tableWidget_sensors.setVerticalHeaderItem(row, item)
            font.setPointSize(14)
            for column in range (columnCount):
                item = QtWidgets.QTableWidgetItem(str(sensor[row][column+1]))
                item.setFont(font)
                item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                self.tableWidget_sensors.setItem(row, column, item)
        self.pushButton_sensors_back = QtWidgets.QPushButton(self.Sensors)
        self.pushButton_sensors_back.setGeometry(QtCore.QRect(414, 550, 141, 31))
        self.pushButton_sensors_back.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_sensors_back.setObjectName("pushButton_sensors_back")
        self.pushButton_sensors_back.clicked.connect(self.back_tab)
        self.pushButton_sensor_add = QtWidgets.QPushButton(self.Sensors)
        self.pushButton_sensor_add.setGeometry(QtCore.QRect(50, 550, 141, 31))
        self.pushButton_sensor_add.setObjectName("pushButton_sensor_add")
        self.pushButton_sensor_add.clicked.connect(lambda: self.add_sensor_to_list(-1))
        self.tableWidget_sensors.cellClicked.connect(self.add_sensor_to_list)
        self.tab_sensors = self.stackedWidget_main.addWidget(self.Sensors)
        #   -----------------------------------------------------------------
        #                       вкладка РУЧНОЕ ИЗМЕРЕНИЕ
        #   -----------------------------------------------------------------
        self.Manual = QtWidgets.QWidget()
        self.Manual.setObjectName("Manual")
        self.pushButton_start_stop_manual = QtWidgets.QPushButton(self.Manual)
        self.pushButton_start_stop_manual.setGeometry(QtCore.QRect(60, 530, 191, 51))
        self.pushButton_start_stop_manual.clicked.connect(self.start_stop)
        font = QtGui.QFont()
        font.setPointSize(18)
        self.pushButton_start_stop_manual.setFont(font)
        self.pushButton_start_stop_manual.setStyleSheet("background-color: rgb(31, 100, 10); color: rgb(255, 255, 255);")
        self.pushButton_start_stop_manual.setObjectName("pushButton_start_stop_manual")
        self.groupBox_manual_path = QtWidgets.QGroupBox(self.Manual)
        self.groupBox_manual_path.setGeometry(QtCore.QRect(350, 530, 571, 51))
        self.groupBox_manual_path.setTitle("")
        self.groupBox_manual_path.setObjectName("groupBox_manual_path")
        self.pushButton_manual_channels_settings = QtWidgets.QPushButton(self.groupBox_manual_path)
        self.pushButton_manual_channels_settings.setGeometry(QtCore.QRect(0, 0, 191, 51))
        font.setPointSize(10)
        self.pushButton_manual_channels_settings.setFont(font)
        self.pushButton_manual_channels_settings.setObjectName("pushButton_manual_channels_settings")
        self.pushButton_manual_channels_settings.clicked.connect(lambda :self.stackedWidget_manual_measuring_tabs.setCurrentIndex(0))
        self.pushButton_manual_measuring = QtWidgets.QPushButton(self.groupBox_manual_path)
        self.pushButton_manual_measuring.setGeometry(QtCore.QRect(190, 0, 191, 51))
        self.pushButton_manual_measuring.setFont(font)
        self.pushButton_manual_measuring.setObjectName("pushButton_manual_measuring")
        self.pushButton_manual_measuring.clicked.connect(lambda :self.stackedWidget_manual_measuring_tabs.setCurrentIndex(1))
        self.pushButton_manual_result = QtWidgets.QPushButton(self.groupBox_manual_path)
        self.pushButton_manual_result.setGeometry(QtCore.QRect(380, 0, 191, 51))
        self.pushButton_manual_result.setFont(font)
        self.pushButton_manual_result.setObjectName("pushButton_manual_result")
        self.pushButton_manual_result.clicked.connect(lambda :self.stackedWidget_manual_measuring_tabs.setCurrentIndex(2))
        
        self.stackedWidget_manual_measuring_tabs = QtWidgets.QStackedWidget(self.Manual)
        self.stackedWidget_manual_measuring_tabs.setGeometry(QtCore.QRect(0, 0, 971, 521))
        self.stackedWidget_manual_measuring_tabs.setObjectName("stackedWidget_manual_measuring_tabs")

        #   -----------------------------------------------------------------
        #   Раздел Настройка в ручном режиме.
        #   -----------------------------------------------------------------

        self.channels_settings_tab = QtWidgets.QWidget()
        self.channels_settings_tab.setObjectName("channels_settings_tab")

        #   -----------------------------------------------------------------
        #   Таблица уставок
        #   -----------------------------------------------------------------

        self.tableWidget_ustavka_manual = QtWidgets.QTableWidget(self.channels_settings_tab)
        self.tableWidget_ustavka_manual.setGeometry(QtCore.QRect(0, 0, 141, 521))
        self.tableWidget_ustavka_manual.setObjectName("tableWidget_ustavka_manual")
        self.tableWidget_ustavka_manual.setColumnCount(2)
        self.tableWidget_ustavka_manual.setRowCount(1)
        for column in range(2):
            item = QtWidgets.QTableWidgetItem()
            self.tableWidget_ustavka_manual.setHorizontalHeaderItem(column, item)
            item = QtWidgets.QTableWidgetItem()
            item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
            item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.tableWidget_ustavka_manual.setItem(0, column, item)
        self.tableWidget_ustavka_manual.cellClicked.connect(self.add_ustavka_to_tab)
        self.tableWidget_ustavka_manual.horizontalHeader().setDefaultSectionSize(60)

        #   -----------------------------------------------------------------
        #   Таблица каналов
        #   -----------------------------------------------------------------

        self.tableWidget_manual_channels_settings = QtWidgets.QTableWidget(self.channels_settings_tab)
        self.tableWidget_manual_channels_settings.setGeometry(QtCore.QRect(140, 0, 831, 521))
        self.tableWidget_manual_channels_settings.setObjectName("tableWidget_manual_channels_settings")
        self.tableWidget_manual_channels_settings.cellClicked.connect(self.add_sensor_to_measuring_tab)
        columnWidths = [50, 120, 78, 98, 100, 130, 78]
        columnCount = len(columnWidths)
        rowCount = 8
        self.tableWidget_manual_channels_settings.setColumnCount(columnCount)
        self.tableWidget_manual_channels_settings.setRowCount(rowCount)
        font.setPointSize(13)
        for row in range(rowCount):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_manual_channels_settings.setVerticalHeaderItem(row, item)
            for column in range(columnCount):
                if row == 0:
                    item = QtWidgets.QTableWidgetItem()
                    item.setFont(font)
                    self.tableWidget_manual_channels_settings.setHorizontalHeaderItem(column, item)
                    self.tableWidget_manual_channels_settings.setColumnWidth(column, columnWidths[column])
                if column == 0:
                    item = QtWidgets.QTableWidgetItem("+")
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                    item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                else:
                    item = QtWidgets.QTableWidgetItem("")
                    item.setFlags(Qt.ItemFlag.NoItemFlags)
                font.setPointSize(14)
                item.setFont(font)
                self.tableWidget_manual_channels_settings.setItem(row, column, item)
        self.tableWidget_manual_channels_settings.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.stackedWidget_manual_measuring_tabs.addWidget(self.channels_settings_tab)

        #   -----------------------------------------------------------------
        #   Раздел Измерение
        #   -----------------------------------------------------------------

        self.measuring_tab = QtWidgets.QWidget()
        self.measuring_tab.setObjectName("measuring_tab")
        self.tableWidget_progress_bar_manual = QtWidgets.QTableWidget(self.measuring_tab)
        self.tableWidget_progress_bar_manual.setGeometry(QtCore.QRect(0, 0, 921, 21))
        self.tableWidget_progress_bar_manual.setObjectName("tableWidget_progress_bar_manual")
        self.tableWidget_progress_bar_manual.setColumnCount(0)
        self.tableWidget_progress_bar_manual.setRowCount(1)
        self.tableWidget_progress_bar_manual.verticalHeader().hide()
        self.tableWidget_progress_bar_manual.horizontalHeader().hide()
        self.tableWidget_progress_bar_manual.horizontalHeader().setStretchLastSection(True)
        self.tableWidget_progress_bar_manual.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.tableWidget_progress_bar_manual.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget_progress_bar_manual.setVerticalHeaderItem(0, item)

        #   -----------------------------------------------------------------
        #   Переключатель график\таблица
        #   -----------------------------------------------------------------

        self.measuring_view = QtWidgets.QStackedWidget(self.measuring_tab)
        self.measuring_view.setGeometry(QtCore.QRect(0, 30, 981, 491))
        self.measuring_view.setObjectName("measuring_view")
        self.graph = QtWidgets.QWidget()
        self.graph.setObjectName("graph")
        self.groupBox_channels_buttons = QtWidgets.QGroupBox(self.graph)
        self.groupBox_channels_buttons.setGeometry(QtCore.QRect(0, 0, 110, 441))
        # self.groupBox_channels_buttons.setStyleSheet('border-style: none')
        self.groupBox_channels_buttons.setObjectName("groupBox_channels_buttons")
        font = QtGui.QFont()
        self.color_buttons = ["24, 181, 59", "130, 109, 255", "84, 114, 47", "242, 80, 80",
                        "0, 255, 0", "224, 192, 79", "43, 209, 199", "60, 84, 119", "0, 0, 255"]
        for i in range (9):
            exec(f"self.pushButton_{str(i+1)} = QtWidgets.QPushButton(self.groupBox_channels_buttons)")
            exec(f"self.pushButton_{str(i+1)}.setGeometry(QtCore.QRect(0, "+str(i*50)+", 41, 41))")
            font.setPointSize(20)
            exec("self.pushButton_"+str(i+1)+".setFont(font)")
            if i != 8: exec("self.pushButton_"+str(i+1)+".setEnabled(False)")
            exec("self.pushButton_"+str(i+1)+".setCheckable(True)")
            exec("self.pushButton_"+str(i+1)+".setStyleSheet('color: rgb(255, 255, 255); background-color: rgb("+self.color_buttons[i]+")')")
            exec("self.pushButton_"+str(i+1)+".setObjectName('pushButton_"+str(i+1)+"')")
            exec("self.pushButton_"+str(i+1)+".clicked.connect(self.graph_btn_toggle)")
            exec("self.label_temp_of_chan_"+str(i+1)+"= QtWidgets.QLabel(self.groupBox_channels_buttons)")
            exec("self.label_temp_of_chan_"+str(i+1)+".setObjectName('label_temp_of_chan_"+str(i+1)+"')")
            exec("self.label_temp_of_chan_"+str(i+1)+".setGeometry(QtCore.QRect(45, "+str(i*50)+", 65, 41))")
            font.setPointSize(14)
            exec("self.label_temp_of_chan_"+str(i+1)+".setFont(font)")
        self.pushButton_9.setChecked(True)
        self.graphWidget = pg.PlotWidget(self.graph, axisItems={'bottom': pg.DateAxisItem()}) # pg.DateAxisItem(utcOffset=0)
        self.graphWidget.setGeometry(QtCore.QRect(110, 0, 811, 491))
        self.graphWidget.setObjectName("graphWidget")
        self.graphWidget.setBackground('w')
        self.graphWidget.showGrid(x=True, y=True)
        self.groupBox_navigate_buttons = QtWidgets.QGroupBox(self.graph)
        self.groupBox_navigate_buttons.setGeometry(QtCore.QRect(930, 0, 41, 441))
        self.groupBox_navigate_buttons.setObjectName("groupBox_navigate_buttons")
        buttons = ["move_graph_left", "move_graph_right", "move_graph_up", "move_graph_down", "scale_vertical_up", "scale_vertical_down", "scale_horizontal_up", "scale_horizontal_down", "scale_auto"]
        font = QtGui.QFont()
        font.setPointSize(12)
        for i in range(len(buttons)):
            exec("self.pushButton_"+buttons[i]+" = QtWidgets.QPushButton(self.groupBox_navigate_buttons)")
            exec("self.pushButton_"+buttons[i]+".setGeometry(QtCore.QRect(0, "+str(i*50)+", 41, 41))")
            exec("self.pushButton_"+buttons[i]+".setFont(font)")
            # self.pushButton_left.setStyleSheet("")
            exec("self.pushButton_"+buttons[i]+".setObjectName('pushButton_"+buttons[i]+"')")
            exec("self.pushButton_"+buttons[i]+".clicked.connect(self."+buttons[i]+")")

        self.pushButton_view_table = QtWidgets.QPushButton(self.graph)
        self.pushButton_view_table.setGeometry(QtCore.QRect(930, 450, 41, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.pushButton_view_table.setFont(font)
        self.pushButton_view_table.setStyleSheet("")
        self.pushButton_view_table.setObjectName("pushButton_view_table")
        self.pushButton_view_table.clicked.connect(lambda :self.measuring_view.setCurrentIndex(1))
        self.measuring_view.addWidget(self.graph)

        #   -----------------------------------------------------------------
        #   Таблица показаний в ходе ручного измерения.
        #   -----------------------------------------------------------------

        self.table = QtWidgets.QWidget()
        self.table.setObjectName("table")
        self.tableWidget_manual_measuring_result = QtWidgets.QTableWidget(self.table)
        self.tableWidget_manual_measuring_result.setGeometry(QtCore.QRect(0, 0, 921, 212))
        self.tableWidget_manual_measuring_result.setObjectName("tableWidget")
        self.tableWidget_manual_measuring_result.horizontalHeader().hide()
        self.tableWidget_manual_measuring_result.setColumnCount(8)
        self.tableWidget_manual_measuring_result.setRowCount(7)
        for row in range(7):
            item = QtWidgets.QTableWidgetItem()
            self.tableWidget_manual_measuring_result.setVerticalHeaderItem(row, item)
        for col in range(8):
            item = QtWidgets.QTableWidgetItem(str(col+1))
            self.tableWidget_manual_measuring_result.setItem(0, col, item)
            for row in range(6):
                item = QtWidgets.QTableWidgetItem()
                self.tableWidget_manual_measuring_result.setItem(row+1, col, item)
        self.tableWidget_manual_measuring_result.horizontalHeader().setDefaultSectionSize(92)
        
        self.textEdit_log_manual = QtWidgets.QTextEdit(self.table)
        # self.textEdit_log_auto.setEnabled(False)
        self.textEdit_log_manual.setReadOnly(True)
        self.textEdit_log_manual.setGeometry(QtCore.QRect(0, 213, 921, 280))
        self.textEdit_log_manual.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.textEdit_log_manual.setObjectName("textEdit_log_manual")

        self.pushButton_view_graph = QtWidgets.QPushButton(self.table)
        self.pushButton_view_graph.setGeometry(QtCore.QRect(930, 450, 41, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.pushButton_view_graph.setFont(font)
        self.pushButton_view_graph.setStyleSheet("")
        self.pushButton_view_graph.setObjectName("pushButton_view_graph")
        self.pushButton_view_graph.clicked.connect(lambda :self.measuring_view.setCurrentIndex(0))
        self.measuring_view.addWidget(self.table)
        self.stackedWidget_manual_measuring_tabs.addWidget(self.measuring_tab)
        self.tab_manual = self.stackedWidget_main.addWidget(self.Manual)

        #   -----------------------------------------------------------------
        #   Раздел Результат
        #   -----------------------------------------------------------------

        self.manual_result_tab = QtWidgets.QWidget()
        self.manual_result_tab.setObjectName("manual_result_tab")
        self.horizontalLayoutWidget_3 = QtWidgets.QWidget(self.manual_result_tab)
        self.horizontalLayoutWidget_3.setGeometry(QtCore.QRect(0, 0, 971, 521))
        self.horizontalLayoutWidget_3.setObjectName("horizontalLayoutWidget_3")
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_3)
        self.horizontalLayout_3.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        self.groupBox_manual_result = QtWidgets.QGroupBox(self.horizontalLayoutWidget_3)
        self.groupBox_manual_result.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.groupBox_manual_result.setTitle("")
        self.groupBox_manual_result.setObjectName("groupBox_manual_result")
        self.gridLayout_3 = QtWidgets.QGridLayout(self.groupBox_manual_result)
        self.gridLayout_3.setObjectName("gridLayout_3")
        self.label_manual_number = QtWidgets.QLabel(self.groupBox_manual_result)
        self.label_manual_number.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_number.setObjectName("label_manual_number")
        self.gridLayout_3.addWidget(self.label_manual_number, 8, 0, 1, 1)

        spacerItem3 = QtWidgets.QSpacerItem(454, 402, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_3.addItem(spacerItem3, 9, 0, 1, 2)
        self.label_manual_fio = QtWidgets.QLabel(self.groupBox_manual_result)
        self.label_manual_fio.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_fio.setObjectName("label_manual_fio")
        self.gridLayout_3.addWidget(self.label_manual_fio, 4, 0, 1, 1)
        self.label_manual_model = QtWidgets.QLabel(self.groupBox_manual_result)
        self.label_manual_model.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_model.setObjectName("label_manual_model")
        self.gridLayout_3.addWidget(self.label_manual_model, 7, 0, 1, 1)
        self.label_manual_megaommetr = QtWidgets.QLabel(self.groupBox_manual_result)
        self.label_manual_megaommetr.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_megaommetr.setObjectName("label_manual_megaommetr")
        self.gridLayout_3.addWidget(self.label_manual_megaommetr, 6, 0, 1, 1)
        self.label_manual_customer = QtWidgets.QLabel(self.groupBox_manual_result)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_manual_customer.setFont(font)
        self.label_manual_customer.setStyleSheet("")
        self.label_manual_customer.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.label_manual_customer.setObjectName("label_manual_customer")
        self.gridLayout_3.addWidget(self.label_manual_customer, 0, 0, 1, 2)
        self.label_manual_measurement_number_text = QtWidgets.QLabel(self.groupBox_manual_result)
        self.label_manual_measurement_number_text.setFont(font)
        self.label_manual_measurement_number_text.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_measurement_number_text.setObjectName("label_manual_number")
        self.gridLayout_3.addWidget(self.label_manual_measurement_number_text, 1, 0, 1, 1)
        self.label_manual_measurement_number = QtWidgets.QLabel(self.groupBox_manual_result)
        self.label_manual_measurement_number.setFont(font)
        self.label_manual_measurement_number.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_measurement_number.setObjectName("label_manual_number")
        self.gridLayout_3.addWidget(self.label_manual_measurement_number, 1, 1, 1, 1)
        self.lineEdit_manual_customer = QtWidgets.QLineEdit(self.groupBox_manual_result)
        self.lineEdit_manual_customer.setObjectName("lineEdit_manual_customer")
        self.gridLayout_3.addWidget(self.lineEdit_manual_customer, 3, 0, 1, 2)
        self.lineEdit_manual_operators_name = QtWidgets.QLineEdit(self.groupBox_manual_result)
        self.lineEdit_manual_operators_name.setObjectName("lineEdit_manual_operators_name")
        self.gridLayout_3.addWidget(self.lineEdit_manual_operators_name, 5, 0, 1, 2)
        self.lineEdit_manual_model = QtWidgets.QLineEdit(self.groupBox_manual_result)
        self.lineEdit_manual_model.setObjectName("lineEdit_manual_model")
        self.gridLayout_3.addWidget(self.lineEdit_manual_model, 7, 1, 1, 1)
        self.lineEdit_manual_number = QtWidgets.QLineEdit(self.groupBox_manual_result)
        self.lineEdit_manual_number.setObjectName("lineEdit_manual_number")
        self.gridLayout_3.addWidget(self.lineEdit_manual_number, 8, 1, 1, 1)
        self.label_manual_customer_2 = QtWidgets.QLabel(self.groupBox_manual_result)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_manual_customer_2.setFont(font)
        self.label_manual_customer_2.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_customer_2.setObjectName("label_manual_customer_2")
        self.gridLayout_3.addWidget(self.label_manual_customer_2, 2, 0, 1, 2)
        self.horizontalLayout_3.addWidget(self.groupBox_manual_result)
        self.groupBox_manual_measuring_conditions = QtWidgets.QGroupBox(self.horizontalLayoutWidget_3)
        self.groupBox_manual_measuring_conditions.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.groupBox_manual_measuring_conditions.setTitle("")
        self.groupBox_manual_measuring_conditions.setObjectName("groupBox_manual_measuring_conditions")
        self.gridLayout_4 = QtWidgets.QGridLayout(self.groupBox_manual_measuring_conditions)
        self.gridLayout_4.setObjectName("gridLayout_4")
        self.label_manual_t_atm = QtWidgets.QLabel(self.groupBox_manual_measuring_conditions)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_manual_t_atm.sizePolicy().hasHeightForWidth())
        self.label_manual_t_atm.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_manual_t_atm.setFont(font)
        self.label_manual_t_atm.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_t_atm.setObjectName("label_manual_t_atm")
        self.gridLayout_4.addWidget(self.label_manual_t_atm, 1, 0, 1, 1)
        self.lineEdit_manual_temp = QtWidgets.QLineEdit(self.groupBox_manual_measuring_conditions)
        self.lineEdit_manual_temp.setText("")
        self.lineEdit_manual_temp.setObjectName("lineEdit_manual_temp")
        self.gridLayout_4.addWidget(self.lineEdit_manual_temp, 1, 1, 1, 1)
        self.label_manual_p_atm = QtWidgets.QLabel(self.groupBox_manual_measuring_conditions)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_manual_p_atm.sizePolicy().hasHeightForWidth())
        self.label_manual_p_atm.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_manual_p_atm.setFont(font)
        self.label_manual_p_atm.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_p_atm.setObjectName("label_manual_p_atm")
        self.gridLayout_4.addWidget(self.label_manual_p_atm, 2, 0, 1, 1)
        self.lineEdit_manual_pressure = QtWidgets.QLineEdit(self.groupBox_manual_measuring_conditions)
        self.lineEdit_manual_pressure.setText("")
        self.lineEdit_manual_pressure.setObjectName("lineEdit_manual_pressure")
        self.gridLayout_4.addWidget(self.lineEdit_manual_pressure, 2, 1, 1, 1)
        self.label_manual_hydro = QtWidgets.QLabel(self.groupBox_manual_measuring_conditions)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_manual_hydro.sizePolicy().hasHeightForWidth())
        self.label_manual_hydro.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_manual_hydro.setFont(font)
        self.label_manual_hydro.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_manual_hydro.setObjectName("label_manual_hydro")
        self.gridLayout_4.addWidget(self.label_manual_hydro, 4, 0, 1, 1)
        spacerItem4 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_4.addItem(spacerItem4, 5, 0, 1, 2)
        self.lineEdit_manual_hydro = QtWidgets.QLineEdit(self.groupBox_manual_measuring_conditions)
        self.lineEdit_manual_hydro.setText("")
        self.lineEdit_manual_hydro.setObjectName("lineEdit_manual_hydro")
        self.gridLayout_4.addWidget(self.lineEdit_manual_hydro, 4, 1, 1, 1)
        self.label_manual_article_name = QtWidgets.QLabel(self.groupBox_manual_measuring_conditions)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_manual_article_name.setFont(font)
        self.label_manual_article_name.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.label_manual_article_name.setObjectName("label_manual_article_name")
        self.gridLayout_4.addWidget(self.label_manual_article_name, 0, 0, 1, 2)
        self.horizontalLayout_3.addWidget(self.groupBox_manual_measuring_conditions)
        self.pushButton_save_data_manual = QtWidgets.QPushButton(self.groupBox_manual_measuring_conditions)
        self.pushButton_save_data_manual.setGeometry(QtCore.QRect(100, 170, 201, 41))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.pushButton_save_data_manual.setFont(font)
        self.pushButton_save_data_manual.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(163, 194, 194);")
        self.pushButton_save_data_manual.setObjectName("pushButton_save_measurement_data")
        self.pushButton_save_data_manual.setEnabled(False)
        self.pushButton_save_data_manual.clicked.connect(self.save_measurement_data)
        self.gridLayout_4.addWidget(self.pushButton_save_data_manual)
        self.stackedWidget_manual_measuring_tabs.addWidget(self.manual_result_tab)

        #   -----------------------------------------------------------------
        #                       вкладка АВТОМАТИЧЕСКОЕ ИЗМЕРЕНИЕ
        #   -----------------------------------------------------------------

        # Кнопки разделов измерения

        self.Auto = QtWidgets.QWidget()
        self.Auto.setObjectName("Auto")
        self.groupBox_bottom_buttons = QtWidgets.QGroupBox(self.Auto)
        self.groupBox_bottom_buttons.setGeometry(QtCore.QRect(30, 490, 901, 51))
        self.groupBox_bottom_buttons.setObjectName("groupBox_bottom_buttons")
        self.pushButton_measuring_type = QtWidgets.QPushButton(self.groupBox_bottom_buttons)
        self.pushButton_measuring_type.setGeometry(QtCore.QRect(0, 0, 180, 51))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.pushButton_measuring_type.setFont(font)
        self.pushButton_measuring_type.setStyleSheet("color: rgb(154, 154, 154);")
        self.pushButton_measuring_type.setObjectName("pushButton_measuring_type")
        self.pushButton_measuring_type.clicked.connect(lambda :self.stackedWidget_auto_measuring_tabs.setCurrentIndex(0))
        self.pushButton_tpoints_sensors = QtWidgets.QPushButton(self.groupBox_bottom_buttons)
        self.pushButton_tpoints_sensors.setGeometry(QtCore.QRect(180, 0, 180, 51))
        self.pushButton_tpoints_sensors.setFont(font)
        self.pushButton_tpoints_sensors.setStyleSheet("color: rgb(154, 154, 154);")
        self.pushButton_tpoints_sensors.setObjectName("pushButton_tpoints_sensors")
        self.pushButton_tpoints_sensors.clicked.connect(lambda :self.stackedWidget_auto_measuring_tabs.setCurrentIndex(1))
        self.pushButton_measuring_settings = QtWidgets.QPushButton(self.groupBox_bottom_buttons)
        self.pushButton_measuring_settings.setGeometry(QtCore.QRect(360, 0, 180, 51))
        self.pushButton_measuring_settings.setFont(font)
        self.pushButton_measuring_settings.setStyleSheet("color: rgb(154, 154, 154);")
        self.pushButton_measuring_settings.setObjectName("pushButton_measuring_settings")
        self.pushButton_measuring_settings.clicked.connect(lambda :self.stackedWidget_auto_measuring_tabs.setCurrentIndex(2))
        self.pushButton_start_stop_tab = QtWidgets.QPushButton(self.groupBox_bottom_buttons)
        self.pushButton_start_stop_tab.setGeometry(QtCore.QRect(540, 0, 180, 51))
        self.pushButton_start_stop_tab.setFont(font)
        self.pushButton_start_stop_tab.setStyleSheet("color: rgb(154, 154, 154);")
        self.pushButton_start_stop_tab.setObjectName("pushButton_start_stop_tab")
        self.pushButton_start_stop_tab.clicked.connect(lambda :self.stackedWidget_auto_measuring_tabs.setCurrentIndex(3))
        self.pushButton_protocol = QtWidgets.QPushButton(self.groupBox_bottom_buttons)
        self.pushButton_protocol.setGeometry(QtCore.QRect(720, 0, 180, 51))
        self.pushButton_protocol.setFont(font)
        self.pushButton_protocol.setStyleSheet("color: rgb(154, 154, 154);")
        self.pushButton_protocol.setObjectName("pushButton_protocol")
        self.pushButton_protocol.setEnabled(False)
        self.pushButton_protocol.clicked.connect(lambda :self.stackedWidget_auto_measuring_tabs.setCurrentIndex(4))

        #   -----------------------------------------------------------------
        #   Раздел Тип измерения в автоматическом режиме.
        #   -----------------------------------------------------------------

        self.stackedWidget_auto_measuring_tabs = QtWidgets.QStackedWidget(self.Auto)
        self.stackedWidget_auto_measuring_tabs.setGeometry(QtCore.QRect(0, 0, 971, 481))
        self.stackedWidget_auto_measuring_tabs.setObjectName("stackedWidget_auto_measuring_tabs")
        self.measuring_type_tab = QtWidgets.QWidget()
        self.measuring_type_tab.setObjectName("measuring_type_tab")
        self.pushButton_verification_TS = QtWidgets.QPushButton(self.measuring_type_tab)
        self.pushButton_verification_TS.setGeometry(QtCore.QRect(520, 160, 341, 71))
        font = QtGui.QFont()
        font.setPointSize(16)
        self.pushButton_verification_TS.setFont(font)
        self.pushButton_verification_TS.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_verification_TS.setObjectName("pushButton_verification_TS")
        self.pushButton_verification_TS.setCheckable(True)
        self.pushButton_verification_TS.clicked.connect(self.set_auto_verification_TS_mode)
        self.pushButton_graduation_TS = QtWidgets.QPushButton(self.measuring_type_tab)
        self.pushButton_graduation_TS.setGeometry(QtCore.QRect(520, 270, 341, 71))
        self.pushButton_graduation_TS.setFont(font)
        self.pushButton_graduation_TS.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_graduation_TS.setObjectName("pushButton_graduation_TS")
        self.pushButton_graduation_TS.setCheckable(True)
        self.pushButton_graduation_TS.clicked.connect(self.set_auto_graduation_TS_mode)
        self.pushButton_calibration_TP = QtWidgets.QPushButton(self.measuring_type_tab)
        self.pushButton_calibration_TP.setGeometry(QtCore.QRect(100, 270, 341, 71))
        self.pushButton_calibration_TP.setFont(font)
        self.pushButton_calibration_TP.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_calibration_TP.setObjectName("pushButton_calibration_TP")
        self.pushButton_calibration_TP.setCheckable(True)
        self.pushButton_calibration_TP.clicked.connect(self.set_auto_calibration_TP_mode)
        self.pushButton_verification_TP = QtWidgets.QPushButton(self.measuring_type_tab)
        self.pushButton_verification_TP.setGeometry(QtCore.QRect(100, 160, 341, 71))
        self.pushButton_verification_TP.setFont(font)
        self.pushButton_verification_TP.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_verification_TP.setObjectName("pushButton_verification_TP")
        self.pushButton_verification_TP.setCheckable(True)
        self.pushButton_verification_TP.clicked.connect(self.set_auto_verification_TP_mode)
        self.stackedWidget_auto_measuring_tabs.addWidget(self.measuring_type_tab)

        self.tpoints_sensors_tab = QtWidgets.QWidget()
        self.tpoints_sensors_tab.setObjectName("tpoints_sensors_tab")
        self.stackedWidget_measuring_types_tables = QtWidgets.QStackedWidget(self.tpoints_sensors_tab)
        self.stackedWidget_measuring_types_tables.setGeometry(QtCore.QRect(60, 0, 911, 481))
        self.stackedWidget_measuring_types_tables.setObjectName("stackedWidget_measuring_types_tables")

        #   -----------------------------------------------------------------
        #   Подраздел измерения термопар в автоматическом режиме.
        #   -----------------------------------------------------------------
        
        self.TP = QtWidgets.QWidget()
        self.TP.setObjectName("TP")
        self.tableWidget_auto_channels_settings_TP = QtWidgets.QTableWidget(self.TP)
        self.tableWidget_auto_channels_settings_TP.setGeometry(QtCore.QRect(0, 0, 911, 481))
        self.tableWidget_auto_channels_settings_TP.setObjectName("tableWidget_auto_channels_settings_TP")
        self.tableWidget_auto_channels_settings_TP.cellClicked.connect(self.add_sensor_to_measuring_tab)
        columnWidths = [178, 100, 100, 70, 70, 80, 100, 70, 130, 130, 130]
        columnCount = len(columnWidths)
        rowCount = 8
        self.tableWidget_auto_channels_settings_TP.setColumnCount(columnCount)
        self.tableWidget_auto_channels_settings_TP.setRowCount(rowCount)
        font.setPointSize(10)
        for row in range(rowCount):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_auto_channels_settings_TP.setVerticalHeaderItem(row, item)
            for column in range(columnCount):
                if row == 0:
                    item = QtWidgets.QTableWidgetItem()
                    item.setFont(font)
                    self.tableWidget_auto_channels_settings_TP.setHorizontalHeaderItem(column, item)
                    self.tableWidget_auto_channels_settings_TP.setColumnWidth(column, columnWidths[column])
                if column == 0:
                    if row == 0:
                        item = QtWidgets.QTableWidgetItem("Перемычка")
                        item.setFlags(Qt.ItemFlag.NoItemFlags)
                    else:
                        item = QtWidgets.QTableWidgetItem("+")
                        item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                else:
                    item = QtWidgets.QTableWidgetItem("")
                    item.setFlags(Qt.ItemFlag.NoItemFlags)
                item.setFont(font)
                self.tableWidget_auto_channels_settings_TP.setItem(row, column, item)
        self.tableWidget_auto_channels_settings_TP.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.stackedWidget_measuring_types_tables.addWidget(self.TP)

        #   -----------------------------------------------------------------
        #   Подраздел поверки термоспоротивлений в автоматическом режиме.
        #   -----------------------------------------------------------------

        self.TS_verifications = QtWidgets.QWidget()
        self.TS_verifications.setObjectName("TS_verifications")
        self.tableWidget_auto_channels_settings_TS_verifications = QtWidgets.QTableWidget(self.TS_verifications)
        self.tableWidget_auto_channels_settings_TS_verifications.setGeometry(QtCore.QRect(0, 0, 911, 481))
        self.tableWidget_auto_channels_settings_TS_verifications.setObjectName("tableWidget_auto_channels_settings_TS_verifications")
        self.tableWidget_auto_channels_settings_TS_verifications.cellClicked.connect(self.add_sensor_to_measuring_tab)
        columnWidths = [178, 100, 100, 70, 70, 80, 100, 70, 130, 130]
        columnCount = len(columnWidths)
        rowCount = 8
        self.tableWidget_auto_channels_settings_TS_verifications.setColumnCount(columnCount)
        self.tableWidget_auto_channels_settings_TS_verifications.setRowCount(rowCount)
        font.setPointSize(10)
        for row in range(rowCount):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_auto_channels_settings_TS_verifications.setVerticalHeaderItem(row, item)
            for column in range(columnCount):
                if row == 0:
                    item = QtWidgets.QTableWidgetItem()
                    item.setFont(font)
                    self.tableWidget_auto_channels_settings_TS_verifications.setHorizontalHeaderItem(column, item)
                    self.tableWidget_auto_channels_settings_TS_verifications.setColumnWidth(column, columnWidths[column])
                if column == 0:
                    item = QtWidgets.QTableWidgetItem("+")
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                    item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)        
                else:
                    item = QtWidgets.QTableWidgetItem("")
                    item.setFlags(Qt.ItemFlag.NoItemFlags)    
                item.setFont(font)
                self.tableWidget_auto_channels_settings_TS_verifications.setItem(row, column, item)
        # self.tableWidget_auto_channels_settings_TS_verifications.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.stackedWidget_measuring_types_tables.addWidget(self.TS_verifications)

        #   -----------------------------------------------------------------
        #   Подраздел градуировки термосопротивлений в автоматическом режиме.
        #   -----------------------------------------------------------------

        self.TS_gradiations = QtWidgets.QWidget()
        self.TS_gradiations.setObjectName("TS_gradiations")
        self.tableWidget_auto_channels_settings_TS_gradiations = QtWidgets.QTableWidget(self.TS_gradiations)
        self.tableWidget_auto_channels_settings_TS_gradiations.setGeometry(QtCore.QRect(0, 0, 911, 481))
        self.tableWidget_auto_channels_settings_TS_gradiations.setObjectName("tableWidget_auto_channels_settings_TS_gradiations")
        self.tableWidget_auto_channels_settings_TS_gradiations.cellClicked.connect(self.add_sensor_to_measuring_tab)
        columnWidths = [178, 100, 100, 70, 70, 80, 100, 70, 130, 130]
        columnCount = len(columnWidths)
        rowCount = 8
        self.tableWidget_auto_channels_settings_TS_gradiations.setColumnCount(columnCount)
        self.tableWidget_auto_channels_settings_TS_gradiations.setRowCount(rowCount)
        font.setPointSize(10)
        for row in range(rowCount):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_auto_channels_settings_TS_gradiations.setVerticalHeaderItem(row, item)
            for column in range(columnCount):
                if row == 0:
                    item = QtWidgets.QTableWidgetItem()
                    item.setFont(font)
                    self.tableWidget_auto_channels_settings_TS_gradiations.setHorizontalHeaderItem(column, item)
                    self.tableWidget_auto_channels_settings_TS_gradiations.setColumnWidth(column, columnWidths[column])
                if column == 0:
                    item = QtWidgets.QTableWidgetItem("+")
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                    item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)        
                else:
                    item = QtWidgets.QTableWidgetItem("")
                    item.setFlags(Qt.ItemFlag.NoItemFlags)    
                item.setFont(font)
                self.tableWidget_auto_channels_settings_TS_gradiations.setItem(row, column, item)
        # self.tableWidget_auto_channels_settings_TS_gradiations.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.stackedWidget_measuring_types_tables.addWidget(self.TS_gradiations)

        #   Таблица температурных точек раздела Температурные точки и данные датчиков в автоматическом режиме.

        self.tableWidget_ustavka_auto = QtWidgets.QTableWidget(self.tpoints_sensors_tab)
        self.tableWidget_ustavka_auto.setGeometry(QtCore.QRect(0, 0, 61, 481))
        self.tableWidget_ustavka_auto.setObjectName("tableWidget_ustavka_auto")
        self.tableWidget_ustavka_auto.setColumnCount(1)
        self.tableWidget_ustavka_auto.setRowCount(1)
        item = QtWidgets.QTableWidgetItem()
        brush = QtGui.QBrush(QtGui.QColor(0, 0, 0))
        brush.setStyle(QtCore.Qt.BrushStyle.NoBrush)
        item.setForeground(brush)
        self.tableWidget_ustavka_auto.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        font = QtGui.QFont()
        font.setPointSize(16)
        item.setFont(font)
        self.tableWidget_ustavka_auto.setItem(0, 0, item)
        self.tableWidget_ustavka_auto.horizontalHeader().setDefaultSectionSize(50)
        self.tableWidget_ustavka_auto.cellClicked.connect(self.add_temp_to_tab)
        self.stackedWidget_auto_measuring_tabs.addWidget(self.tpoints_sensors_tab)

        #   -----------------------------------------------------------------
        #   Раздел Параметры измерения в автоматическом режиме.
        #   -----------------------------------------------------------------

        self.measuring_settings_tab = QtWidgets.QWidget()
        self.measuring_settings_tab.setObjectName("measuring_settings_tab")
        self.horizontalLayoutWidget = QtWidgets.QWidget(self.measuring_settings_tab)
        self.horizontalLayoutWidget.setGeometry(QtCore.QRect(0, 0, 971, 481))
        self.horizontalLayoutWidget.setObjectName("horizontalLayoutWidget")
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget)
        self.horizontalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.groupBox_measuring_data = QtWidgets.QGroupBox(self.horizontalLayoutWidget)
        self.groupBox_measuring_data.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.groupBox_measuring_data.setTitle("")
        self.groupBox_measuring_data.setObjectName("groupBox_measuring_data")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.groupBox_measuring_data)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.label_auto_number = QtWidgets.QLabel(self.groupBox_measuring_data)
        self.label_auto_number.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_number.setObjectName("label_auto_number")
        self.gridLayout_2.addWidget(self.label_auto_number, 8, 0, 1, 1)
        spacerItem1 = QtWidgets.QSpacerItem(454, 402, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_2.addItem(spacerItem1, 9, 0, 1, 2)
        self.label_auto_fio = QtWidgets.QLabel(self.groupBox_measuring_data)
        self.label_auto_fio.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_fio.setObjectName("label_auto_fio")
        self.gridLayout_2.addWidget(self.label_auto_fio, 4, 0, 1, 1)
        self.label_auto_model = QtWidgets.QLabel(self.groupBox_measuring_data)
        self.label_auto_model.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_model.setObjectName("label_auto_model")
        self.gridLayout_2.addWidget(self.label_auto_model, 7, 0, 1, 1)
        self.label_auto_megaommetr = QtWidgets.QLabel(self.groupBox_measuring_data)
        self.label_auto_megaommetr.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_megaommetr.setObjectName("label_auto_megaommetr")
        self.gridLayout_2.addWidget(self.label_auto_megaommetr, 6, 0, 1, 1)
        self.label_auto_customer = QtWidgets.QLabel(self.groupBox_measuring_data)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_auto_customer.setFont(font)
        self.label_auto_customer.setStyleSheet("")
        self.label_auto_customer.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.label_auto_customer.setObjectName("label_auto_customer")
        self.gridLayout_2.addWidget(self.label_auto_customer, 0, 0, 1, 2)
        self.label_auto_measurement_number_text = QtWidgets.QLabel(self.groupBox_measuring_data)
        self.label_auto_measurement_number_text.setFont(font)
        self.label_auto_measurement_number_text.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_measurement_number_text.setObjectName("label_manual_number")
        self.gridLayout_2.addWidget(self.label_auto_measurement_number_text, 1, 0, 1, 1)
        self.label_auto_measurement_number = QtWidgets.QLabel(self.groupBox_measuring_data)
        self.label_auto_measurement_number.setFont(font)
        self.label_auto_measurement_number.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_measurement_number.setObjectName("label_manual_number")
        self.gridLayout_2.addWidget(self.label_auto_measurement_number, 1, 1, 1, 1)
        self.lineEdit_auto_customer = QtWidgets.QLineEdit(self.groupBox_measuring_data)
        self.lineEdit_auto_customer.setObjectName("lineEdit_auto_customer")
        self.gridLayout_2.addWidget(self.lineEdit_auto_customer, 3, 0, 1, 2)
        self.lineEdit_auto_operators_name = QtWidgets.QLineEdit(self.groupBox_measuring_data)
        self.lineEdit_auto_operators_name.setObjectName("lineEdit_auto_operators_name")
        self.gridLayout_2.addWidget(self.lineEdit_auto_operators_name, 5, 0, 1, 2)
        self.lineEdit_auto_model = QtWidgets.QLineEdit(self.groupBox_measuring_data)
        self.lineEdit_auto_model.setObjectName("lineEdit_auto_model")
        self.gridLayout_2.addWidget(self.lineEdit_auto_model, 7, 1, 1, 1)
        self.lineEdit_auto_number = QtWidgets.QLineEdit(self.groupBox_measuring_data)
        self.lineEdit_auto_number.setObjectName("lineEdit_auto_number")
        self.gridLayout_2.addWidget(self.lineEdit_auto_number, 8, 1, 1, 1)
        self.label_auto_customer_2 = QtWidgets.QLabel(self.groupBox_measuring_data)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_auto_customer_2.setFont(font)
        self.label_auto_customer_2.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_customer_2.setObjectName("label_auto_customer_2")
        self.gridLayout_2.addWidget(self.label_auto_customer_2, 2, 0, 1, 2)
        self.horizontalLayout_2.addWidget(self.groupBox_measuring_data)
        self.groupBox_measuring_conditions = QtWidgets.QGroupBox(self.horizontalLayoutWidget)
        self.groupBox_measuring_conditions.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.groupBox_measuring_conditions.setTitle("")
        self.groupBox_measuring_conditions.setObjectName("groupBox_measuring_conditions")
        self.gridLayout = QtWidgets.QGridLayout(self.groupBox_measuring_conditions)
        self.gridLayout.setObjectName("gridLayout")
        self.label_auto_t_atm = QtWidgets.QLabel(self.groupBox_measuring_conditions)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_auto_t_atm.sizePolicy().hasHeightForWidth())
        self.label_auto_t_atm.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_auto_t_atm.setFont(font)
        self.label_auto_t_atm.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_t_atm.setObjectName("label_auto_t_atm")
        self.gridLayout.addWidget(self.label_auto_t_atm, 1, 0, 1, 1)
        self.lineEdit_auto_temp = QtWidgets.QLineEdit(self.groupBox_measuring_conditions)
        self.lineEdit_auto_temp.setText("")
        self.lineEdit_auto_temp.setObjectName("lineEdit_auto_temp")
        self.gridLayout.addWidget(self.lineEdit_auto_temp, 1, 1, 1, 1)
        self.label_auto_p_atm = QtWidgets.QLabel(self.groupBox_measuring_conditions)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_auto_p_atm.sizePolicy().hasHeightForWidth())
        self.label_auto_p_atm.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_auto_p_atm.setFont(font)
        self.label_auto_p_atm.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_p_atm.setObjectName("label_auto_p_atm")
        self.gridLayout.addWidget(self.label_auto_p_atm, 2, 0, 1, 1)
        self.lineEdit_auto_pressure = QtWidgets.QLineEdit(self.groupBox_measuring_conditions)
        self.lineEdit_auto_pressure.setText("")
        self.lineEdit_auto_pressure.setObjectName("lineEdit_auto_pressure")
        self.gridLayout.addWidget(self.lineEdit_auto_pressure, 2, 1, 1, 1)
        self.label_auto_hydro = QtWidgets.QLabel(self.groupBox_measuring_conditions)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.label_auto_hydro.sizePolicy().hasHeightForWidth())
        self.label_auto_hydro.setSizePolicy(sizePolicy)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_auto_hydro.setFont(font)
        self.label_auto_hydro.setStyleSheet("color: rgb(154, 154, 154);")
        self.label_auto_hydro.setObjectName("label_auto_hydro")
        self.gridLayout.addWidget(self.label_auto_hydro, 4, 0, 1, 1)
        spacerItem2 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout.addItem(spacerItem2, 5, 0, 1, 2)
        self.lineEdit_auto_hydro = QtWidgets.QLineEdit(self.groupBox_measuring_conditions)
        self.lineEdit_auto_hydro.setText("")
        self.lineEdit_auto_hydro.setObjectName("lineEdit_auto_hydro")
        self.gridLayout.addWidget(self.lineEdit_auto_hydro, 4, 1, 1, 1)
        self.label_auto_article_name = QtWidgets.QLabel(self.groupBox_measuring_conditions)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_auto_article_name.setFont(font)
        self.label_auto_article_name.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.label_auto_article_name.setObjectName("label_auto_article_name")
        self.gridLayout.addWidget(self.label_auto_article_name, 0, 0, 1, 2)
        self.horizontalLayout_2.addWidget(self.groupBox_measuring_conditions)
        self.stackedWidget_auto_measuring_tabs.addWidget(self.measuring_settings_tab)

        #   -----------------------------------------------------------------
        #   Раздел Запуск/остановка в автоматическом режиме.
        #   -----------------------------------------------------------------

        self.start_stop_tab = QtWidgets.QWidget()
        self.start_stop_tab.setObjectName("start_stop_tab")
        self.textEdit_log_auto = QtWidgets.QTextEdit(self.start_stop_tab)
        # self.textEdit_log_auto.setEnabled(False)
        self.textEdit_log_auto.setReadOnly(True)
        self.textEdit_log_auto.setGeometry(QtCore.QRect(0, 0, 971, 395))
        self.textEdit_log_auto.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.textEdit_log_auto.setObjectName("textEdit_log_auto")
        self.tableWidget_progress_bar_auto = QtWidgets.QTableWidget(self.start_stop_tab)
        self.tableWidget_progress_bar_auto.setGeometry(QtCore.QRect(0, 400, 971, 21))
        self.tableWidget_progress_bar_auto.setObjectName("tableWidget_progress_bar_auto")
        self.tableWidget_progress_bar_auto.setColumnCount(0)
        self.tableWidget_progress_bar_auto.setRowCount(1)
        self.tableWidget_progress_bar_auto.verticalHeader().hide()
        self.tableWidget_progress_bar_auto.horizontalHeader().hide()
        self.tableWidget_progress_bar_auto.horizontalHeader().setStretchLastSection(True)
        self.tableWidget_progress_bar_auto.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.tableWidget_progress_bar_auto.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget_progress_bar_auto.setVerticalHeaderItem(0, item)
        self.pushButton_start_stop_auto = QtWidgets.QPushButton(self.start_stop_tab)
        self.pushButton_start_stop_auto.setGeometry(QtCore.QRect(390, 430, 191, 51))
        font = QtGui.QFont()
        font.setPointSize(16)
        self.pushButton_start_stop_auto.setFont(font)
        self.pushButton_start_stop_auto.setStyleSheet("background-color: rgb(31, 100, 10); color: rgb(255, 255, 255);")
        self.pushButton_start_stop_auto.setObjectName("pushButton_start_stop_auto")
        self.pushButton_start_stop_auto.clicked.connect(self.start_stop)
        self.stackedWidget_auto_measuring_tabs.addWidget(self.start_stop_tab)

        #   -----------------------------------------------------------------
        #   Раздел Протокол в автоматическом режиме.
        #   -----------------------------------------------------------------        

        self.protocol_tab = QtWidgets.QWidget()
        self.protocol_tab.setObjectName("protocol_tab")
        self.label_measuring_finished = QtWidgets.QLabel(self.protocol_tab)
        self.label_measuring_finished.setGeometry(QtCore.QRect(370, 80, 321, 31))
        font = QtGui.QFont()
        font.setPointSize(18)
        self.label_measuring_finished.setFont(font)
        self.label_measuring_finished.setObjectName("label_measuring_finished")
        self.pushButton_print_auto = QtWidgets.QPushButton(self.protocol_tab)
        self.pushButton_print_auto.setGeometry(QtCore.QRect(100, 170, 201, 41))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.pushButton_print_auto.setFont(font)
        self.pushButton_print_auto.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(96, 121, 255);")
        self.pushButton_print_auto.setObjectName("pushButton_print_auto")
        self.pushButton_preview_auto = QtWidgets.QPushButton(self.protocol_tab)
        self.pushButton_preview_auto.setGeometry(QtCore.QRect(100, 230, 201, 41))
        self.pushButton_preview_auto.setFont(font)
        self.pushButton_preview_auto.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(96, 121, 255);")
        self.pushButton_preview_auto.setObjectName("pushButton_preview_auto")
        self.pushButton_preview_auto.clicked.connect(self.make_protocol)
        self.pushButton_save_data_auto = QtWidgets.QPushButton(self.protocol_tab)
        self.pushButton_save_data_auto.setGeometry(QtCore.QRect(100, 290, 201, 41))
        self.pushButton_save_data_auto.setFont(font)
        self.pushButton_save_data_auto.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(96, 121, 255);")
        self.pushButton_save_data_auto.setObjectName("pushButton_save_data_auto")
        self.pushButton_save_data_auto.clicked.connect(self.save_measurement_data)
        self.stackedWidget_auto_measuring_tabs.addWidget(self.protocol_tab)
        self.tab_auto = self.stackedWidget_main.addWidget(self.Auto)

        #   -----------------------------------------------------------------
        #                       вкладка АРХИВ
        #   -----------------------------------------------------------------

        self.Archive = QtWidgets.QWidget()
        self.Archive.setObjectName("Archive")
        self.tableWidget_archive = QtWidgets.QTableWidget(self.Archive)
        self.tableWidget_archive.setGeometry(QtCore.QRect(0, 0, 811, 501))
        self.tableWidget_archive.setObjectName("tableWidget_archive")
        self.tableWidget_archive.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows);
        self.tableWidget_archive.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection);
        columnWidths = [431, 150, 180]
        columnCount = len(columnWidths)
        self.tableWidget_archive.setColumnCount(columnCount)
        measurements = self.db_query("SELECT * FROM measurements")
        count_of_measurements = len(measurements)
        self.tableWidget_archive.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.tableWidget_archive.setRowCount(count_of_measurements)
        for column in range(columnCount):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_archive.setHorizontalHeaderItem(column, item)
            self.tableWidget_archive.setColumnWidth(column, columnWidths[column])
        columns = (3, 1, 2)
        if count_of_measurements:
            font.setPointSize(14)
            for row,measurement in enumerate(measurements):
                measurement = list(measurement)
                match measurement[1]:
                    case "manual":
                        measurement[1] = 'Ручное'
                    case "verification_TS":
                        measurement[1] = 'Поверка ТС'
                    case "graduation_TS":
                        measurement[1] = 'Градуировка ТС'
                    case "calibration_TP":
                        measurement[1] = 'Калибровка ТП'
                    case "verification_TP":
                        measurement[1] = 'Поверка ТП'
                item = QtWidgets.QTableWidgetItem()
                item.setFont(font)
                item.setText(str(measurement[0]))
                self.tableWidget_archive.setVerticalHeaderItem(row, item)
                for column in range (columnCount):
                    item = QtWidgets.QTableWidgetItem(str(measurement[columns[column]]))
                    item.setFont(font)
                    item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                    self.tableWidget_archive.setItem(row, column, item)
        self.progressBar_memory = QtWidgets.QProgressBar(self.Archive)
        self.progressBar_memory.setGeometry(QtCore.QRect(200, 540, 771, 23))
        self.progressBar_memory.setProperty("value", 10)
        self.progressBar_memory.setObjectName("progressBar_memory")
        self.label_memory = QtWidgets.QLabel(self.Archive)
        self.label_memory.setGeometry(QtCore.QRect(10, 540, 171, 21))
        self.label_memory.setFont(font)
        self.label_memory.setObjectName("label_memory")
        self.pushButton_archive_export = QtWidgets.QPushButton(self.Archive)
        self.pushButton_archive_export.setGeometry(QtCore.QRect(830, 20, 141, 31))
        self.pushButton_archive_export.setFont(font)
        self.pushButton_archive_export.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_archive_export.setObjectName("pushButton_export")
        self.pushButton_archive_delete = QtWidgets.QPushButton(self.Archive)
        self.pushButton_archive_delete.setGeometry(QtCore.QRect(830, 70, 141, 31))
        self.pushButton_archive_delete.setFont(font)
        self.pushButton_archive_delete.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_archive_delete.setObjectName("pushButton_delete")
        self.pushButton_archive_delete.clicked.connect(self.archive_delete_record)
        self.pushButton_archive_view = QtWidgets.QPushButton(self.Archive)
        self.pushButton_archive_view.setGeometry(QtCore.QRect(830, 120, 141, 31))
        self.pushButton_archive_view.setFont(font)
        self.pushButton_archive_view.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_archive_view.setObjectName("pushButton_archive_view")
        self.pushButton_archive_view.clicked.connect(self.archive_view_record)
        self.pushButton_archive_back = QtWidgets.QPushButton(self.Archive)
        self.pushButton_archive_back.setGeometry(QtCore.QRect(830, 170, 141, 31))
        self.pushButton_archive_back.setFont(font)
        self.pushButton_archive_back.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.pushButton_archive_back.setObjectName("pushButton_archive_back")
        self.pushButton_archive_back.clicked.connect(self.back_tab)
        self.tab_archive = self.stackedWidget_main.addWidget(self.Archive)

        #   -----------------------------------------------------------------
        #                       Боковое меню
        #   -----------------------------------------------------------------

        self.toolButton_side_menu = QtWidgets.QToolButton(self.centralwidget)
        self.toolButton_side_menu.setEnabled(True)
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.toolButton_side_menu.sizePolicy().hasHeightForWidth())
        self.toolButton_side_menu.setSizePolicy(sizePolicy)
        self.toolButton_side_menu.setStyleSheet("background-color: rgb(38, 0, 51); color: rgb(255, 255, 255);")
        self.toolButton_side_menu.setObjectName("toolButton_side_menu")
        self.toolButton_side_menu.clicked.connect(self.open_side_menu)
        self.widget_side_menu = QtWidgets.QWidget(self.centralwidget)
        self.widget_side_menu.setEnabled(True)
        self.widget_side_menu.setGeometry(QtCore.QRect(824, 9, 186, 582))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.widget_side_menu.sizePolicy().hasHeightForWidth())
        self.widget_side_menu.setSizePolicy(sizePolicy)
        self.widget_side_menu.setFocusPolicy(QtCore.Qt.FocusPolicy.TabFocus)
        self.widget_side_menu.setStyleSheet("background-color: rgb(38, 0, 51);")
        self.widget_side_menu.setObjectName("widget_side_menu")
        self.widget_side_menu.close()
        font = QtGui.QFont()
        font.setPointSize(14)
        self.pushButton_mode_mnual = QtWidgets.QPushButton(self.widget_side_menu)
        self.pushButton_mode_mnual.setGeometry(QtCore.QRect(6, 10, 171, 41))
        self.pushButton_mode_mnual.setFont(font)
        self.pushButton_mode_mnual.setStyleSheet("background-color: rgb(74, 0, 99); color: rgb(255, 255, 255);")
        self.pushButton_mode_mnual.setObjectName("pushButton_mode_mnual")
        self.pushButton_mode_mnual.clicked.connect(self.open_mode_manual)
        self.pushButton_mode_auto = QtWidgets.QPushButton(self.widget_side_menu)
        self.pushButton_mode_auto.setGeometry(QtCore.QRect(6, 60, 171, 41))
        self.pushButton_mode_auto.setFont(font)
        self.pushButton_mode_auto.setStyleSheet("background-color: rgb(74, 0, 99); color: rgb(255, 255, 255);")
        self.pushButton_mode_auto.setObjectName("pushButton_mode_auto")
        self.pushButton_mode_auto.clicked.connect(self.open_mode_auto)
        self.pushButton_sensors = QtWidgets.QPushButton(self.widget_side_menu)
        self.pushButton_sensors.setGeometry(QtCore.QRect(6, 110, 171, 41))
        self.pushButton_sensors.setFont(font)
        self.pushButton_sensors.setStyleSheet("background-color: rgb(74, 0, 99); color: rgb(255, 255, 255);")
        self.pushButton_sensors.setObjectName("pushButton_sensors")
        self.pushButton_sensors.clicked.connect(self.open_sensors)
        self.pushButton_settings_ISH = QtWidgets.QPushButton(self.widget_side_menu)
        self.pushButton_settings_ISH.setGeometry(QtCore.QRect(6, 160, 171, 41))
        self.pushButton_settings_ISH.setFont(font)
        self.pushButton_settings_ISH.setStyleSheet("background-color: rgb(74, 0, 99); color: rgb(255, 255, 255);")
        self.pushButton_settings_ISH.setObjectName("pushButton_sensors")
        self.pushButton_settings_ISH.clicked.connect(self.open_settings_ISH)
        self.pushButton_archive = QtWidgets.QPushButton(self.widget_side_menu)
        self.pushButton_archive.setGeometry(QtCore.QRect(6, 210, 171, 41))
        self.pushButton_archive.setFont(font)
        self.pushButton_archive.setStyleSheet("background-color: rgb(74, 0, 99); color: rgb(255, 255, 255);")
        self.pushButton_archive.setObjectName("pushButton_archive")
        self.pushButton_archive.clicked.connect(self.open_archive)
        self.pushButton_settings = QtWidgets.QPushButton(self.widget_side_menu)
        self.pushButton_settings.setGeometry(QtCore.QRect(6, 530, 171, 41))
        self.pushButton_settings.setFont(font)
        self.pushButton_settings.setStyleSheet("background-color: rgb(74, 0, 99); color: rgb(255, 255, 255);")
        self.pushButton_settings.setObjectName("pushButton_settings")
        self.pushButton_settings.clicked.connect(self.open_settings)
        MainWindow.setCentralWidget(self.centralwidget)
        self.retranslateUi(MainWindow)
        self.stackedWidget_main.setCurrentIndex(self.tab_logo)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        #   -----------------------------------------------------------------
        #                       инициализация
        #   -----------------------------------------------------------------

        #   ===     загрузка списка измерений в раздел Архив    ===
        default_t = self.db_query('SELECT value FROM options WHERE option="default_t"')[0][0]
        self.lineEdit_manual_temp.setText(default_t)
        self.lineEdit_auto_temp.setText(default_t)
        default_p = self.db_query('SELECT value FROM options WHERE option="default_p"')[0][0]
        self.lineEdit_manual_pressure.setText(default_p)
        self.lineEdit_auto_pressure.setText(default_p)
        default_h = self.db_query('SELECT value FROM options WHERE option="default_h"')[0][0]
        self.lineEdit_manual_hydro.setText(default_h)
        self.lineEdit_auto_hydro.setText(default_h)
        match platform:                 # определяем на какой платформе запускается приложение
            case "win32":               # под Win запускаемся в оконном режиме и соответвующими COM портами
                MainWindow.showNormal()
                self.template_path = '\\templates\\'
            case "linux":               # под Linux запускаемся в полном экране
                MainWindow.showFullScreen()
                self.template_path = '//templates//'
        self.portname_bu7 = self.db_query(f"SELECT value FROM options WHERE option='com_port_bu7_{platform}'")[0][0]
        self.portname_itm = self.db_query(f"SELECT value FROM options WHERE option='com_port_itm_{platform}'")[0][0]
        self.lineEdit_comport_bu7.setText(self.portname_bu7)
        self.lineEdit_comport_itm.setText(self.portname_itm)
        self.timer_measuring = QTimer()
        self.timer_measuring.setInterval(5000)                  # Измерение раз в 5 сек.
        self.timer_measuring.timeout.connect(self.do_measuring_step)
        self.timer_measuring.start()
        self.timer_ustavka = QTimer()
        self.timer_ustavka.timeout.connect(self.do_ustavka_step)
        self.timer_auto_switch_logo = QTimer()
        self.timer_auto_switch_logo.setInterval(1000)
        self.timer_auto_switch_logo.setSingleShot(True)
        self.timer_auto_switch_logo.timeout.connect(self.open_mode_manual)
        self.timer_auto_switch_logo.start()
        self.graph_x_points = 10
        self.graphWidget.clear()                            # очистка графиков
        self.data_for_graph = []                            # список всех измеренных данных для графиков
        self.sensor_pen = []                                # список свойств карандашей для отображения графиков по каналам
        self.sensor_graph = []                              # указатели на графики по каналам
        self.etalon_channel = -1                            # канал подключенного эталонного датчика
        self.thread_itm_exchange = threading.Thread(target=self.itm_exchange_worker, daemon=True)
        self.itm_meas_request = self.itm_request_channels = self.itm_command = self.itm_result = ''
        self.itm_exchange_stop = False
        self.itm_run = False
        self.thread_itm_exchange.start()
        self.tab_path = []                                  # предыдущие вкладки для кнопки "назад"
        for i in range(9):
            self.data_for_graph.append(list())
            self.data_for_graph[i].append(list())   # список измеренных данных на канал для графика
            self.data_for_graph[i].append(list())   # время, в которое сделан замер для графика
            exec("self.sensor_pen.append(pg.mkPen(color=("+self.color_buttons[i]+"), width=3))")    # создаем карандаш графика из цвета кнопки
            self.sensor_graph.append(self.graphWidget.plot(list(range(0)),list(range(0)), pen=self.sensor_pen[i]) )
            if not eval("self.pushButton_"+str(i+1)+".isChecked()"):
                self.graphWidget.removeItem(self.sensor_graph[i])

    def itm_exchange_worker(self):
        while True:
            self.itm_exchange_stop = False
            if self.itm_command:
                if self.itm_command == 'W,stopmeas=1': self.itm_run = False
                elif self.itm_command == 'W,startmeas=1': self.itm_run = True
                print(self.itm_command)
                self.itm_result = exchange_data_itm(self, self.itm_command, self.portname_itm)
                print(self.itm_result)
                self.itm_command = None
            elif self.itm_meas_request and self.itm_request_channels and eval('self.pushButton_start_stop_'+ self.measuring_mode +'.text()') == 'Стоп' and self.itm_run:
                print(f'worker request: {self.itm_meas_request}')
                self.itm_meas_result = exchange_data_itm(self, self.itm_meas_request, self.portname_itm)
                print(f'worker result: {self.itm_meas_result}')
            else: self.itm_meas_result = False

    def itm_send_command(self, command, check, ok_message, err_message):
        self.itm_exchange_stop = True
        self.itm_command = command
        self.itm_result = None
        while not self.itm_result:
            time.sleep(0.1)
        if self.itm_result == check:
            exec(f'self.textEdit_log_{self.measuring_mode}.append(f"{ok_message}. {self.itm_result}")')
            return True
        else:
            exec(f'self.textEdit_log_{self.measuring_mode}.append(f"{err_message}. {self.itm_result}")')
            return False

    def start_stop(self):
        button = QApplication.instance().sender()
        button_name = "self." + button.objectName()
        match eval(f'{button_name}.text()'):
            case "Старт":
                self.start_measuring()
            case "Стоп":
                self.stop_measuring()
            case _:
                exec(f'self.textEdit_log_{self.measuring_mode}.append("{button_name.text()}")')

    def start_measuring(self):
        rowCount = eval(f'self.tableWidget_ustavka_{self.measuring_mode}.rowCount()-1')     #т.к. последняя строка с "+"
        if self.measuring_mode == 'auto':
            if rowCount <= 0:
                self.textEdit_log_auto.append("Не установлено ни одной температурной точки.")
                return
            elif self.etalon_channel < 0:
                self.textEdit_log_auto.append("Не подключен эталонный датчик.")
                return
        self.stab_time = None
        self.itm_meas_result = ''
        self.auto_measuring_step = 0
        self.measured_data_channel = []                     # список измеренных достоверных (в режиме стабильности) данных по каналам
        self.measured_data_counter_channel = []             # список счетчиков измеренных достоверных данных текущей уствки на канал (временные данные)
        self.time_of_start = QtCore.QDateTime.currentDateTime() # время начала измерений
        for i in range(8):
            self.measured_data_channel.append(list())       # список содержит списки значений и времени
            self.measured_data_channel[i].append(list())    # 0 - список измеренных данных на канал
            self.measured_data_channel[i].append(list())    # 1 - время, в которое сделан замер
            self.measured_data_channel[i].append(list())    # 2 - номер уставки, если уставок нет, то 0
            self.measured_data_channel[i].append(list())    # 3 - сопротивыление, если ТС или ЭДС, если ТП
            self.measured_data_counter_channel.append(0)
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setEnabled(False)')
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(163, 194, 194);")')
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setText("Сохранить")')
        self.toolButton_side_menu.setEnabled(False)
        if self.measuring_mode == 'auto': self.auto_measuring_step_buttons(False)
        self.counter_of_measurements = int(self.db_query("SELECT value FROM options WHERE option='counter_of_measurements'")[0][0])
        while self.db_query(f"SELECT COUNT(*) FROM measurements WHERE id={self.counter_of_measurements}")[0][0]>0: self.counter_of_measurements+=1
        exec(f'self.textEdit_log_{self.measuring_mode}.append("Измерение № {str(self.counter_of_measurements)}")')
        exec(f'self.label_{self.measuring_mode}_measurement_number.setText(str(self.counter_of_measurements))')
        if rowCount > 0:                                            # если в таблице уставки есть хоть одна уставка, то обрабатываем её.
            self.ustavka_step = 0
            total_time = 0
            exec(f'self.textEdit_log_{self.measuring_mode}.append("Количество уставок: {str(rowCount)}")')
            exec(f'self.tableWidget_progress_bar_{self.measuring_mode}.setColumnCount(rowCount)')
            match self.measuring_mode:
                case 'auto':
                    table_width = 971
                    # time = QtCore.QTime.fromString("00:00:56", "hh:mm:ss").msecsSinceStartOfDay()
                    time = 8
                    total_time = rowCount * time
                case 'manual':
                    table_width = 921
                    for row in range(rowCount): total_time += QtCore.QTime.fromString(self.tableWidget_ustavka_manual.item(row,1).text(), "hh:mm").msecsSinceStartOfDay()
            for col in range(rowCount):
                if self.measuring_mode == 'manual': time = QtCore.QTime.fromString(self.tableWidget_ustavka_manual.item(col,1).text(), "hh:mm").msecsSinceStartOfDay()
                cell_width = int(table_width*time/total_time)
                exec(f'self.tableWidget_progress_bar_{self.measuring_mode}.setColumnWidth(col, cell_width)')
                progressBar_name = "self.progressBar_"+str(col)
                exec(progressBar_name+" = QtWidgets.QProgressBar()")
                exec(progressBar_name+".setObjectName('"+progressBar_name+"')")
                exec(progressBar_name+".setAlignment(Qt.AlignmentFlag.AlignHCenter)")
                exec(progressBar_name+".setStyleSheet('QProgressBar{max-height: 19px;}')") #padding: 1px;
                exec(progressBar_name+".setRange(0,"+str(time)+")")
                exec(f'self.tableWidget_progress_bar_{self.measuring_mode}.setCellWidget(0, col, {progressBar_name})')
            self.do_ustavka() # запустили нагрев на данную уставку
        else:                                                       # Иначе просто создаем один прогрессбар с максимальным значением и показываем на нем "Замер".
            self.textEdit_log_manual.append("Уставки не найдены, выполняем измерение каналов.")
            self.progressBar_0 = QtWidgets.QProgressBar()
            self.progressBar_0.setObjectName("self.progressBar_0")
            self.progressBar_0.setAlignment(Qt.AlignmentFlag.AlignHCenter)
            self.progressBar_0.setStyleSheet('QProgressBar{max-height: 19px;}')
            self.progressBar_0.setFormat('Замер')
            self.progressBar_0.setRange(0,100)
            self.progressBar_0.setValue(100)
            self.tableWidget_progress_bar_manual.setColumnCount(1)
            self.tableWidget_progress_bar_manual.setCellWidget(0, 0, self.progressBar_0)
        self.itm_meas_request = "R"
        self.itm_request_channels = []
        for i in range(8):    # перебираем каналы в таблице датчиков, если указано, что датчик подключен, то добавляем его в запрос для измерения.
            cell_text = eval(f'{self.table_name}.item(i, 0).text()')
            if cell_text != "+" and cell_text != "Перемычка":
                self.itm_meas_request += ',valch'+str(i+1) + ',dopch'+str(i+1)
                self.itm_request_channels.append(i+1)
        self.itm_send_command("W,startmeas=1", "W,startmeas-OK", "Запуск измерений в ИТМ", "Невозможно запустить измерения в ИТМ")
        exec(f'self.pushButton_start_stop_{self.measuring_mode}.setStyleSheet("background-color: rgb(255, 0, 0); color: rgb(255, 255, 255);")')
        exec(f'self.pushButton_start_stop_{self.measuring_mode}.setText("Стоп")')

    def stop_measuring(self):
        exec('self.tableWidget_progress_bar_' + self.measuring_mode + '.setColumnCount(0)')
        self.timer_ustavka.stop()
        self.itm_meas_request = self.itm_request_channels = ''
        exec('self.pushButton_start_stop_'+ self.measuring_mode +'.setText("Старт")')
        exec('self.pushButton_start_stop_'+ self.measuring_mode +'.setStyleSheet("background-color: rgb(31, 100, 10); color: rgb(255, 255, 255);")')
        self.itm_send_command("W,stopmeas=1", "W,stopmeas-OK", "Останов измерений в ИТМ", "Невозможно остановить измерения в ИТМ")
        self.toolButton_side_menu.setEnabled(True)
        if self.measuring_mode == 'auto': self.auto_measuring_step_buttons()
        result = exchange_data_bu7(self, "W,powerset=OFF", self.portname_bu7)
        if result == "W,powerset-OK":
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Остановили нагрев")')
        else:
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Нет связи с БУ7: '+str(result)+'")')
        if 'measured_data_channel' in self.__dict__:
            empty = True
            for chan in range (8):
                if len (self.measured_data_channel[chan][0]): empty = False
            if not empty: 
                exec(f'self.pushButton_save_data_{self.measuring_mode}.setEnabled(True)')
                exec(f'self.pushButton_save_data_{self.measuring_mode}.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(96, 121, 255);")')
                if self.measuring_mode == 'auto':
                    self.stackedWidget_auto_measuring_tabs.setCurrentIndex(4)
                    self.calculating_protocol_auto()
                    return
                else:
                    return

    def do_ustavka_step(self):
        if self.measuring_mode == 'manual': # доводим текущий progressbar до 100%
            exec(f'self.progressBar_{str(self.ustavka_step)}.setValue({str(self.timer_ustavka.interval())})')
        else:
            exec(f'self.progressBar_{str(self.ustavka_step)}.setValue(8)')
        self.ustavka_step +=1
        if self.ustavka_step >= eval('self.tableWidget_ustavka_'+ str(self.measuring_mode) + '.rowCount()-1'):
            self.stop_measuring()
            return
        self.do_ustavka()

    def do_ustavka(self):
        temperature = int(eval(f'self.tableWidget_ustavka_{self.measuring_mode}.item(self.ustavka_step,0).text()'))
        exec('self.textEdit_log_'+ self.measuring_mode +'.append("W,Ts2='+str(temperature)+'.0")')
        result = exchange_data_bu7(self, "W,Ts2="+str(temperature)+".0", self.portname_bu7)               #Запускаем нагрев на заданную температуру
        if result == "W,Ts2-OK":
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Установлена температуру нагрева на:'+str(temperature)+' градусов")')
        else:
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Нет связи с БУ7: '+str(result)+'")')
        result = exchange_data_bu7(self, "W,powerset=ON", self.portname_bu7)
        if result == "W,powerset-OK":
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Запустили нагрев на:'+str(temperature)+' градусов")')
        else:
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Нет связи с БУ7: '+str(result)+'")')
        for row in range(8): self.measured_data_counter_channel[row] = 0
        self.stab_prev_temp = 0
        if self.measuring_mode == 'manual':
            time = QtCore.QTime.fromString(self.tableWidget_ustavka_manual.item(self.ustavka_step,1).text(), "hh:mm").msecsSinceStartOfDay()
            self.timer_ustavka.setInterval(time)
            self.timer_ustavka.setSingleShot(True)
        else:
            self.auto_measuring_step = 0
        self.progressbar_show_heating()

    def do_measuring_step(self):
        self.do_measuring_step_body()   #приходится разбивать функцию на 2, т.к. мы не знаем на каком этапе завершится функция, а нам надо обновить графики
        if self.measuring_mode == 'manual': self.show_graph()

    def do_measuring_step_body(self):
        result = exchange_data_bu7(self, 'R,T2', self.portname_bu7)
        # print(result)
        current_temp = extract_val_param(result, 'T2')
        if current_temp == 'errval=':
            # exec('self.textEdit_log_'+ self.measuring_mode +'.append("Нет связи с БУ7")')
            current_temp = random()+27
        else:
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Текущая температура: '+str(current_temp)+'")')
        if self.measuring_mode == 'manual':                                                 # Если режим - ручное измерение, то рисуем график температуры БУ7
            interval = QtCore.QDateTime.currentDateTime().toTime_t()                        # - self.time_of_start.toTime_t(); кажется это текущее время.
            self.label_temp_of_chan_9.setText(str(f"{float(current_temp):.2f}"))
            self.data_for_graph[8][0].append(float(current_temp))                   # заполняем список данных из нагревателя БУ7, температуру
            self.data_for_graph[8][1].append(interval)                              # и время
        if eval('self.pushButton_start_stop_'+ self.measuring_mode +'.text()') == 'Старт': return
        if eval('self.tableWidget_ustavka_' + self.measuring_mode + '.rowCount()')-1 > 0:   # если задана хоть одна уставка, то проверяем стабильность
            if self.stab_prev_temp: exec('self.textEdit_log_'+ self.measuring_mode +'.append("Дрейф:'+str(abs(float(current_temp) - float(self.stab_prev_temp)))+'")')
            if abs(float(current_temp) - float(self.stab_prev_temp)) >= 1:  # 0.1 должно быть в релизе
                self.stab_time = None
                self.auto_measuring_step = 0
                self.timer_ustavka.stop()
                self.stab_prev_temp = current_temp
                for row in range(8):
                    if self.measured_data_counter_channel[row]:
                        exec('self.textEdit_log_'+ self.measuring_mode +'.append("Очистка временных данных:'+str(self.measured_data_counter_channel[row])+', '+str(self.measured_data_channel[row][0])+', '+str(len(self.measured_data_channel[row][0]))+'")')
                        self.measured_data_channel[row][0]=self.measured_data_channel[row][0][0:len(self.measured_data_channel[row][0])-self.measured_data_counter_channel[row]]
                        self.measured_data_channel[row][1]=self.measured_data_channel[row][1][0:len(self.measured_data_channel[row][1])-self.measured_data_counter_channel[row]]
                        self.measured_data_channel[row][2]=self.measured_data_channel[row][2][0:len(self.measured_data_channel[row][2])-self.measured_data_counter_channel[row]]
                        self.measured_data_channel[row][3]=self.measured_data_channel[row][3][0:len(self.measured_data_channel[row][3])-self.measured_data_counter_channel[row]]
                        exec('self.textEdit_log_'+ self.measuring_mode +'.append("'+str(self.measured_data_channel[row][0])+'")')
                        self.measured_data_counter_channel[row] = 0
                self.progressbar_show_heating()
                exec('self.textEdit_log_'+ self.measuring_mode +'.append("Ждем стабильности.")')
                return
            if self.stab_time == None: 
                self.stab_time = QtCore.QTime.currentTime()
                return
            if abs(QtCore.QTime.currentTime().secsTo(self.stab_time)) < 5: return   # 1800 должно быть в релизе, т.е. 30 мин.
            exec('self.textEdit_log_'+ self.measuring_mode +'.append("Зона стабильности.")')
            if self.measuring_mode == 'manual':
                interval = self.timer_ustavka.interval()
                remainig_time = self.timer_ustavka.remainingTime()
                if  not self.timer_ustavka.isActive():
                    self.timer_ustavka.start()
                    exec('self.textEdit_log_'+ self.measuring_mode +'.append("Запускаем таймер текущей уставки.")')
            else:
                interval = 8
                remainig_time = 8 - self.auto_measuring_step
            exec("self.progressBar_"+str(self.ustavka_step)+".setStyleSheet('QProgressBar{max-height: 19px;}')") #padding: 1px;
            exec("self.progressBar_"+str(self.ustavka_step)+".setValue("+str(interval-remainig_time)+")")
            exec("self.progressBar_"+str(self.ustavka_step)+".resetFormat()")
        itm_meas_result = self.itm_meas_result
        if itm_meas_result and itm_meas_result.find('Err')<=-1:
            self.itm_meas_result = None
            print(itm_meas_result)
            for chan in self.itm_request_channels:
                t = extract_val_param(itm_meas_result, "valch"+str(chan))
                if t == "errval=":
                    exec('self.textEdit_log_'+ self.measuring_mode +'.append("Нет связи с ИТМ:'+str(itm_meas_result)+'")')
                    t = -271    #random()*4+27
                print('chan'+ str(chan) +' t=',t)
                if self.measuring_mode == 'manual':                             # если ручной режим, то рисуем графики
                    val = f"{float(t):.2f}"                                     # оставляем 2 разряда после точки
                    exec(f"self.label_temp_of_chan_{str(chan)}.setText(str(val))")
                    # exec("self.progressBar_"+str(self.ustavka_step)+".setStyleSheet('')")
                    self.data_for_graph[chan-1][0].append(float(val))
                    self.data_for_graph[chan-1][1].append(QtCore.QDateTime.currentDateTime().toTime_t())
                    self.tableWidget_manual_measuring_result.item(4,chan-1).setText(str(t))
                r_or_e = extract_val_param(itm_meas_result, "dopch"+str(chan))
                if r_or_e == "errval=":
                    exec('self.textEdit_log_'+ self.measuring_mode +'.append("Нет связи с ИТМ:'+str(itm_meas_result)+'")')
                    r_or_e = 0      #random()*4+72
                print('chan'+ str(chan) +' r or e =', r_or_e)
                self.measured_data_channel[chan-1][0].append(float(t))
                self.measured_data_channel[chan-1][1].append(QtCore.QDateTime.currentDateTime())
                if eval('self.tableWidget_ustavka_' + self.measuring_mode + '.rowCount()')-1 > 0:
                    self.measured_data_channel[chan-1][2].append(self.ustavka_step+1)
                else:
                    self.measured_data_channel[chan-1][2].append(0)
                self.measured_data_channel[chan-1][3].append(float(r_or_e))
                self.measured_data_counter_channel[chan-1]+=1
                # current_time = QtCore.QDateTime.currentDateTime().toString("d MMM yyyy HH:mm:ss")
                # exec(f'self.textEdit_log_{self.measuring_mode}.append("Измерение канала: {str(chan+1)}: {str(t)}, {current_time}")')
                # exec(f'self.textEdit_log_{self.measuring_mode}.append("Измерение канала: {self.measured_data_channel[chan][0]}, {self.measured_data_channel[chan][2]}")')
            if self.measuring_mode == 'auto':
                self.itm_send_command("W,stopmeas=1", "W,stopmeas-OK", "Останов измерений в ИТМ", "Невозможно остановить измерения в ИТМ")
                if self.auto_measuring_step % 2:
                    self.itm_send_command("W,rewerseChan=FF", "W,rewerseChan-FF", "Измерение каналов в прямой последовательности", "Невозможно включить измерение каналов в прямой последовательности")
                else:
                    self.itm_send_command("W,rewerseChan=RW", "W,rewerseChan-RW", "Измерение каналов в обратной последовательности", "Невозможно включить измерение каналов в обратной последовательности")
                self.itm_send_command("W,startmeas=1", "W,startmeas-OK", "Запуск измерений в ИТМ", "Невозможно запустить измерения в ИТМ")
                print('Шаг измерения---------------', self.auto_measuring_step)
                self.auto_measuring_step += 1
                if self.auto_measuring_step == 8:
                    self.do_ustavka_step()

    def show_graph(self):
        for chan in range(9):
            if len(self.data_for_graph[chan][0]):
                self.sensor_graph[chan].setData(self.data_for_graph[chan][1], self.data_for_graph[chan][0])
                # exec('self.textEdit_log_'+ self.measuring_mode +'.append(f"{self.data_for_graph[chan][1]}, {self.data_for_graph[chan][0]}")')

    def graph_btn_toggle(self):
        button = QApplication.instance().sender()
        button_name = "self."+button.objectName()
        button_num = int(button_name.split("_")[1])
        if eval("self.pushButton_"+str(button_num)+".isChecked()"):
            self.graphWidget.addItem(self.sensor_graph[button_num-1])
        else:
            self.graphWidget.removeItem(self.sensor_graph[button_num-1])
    
    def move_graph_left(self):
        graph_range = self.graphWidget.viewRange()
        delta = (graph_range[0][1]-graph_range[0][0])/10
        self.graphWidget.setXRange(graph_range[0][0]-delta,graph_range[0][1]-delta, padding = 0)

    def move_graph_right(self):
        graph_range = self.graphWidget.viewRange()
        delta = (graph_range[0][1]-graph_range[0][0])/10
        self.graphWidget.setXRange(graph_range[0][0]+delta,graph_range[0][1]+delta, padding = 0)

    def move_graph_up(self):
        graph_range = self.graphWidget.viewRange()
        delta = (graph_range[1][1]-graph_range[1][0])/10
        self.graphWidget.setYRange(graph_range[1][0]-delta,graph_range[1][1]-delta, padding = 0)

    def move_graph_down(self):
        graph_range = self.graphWidget.viewRange()
        delta = (graph_range[1][1]-graph_range[1][0])/10
        self.graphWidget.setYRange(graph_range[1][0]+delta,graph_range[1][1]+delta, padding = 0)

    def scale_vertical_up(self):
        graph_range = self.graphWidget.viewRange()
        delta = (graph_range[1][1]-graph_range[1][0])/20
        self.graphWidget.setYRange(graph_range[1][0]-delta,graph_range[1][1]+delta, padding = 0)

    def scale_vertical_down(self):
        graph_range = self.graphWidget.viewRange()
        delta = (graph_range[1][1]-graph_range[1][0])/20
        self.graphWidget.setYRange(graph_range[1][0]+delta,graph_range[1][1]-delta, padding = 0)

    def scale_horizontal_up(self):
        graph_range = self.graphWidget.viewRange()
        x_range = (graph_range[0][1]-graph_range[0][0])
        delta = x_range/20
        if x_range >= 10:self.graph_x_points = int(x_range)
        print (self.graph_x_points, x_range)
        self.graphWidget.setXRange(graph_range[0][0]-delta,graph_range[0][1]+delta, padding = 0)

    def scale_horizontal_down(self):
        graph_range = self.graphWidget.viewRange()
        x_range = (graph_range[0][1]-graph_range[0][0])
        delta = x_range/20
        if x_range >= 10:self.graph_x_points = int(x_range)
        print (self.graph_x_points, x_range)
        self.graphWidget.setXRange(graph_range[0][0]+delta,graph_range[0][1]-delta, padding = 0)

    def scale_auto(self):
        self.graphWidget.enableAutoRange()

    def add_ustavka_to_tab(self, selected_row):
        parent = QApplication.instance().sender()
        rowCount = self.tableWidget_ustavka_manual.rowCount()
        # if selected_row !=0:
        #     if selected_row+1 == rowCount and self.tableWidget_ustavka_manual.item(selected_row-1,1).text() == "00:00": return
        temperature = self.tableWidget_ustavka_manual.item(selected_row,0).text()
        time = self.tableWidget_ustavka_manual.item(selected_row,1).text()
        Change_ustavka_dialog_inst = Change_ustavka_dialog(temperature, time, selected_row, rowCount, self)
        Change_ustavka_dialog_inst.show()
        Change_ustavka_dialog_inst.exec()

    def add_temp_to_tab(self, selected_row):
        if self.auto_measuring_mode != 'verification_TS':
            parent = QApplication.instance().sender()
            rowCount = self.tableWidget_ustavka_auto.rowCount()
            temperature = self.tableWidget_ustavka_auto.item(selected_row,0).text()
            Change_temp_dialog_inst = Change_temp_dialog(temperature, selected_row, rowCount, self)
            Change_temp_dialog_inst.show()
            Change_temp_dialog_inst.exec()

    # Добавление и удаление датчиков в общий список всех датчиков.

    def add_sensor_to_list(self, selected_row):
        Sensor_edit_dialog_inst = Sensor_edit_dialog(self, selected_row)
        Sensor_edit_dialog_inst.show()
        Sensor_edit_dialog_inst.exec()

    # Редактирование доверительной погрешности.

    def confidence_error(self):
        Confidence_error_dialog_inst = Confidence_error_dialog(self)
        Confidence_error_dialog_inst.show()
        Confidence_error_dialog_inst.exec()

    def progressbar_show_heating(self):
        exec("self.progressBar_"+str(self.ustavka_step)+".setStyleSheet('QProgressBar{color:rgb(0,0,0);} QProgressBar::chunk {background-color: rgb(255,0,0);}')")
        exec("self.progressBar_"+str(self.ustavka_step)+".setFormat('Нагрев')")
        if self.measuring_mode == 'manual':
            exec(f'self.progressBar_{str(self.ustavka_step)}.setValue({str(self.timer_ustavka.interval())})')
        else:
            exec(f'self.progressBar_{str(self.ustavka_step)}.setValue(8)')


    # Добавление и удаление датчиков в таблицу текущих измерений
    
    def add_sensor_to_measuring_tab(self, selected_row, selected_col):
        table = QApplication.instance().sender()
        table_name = "self."+table.objectName()
        if table_name == "self.tableWidget_auto_channels_settings_TP" and selected_row == 0: return
        if table_name == "self.tableWidget_manual_channels_settings" and eval(table_name + '.item(0,0).text()') == "Перемычка" and selected_row == 0: return
        mode = table_name.split("_")[1]
        item_val = eval(table_name + '.item(selected_row,0).text()')
        items = []
        if item_val != "+":
            data = eval(table_name+'.item(selected_row, selected_col).text()')
            type_of_test = ('Периодич.', 'Первичная', 'Калибровка')
            comp_HK = ('OFF', 'ON')
            input_dialog_items = type_of_test
            match table.objectName():
                case "tableWidget_manual_channels_settings":
                    columns_add_data = [6]
                    input_dialog_items = comp_HK
                    dialog_name = {6:'Компенсация ХК'}
                    input_dialog_mode = {6:'getItem'}
                case "tableWidget_auto_channels_settings_TP":
                    columns_add_data = [1,6,7,8,9,10]
                    if selected_col == 7: input_dialog_items = comp_HK
                    dialog_name = {1:'Регистрационный №', 6:"Вид испытаний", 7:"Компенсация ХК", 8:'Замечания по внеш. осмотру', 9:'Электрич. прочность изоляции', 10:'R изоляции'}
                    input_dialog_mode = {1:'getText', 6:'getItem', 7:'getItem', 8:'getText', 9:'getText', 10:'getText'}
                    default_text = {1:'', 8:'Без замечаний', 9:'В норме', 10:'В норме'}
                case "tableWidget_auto_channels_settings_TS_verifications":
                    columns_add_data = [1,6,7,8,9]
                    dialog_name = {1:'Регистрационный №', 6:"Вид испытаний", 7:'R выводов', 8:'Замечания по внеш. осмотру', 9:'R изоляции'}
                    input_dialog_mode = {1:'getText', 6:'getItem', 7:'getDouble', 8:'getText', 9:'getText'}
                    default_text = {1:'', 7:0, 8:'Без замечаний', 9:'В норме'}
                case "tableWidget_auto_channels_settings_TS_gradiations":
                    columns_add_data = [1,7,8,9]
                    dialog_name = {1:'Регистрационный №', 7:'R выводов', 8:'Замечания по внеш. осмотру', 9:'R изоляции'}
                    input_dialog_mode = {1:'getText', 7:'getDouble', 8:'getText', 9:'getText', 10:'getText'}
                    default_text = {1:'', 7:0, 8:'Без замечаний', 9:'В норме'}
            if selected_col in columns_add_data:
                match input_dialog_mode[selected_col]:
                    case 'getItem':
                        if data: index = input_dialog_items.index(data)
                        else: index = 0
                        data, ok = QInputDialog.getItem(table, dialog_name[selected_col], 'Выберите значение:', input_dialog_items, index, False)
                    case 'getText':
                        if not data: data = default_text[selected_col]
                        data, ok = QInputDialog.getText(table, dialog_name[selected_col], 'Введите новое значение:', QLineEdit.EchoMode.Normal, data)
                    case 'getDouble':
                        if not data: data = default_text[selected_col]
                        data, ok = QInputDialog.getDouble(table, dialog_name[selected_col], 'Введите новое значение:', float(data), 0, 1000, 1)
                if ok:
                    exec(table_name+'.item(selected_row, selected_col).setText(str(data))')
                    if input_dialog_items == comp_HK:                # обновляем вкл\выкл компенсации холодных концов
                        self.itm_send_command(f'W,hk{str(selected_row+1)}={str(data)}', f'W,hk{str(selected_row+1)}-{str(data)}', f"Компенсация ХК для канала {selected_row+1} включена", f"Невозможно включить компенсацию ХК для канала {selected_row+1}")
                return
            items.append("Удалить")
        sensors_qty = self.db_query("SELECT COUNT(*) FROM sensors")[0][0]
        sensor = self.db_query("SELECT * FROM sensors")
        ish_set = self.db_query('SELECT value FROM options WHERE option="sh_set_ish"')
        sh_set_nsh_tp = self.db_query('SELECT value FROM options WHERE option="sh_set_nsh_tp"')
        sh_set_nsh_tr = self.db_query('SELECT value FROM options WHERE option="sh_set_nsh_tr"')
        for sensor_number in range(sensors_qty):    # подготавливаем список датчиков соответствующих типу измерения.
            match table.objectName()[0:37]:
                case "tableWidget_auto_channels_settings_TP":
                    if sh_set_nsh_tr[0][0].find(sensor[sensor_number][4])>=0: continue
                case "tableWidget_auto_channels_settings_TS":
                    if sh_set_nsh_tp[0][0].find(sensor[sensor_number][4])>=0: continue
                case "tableWidget_manual_channels_settings":
                    if sh_set_nsh_tp[0][0].find(sensor[sensor_number][4])>=0:
                        if selected_row == 0: continue
                        if eval(table_name + '.item(0,0).text()') != "+" and eval(table_name + '.item(0,0).text()') != "Перемычка": continue
            is_sensor_already_present = False
            for i in range (8):                                         #исключаем уже существующий в таблице датчик
                if eval(table_name+'.item(i, (self.measuring_mode=="auto")+1).text()')==sensor[sensor_number][2]: is_sensor_already_present = True
            if is_sensor_already_present: continue
            items.append(str(sensor[sensor_number][1])+", №"+str(sensor[sensor_number][2]))
        item, ok = QInputDialog.getItem(table, 'Выбор датчика', 'Выберите датчик:', items, 0, False)

        if ok:
            sensor_previous_sh = eval(table_name + '.item(selected_row, 3).text()')
            for i in range(eval(table_name+'.columnCount()')): exec(table_name + '.item(selected_row, i).setText("")')           # сначала очищаем всю строку датчика
            if mode == 'manual':                                            # если измерение ручное, то очищаем таблицу резултатов ручного измерения
                for row in range(6): self.tableWidget_manual_measuring_result.item(row+1, selected_row).setText('')
            if item == "Удалить":
                exec(table_name + '.item(selected_row, 0).setText("+")')
                if self.etalon_channel == selected_row: self.etalon_channel = -1
                if mode == 'manual':
                    exec(f"self.label_temp_of_chan_{selected_row+1}.setText('-.-')")
                    exec("self.pushButton_"+str(selected_row+1)+".setEnabled(False)")
                    exec("self.pushButton_"+str(selected_row+1)+".setChecked(False)")
                    self.graphWidget.removeItem(self.sensor_graph[selected_row])
                    if sh_set_nsh_tp[0][0].find(sensor_previous_sh)>=0: #если удаляемый датчик является ТП, то проверяем последний-ли он и удаляем Перемычку
                        is_it_last_TP = True
                        print(sh_set_nsh_tp[0][0])
                        for i in range(int(eval(table_name + '.rowCount()'))):
                            sh = eval(table_name+'.item(i, 3).text()')
                            if sh:
                                if sh_set_nsh_tp[0][0].find(sh) >= 0 and selected_row != i: is_it_last_TP = False
                        if is_it_last_TP:
                            exec(table_name + '.item(0,0).setText("+")')
                            exec(table_name + '.item(0,0).setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)')
                self.itm_send_command(f"W,sen{str(selected_row+1)}=non", f"W,sen{str(selected_row+1)}-OK", f'Для канала {selected_row+1} установлен тип датчика non', f'Невозможно установить тип датчика non  для канала {selected_row+1}')
                return
            sensor_sn = item.split(", №")[1]
            sensor = self.db_query(f'SELECT * FROM sensors WHERE sn="{sensor_sn}"')
            match table.objectName():
                case "tableWidget_manual_channels_settings":
                    exec("self.pushButton_"+str(selected_row+1)+".setEnabled(True)")
                    exec("self.pushButton_"+str(selected_row+1)+".setChecked(True)")
                    self.graphWidget.addItem(self.sensor_graph[selected_row])
                    if sh_set_nsh_tp[0][0].find(sensor[0][4])>=0 and eval(table_name + '.item(0,0).text()') != 'Перемычка':
                        exec(table_name + '.item(0,0).setText("Перемычка")')
                        exec(table_name + '.item(0,0).setFlags(Qt.ItemFlag.NoItemFlags)')
                        self.itm_send_command('W,Rref1=R300', 'W,Rref1-OK', 'Для канала 1 установлено опорное сопротивление 300 Ом', 'Невозможно установить опорное сопротивление 300 Ом для канала 1')      # Выбираем опорный резистор для первого канала с заглушкой
                        self.itm_send_command('W,mc1=1.0', 'W,mc1-OK', 'Для канала 1 установлен ток измерения 1 мА', 'Невозможно установить ток измерения 1 мА для канала 1')   # Устанавливаем ток измерения для первого канала с заглушкой
                    columns_order = [0,1,2,3,4,5]
                    default_text = {6:'OFF'}
                case "tableWidget_auto_channels_settings_TP":
                    columns_order = [0,2,4,5,5,3]
                    default_text = {1:'0', 6:'Периодич.', 7:'OFF', 8:'Без замечаний', 9:'В норме', 10:'В норме'}
                    chan1_jumper = True
                case "tableWidget_auto_channels_settings_TS_verifications":
                    columns_order = [0,2,4,5,5,3]
                    default_text = {1:'0', 6:'Периодич.', 7:0, 8:'Без замечаний', 9:'В норме'}
                case "tableWidget_auto_channels_settings_TS_gradiations":
                    columns_order = [0,2,5,4,6,3]
                    default_text = {1:'0', 7:0, 8:'Без замечаний', 9:'В норме'}
            for i in range(len(sensor[0])-1):                               # заносим данные датчика в таблицу
                exec(table_name + '.item(selected_row, columns_order[i]).setText(str(sensor[0][i+1]))')
            for i in default_text:                                          # затем заносим данные по-умолчанию для доп полей
                exec(table_name + '.item(selected_row, i).setText(str(default_text[i]))')
            if eval('self.pushButton_start_stop_'+ self.measuring_mode +'.text()') == "Стоп":   # если идет процесс измерения, то очищаем данные измерения этого канала
                for i in range(4): self.measured_data_channel[selected_row][i].clear()
            exec(f'self.textEdit_log_{mode}.append("{sensor[0]}")')
            if mode == 'manual': self.tableWidget_manual_measuring_result.item(1, selected_row).setText(str(sensor[0][2])) # записываем тип датчика в таблицу результатов ручного измерения
            r = False # опорное сопротивление
            i = False  # измерительный ток
            hk = None # режим холодных концов
            if ish_set[0][0].find(sensor[0][4])>=0:                             # если выбранный датчик является эталонным
                if self.etalon_channel >= 0:                                    # и если эталонный датчик уже существует, то удаляем его.
                    self.itm_send_command(f"W,sen{str(selected_row+1)}=non", f"W,sen{str(selected_row+1)}-OK", f'Для канала {selected_row+1} установлен тип датчика non', f'Невозможно установить тип датчика non  для канала {selected_row+1}')
                    self.graphWidget.removeItem(self.sensor_graph[self.etalon_channel])
                    for i in range(eval(table_name+'.columnCount()')): exec(table_name + '.item(self.etalon_channel, i).setText("")')           # сначала очищаем всю строку
                    exec(table_name + '.item(self.etalon_channel, 0).setText("+")')
                    if mode == 'manual':
                        exec("self.pushButton_"+str(self.etalon_channel+1)+".setEnabled(False)")
                        exec("self.pushButton_"+str(self.etalon_channel+1)+".setChecked(False)")
                        exec(f"self.label_temp_of_chan_{self.etalon_channel+1}.setText('-.-')")
                        for row in range(self.tableWidget_manual_measuring_result.rowCount()-1): self.tableWidget_manual_measuring_result.item(row+1, self.etalon_channel).setText("")
                if mode == 'manual': self.tableWidget_manual_measuring_result.item(2, selected_row).setText("да")
                self.etalon_channel = selected_row
                match str(sensor[0][4]):
                    case "ЭТС":
                        coef_set="coef_set_ish_ets"
                        coef_set_itm = ['ish1A', 'ish1B', 'ish1C', 'ish1D', 'ish1Rttb', 'ish1M']
                        sensor_type = 'ISH1'
                        i = '1.0'
                        r = "R300"
                    case "ПРО":
                        coef_set="coef_set_ish_pro"
                        coef_set_itm = ['ish2tPROAl', 'ish2tPROCu', 'ish2tPROPd', 'ish2tPROPt', 'ish2uPROAl', 'ish2uPROCu', 'ish2uPROPd', 'ish2uPROPt']
                        sensor_type = 'ISH2'
                    case "ППО":
                        coef_set="coef_set_ish_ppo"
                        coef_set_itm = ['ish3tPPOZn', 'ish3tPPOAl', 'ish3tPPOCu', 'ish3uPPOZn', 'ish3uPPOAl', 'ish3uPPOCu']
                        sensor_type = 'ISH3'
                    case _:
                        coef_set="error"
                        sensor_type = 'error'
                coef_set = self.db_query(f'SELECT value FROM options WHERE option="{coef_set}"')[0][0]
                coef_set = coef_set.split(',')
                exec(f'self.textEdit_log_{mode}.append("{coef_set}")')
                for coef in range(len(coef_set)):
                    coef_data = self.db_query(f'SELECT {coef_set[coef]} FROM ish_data WHERE sensor_id="{sensor[0][0]}"')[0][0]
                    if coef_data:
                        exec(f'self.textEdit_log_{mode}.append("{coef_data}")')
                        print(f"W,{coef_set_itm[coef]}={str(coef_data)}", self.portname_itm)
                        self.itm_send_command(f"W,{coef_set_itm[coef]}={str(coef_data)}", f"W,{coef_set_itm[coef]}-OK", f'Коэффициент {coef_set_itm[coef]} установлен в значение {str(coef_data)}', f'Невозможно установить коэффициент {coef_set_itm[coef]} установлен в значение {str(coef_data)}')
            else:
                # if mode == 'manual': self.tableWidget_manual_measuring_result.item(2, selected_row).setText("")         #проверить, вроде уже выше очищали весь столбец
                sensor_type = (sensor[0][4].split("(")[1])[:-1]
                print(sensor_type)
                match sensor_type:
                    case "B":
                        i = '1.0'
                        r = "R300"
                        hk = 'OFF'
                    case "R" | "S" | "A1" | "A2" | "A3" | "K" | "L" | "E" | "J" | "T" | "N" | "B":
                        i = '1.0'
                        r = "R300"
                        if mode == 'auto': hk = eval(table_name + '.item(selected_row, 7).text()')
                        else: hk = eval(table_name + '.item(selected_row, 6).text()')
                    case "10M" | "10P" | "Pt10":
                        r = "R30"
                        i = '3.0'
                    case "50M" | "Pt50" | "50P":
                        r = "R300"
                        i = '1.5'
                    case "100M" | "100P" | "Pt100":
                        r = "R300"
                        i = '0.7'
                    case "500P":
                        r = "R300"
                        i = '0.4'
            exec(f'self.textEdit_log_{mode}.append("Sensor type = {sensor_type}")')                                 # Пишем тип датчика в ИТМ
            self.itm_send_command(f'W,sen{str(selected_row+1)}={sensor_type}', f'W,sen{str(selected_row+1)}-OK', f'Для канала {selected_row+1} установлен тип датчика {sensor_type}', f'Невозможно установить тип датчика {sensor_type}  для канала {selected_row+1}')
            if hk:                                                                                                  # вкл\выкл компенсации холодных концов
                self.itm_send_command(f'W,hk{str(selected_row+1)}={hk}', f'W,hk{str(selected_row+1)}-{hk}', f'Для канала {selected_row+1} включена компенсация ХК', f'Невозможно включить компенсацию ХК для канала {selected_row+1}')
            if r: self.itm_send_command(f'W,Rref{str(selected_row+1)}={r}', f'W,Rref{str(selected_row+1)}-OK', f'Для канала {selected_row+1} установлено опорное сопротивление {r}', f'Невозможно установить опорное сопротивление {r} для канала {selected_row+1}')      # Выбираем опорный резистор
            if i: self.itm_send_command(f'W,mc{str(selected_row+1)}={i}', f'W,mc{str(selected_row+1)}-OK', f'Для канала {selected_row+1} установлен ток измерения {i}', f'Невозможно установить ток измерения {i} для канала {selected_row+1}')   # Устанавливаем ток измерения
            if mode == 'manual': self.tableWidget_manual_measuring_result.item(3, selected_row).setText(str(sensor[0][4]))

    def calculating_protocol_auto(self):
        self.result = []
        for chan in range(8):
            self.result.append(list())
            for measurement in range(eval(f'self.tableWidget_ustavka_{self.measuring_mode}.rowCount()')-1):
                self.result[chan].append(list())
                if len(self.measured_data_channel[chan][0])>0:
                    for j in range (5):
                        self.result[chan][measurement].append(0)         # 0 - средняя температура, 1 - среднее сопротивление или эдс, 2 - отклонение от эталона, 3 - дельта допустимая по ГОСТ, 4 - неопределенность поверки U
                        if j<2:
                            for i in range(8): self.result[chan][measurement][j] += self.measured_data_channel[chan][j*3][i+measurement*8]
                            self.result[chan][measurement][j] = self.result[chan][measurement][j]/8                                           # среднее t и r или e
                            print(self.result[chan][measurement][j])
                    self.result[chan][measurement].append(True)      # 5 - результат поверки
                    if chan != self.etalon_channel:
                        sensor_sn = eval(f'{self.table_name}.item(chan,2).text()')
                        print('sensor_sn = ' + sensor_sn)
                        sensor_type = self.db_query(f'SELECT type FROM sensors WHERE sn="{sensor_sn}"')[0][0]
                        sensor_class = self.db_query(f'SELECT class FROM sensors WHERE sn="{sensor_sn}"')[0][0]
                        print('sensor_type = ' + sensor_type)
                        gost_6616_id = int(self.db_query(f'SELECT id FROM gost_6616 WHERE sensor_type="{sensor_type}"')[0][0])
                        print('gost_6616_id = ' + str(gost_6616_id))
                        t_range = self.db_query(f'SELECT t_min, t_max FROM gost_6616_data WHERE gost_6616_id="{gost_6616_id}" AND class ="{sensor_class}"')
                        if self.auto_measuring_mode == 'verification_TS':
                            data_t = eval(f'self.tableWidget_ustavka_{self.measuring_mode}.item({measurement},0).text()')
                            match sensor_type:
                                case 'Pt10(Pt10)' | '10П(10P)' | '10М(10M)': U0_data_R = 10
                                case 'Pt100(Pt100)' | '100П(100P)' | '100М(100M)': U0_data_R = 50
                                case 'Pt50(Pt50)' | '50П(50P)' | '50M(50M)': U0_data_R = 50
                                case '500П(500P)': U0_data_R = 500
                            U_0 = float(self.db_query(f'SELECT val FROM U0_data WHERE chan="{chan}" AND R="{U0_data_R}" AND t="{data_t}"')[0][0])
                            C1 = float(self.db_query(f'SELECT val FROM C1_data WHERE sensor_type="{sensor_type}" AND t="{data_t}"')[0][0])
                            C2 = float(self.db_query(f'SELECT val FROM C2_data WHERE sensor_type="{sensor_type}" AND t="{data_t}"')[0][0])
                            U_et1 = (1/C1)*(U_0/math.sqrt(10))
                            delta_t = self.measured_data_channel[chan][0][measurement*8 : measurement*8+8]
                            delta_t.sort()
                            delta_t = delta_t[7]-delta_t[0]
                            U_et2 = delta_t/(2*math.sqrt(3))
                            delta_ets = 0.01
                            delta_itm = 0.015
                            U_et3 = delta_ets/2
                            U_et4 = (1/C1)*(delta_itm/3)
                            U_et5 = (1/C1)*(0.00005/math.sqrt(3))
                            notstab_ets = 0.02
                            U_et6 = notstab_ets/math.sqrt(3)
                            U_et = math.sqrt(U_et1**2 + U_et2**2 + U_et3**2 + U_et4**2 + U_et5**2 + U_et6**2)
                            U_ts1 = U_0/math.sqrt(10)
                            U_ts2 = delta_itm/3
                            U_ts3 = 0.00005/math.sqrt(3)
                            U_ts4 = C2 * 0.1 / math.sqrt(3)
                            U_ts5 = C2 * 0.8 / math.sqrt(3)
                            U_ts = math.sqrt(U_ts1**2 + U_ts2**2 + U_ts3**2 + U_ts4**2 + U_ts5**2)
                            U_c = math.sqrt(C2**2*U_et**2+U_ts**2)
                            U_r = 2*U_c
                            U = U_r/C2
                            self.result[chan][measurement][4] = U
                            print('U = ' + str(U))
                        formula = '1'
                        for t in t_range:
                            if self.result[chan][measurement][0] > t[0] and self.result[chan][measurement][0] <= t[1]:
                                formula = self.db_query(f'SELECT formula FROM gost_6616_data WHERE gost_6616_id="{gost_6616_id}" AND class ="{sensor_class}" AND t_min ="{t[0]}" AND t_max ="{t[1]}"')
                                print(formula)
                        t = self.result[chan][measurement][0]    # подставляем в формулу среднюю температуру
                        print(abs(eval(formula[0][0])))
                        self.result[chan][measurement][3] = abs(eval(formula[0][0]))
        for chan in range(8):
            for measurement in range(eval(f'self.tableWidget_ustavka_{self.measuring_mode}.rowCount()')-1):
                if len(self.measured_data_channel[chan][0])>0:
                    self.result[chan][measurement][2]=self.result[chan][measurement][0]-self.result[self.etalon_channel][measurement][0]       # отклонение от эталона
                    if abs(self.result[chan][measurement][2]) > self.result[chan][measurement][3] - self.result[chan][measurement][4]: self.result[chan][measurement][5] = False
                    print(self.result[chan][measurement][5])


    def save_measurement_data(self):
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setEnabled(False)')
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(163, 194, 194);")')
        print('Saving data...')
        if self.measuring_mode == 'manual': measuring_mode = self.measuring_mode
        else: measuring_mode = self.auto_measuring_mode
        data_tuple = (  self.counter_of_measurements,
                        measuring_mode,
                        QtCore.QDateTime.currentDateTime().toString("dd-MM-yyy HH:mm:ss"),
                        eval(f'self.lineEdit_{self.measuring_mode}_customer.text()'),
                        eval(f'self.lineEdit_{self.measuring_mode}_operators_name.text()'),
                        eval(f'self.lineEdit_{self.measuring_mode}_model.text()'),
                        eval(f'self.lineEdit_{self.measuring_mode}_number.text()'),
                        eval(f'self.lineEdit_{self.measuring_mode}_temp.text()'),
                        eval(f'self.lineEdit_{self.measuring_mode}_pressure.text()'),
                        eval(f'self.lineEdit_{self.measuring_mode}_hydro.text()'),
                        eval(f'self.tableWidget_ustavka_{self.measuring_mode}.rowCount()-1')
                        )
        if self.db_query(f"SELECT * FROM measurements WHERE id='{str(self.counter_of_measurements)}'"): # если это измерение уже существует, то мы обновляем поля
            exec(f'self.pushButton_save_data_{self.measuring_mode}.setText("Обновление...")')
            table_fields = ('id', 'type', 'datetime', 'customer', 'operator', 'megaohmmeter_model', 'megaohmmeter_sn', 't', 'p', 'h')
            for i in range(len(data_tuple)-4):  # 4 - т.к. количество уставок тоже не меняется (последнее поле), то заменяем поля с 4-го по предпоследнее
                self.db_query(f"UPDATE measurements SET {table_fields[i+3]}='{str(data_tuple[i+3])}' WHERE id='{str(self.counter_of_measurements)}';", "write")
            exec(f'self.textEdit_log_{self.measuring_mode}.append("Измерение {self.counter_of_measurements} обновлено.")')
        else:
            exec(f'self.pushButton_save_data_{self.measuring_mode}.setText("Сохранение...")')
            self.db_query("INSERT INTO measurements (id, type, datetime, customer, operator, megaohmmeter_model, megaohmmeter_sn, t, p, h, ustavlka_count) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);", "write", data_tuple)
            measurement_data_id = int(self.db_query("SELECT value FROM options WHERE option='counter_of_measurements_data'")[0][0])
            while self.db_query(f"SELECT COUNT(*) FROM measurements_data WHERE id={measurement_data_id}")[0][0]>0: measurement_data_id+=1
            measurement_result_id = int(self.db_query("SELECT value FROM options WHERE option='counter_of_measurements_result'")[0][0])
            while self.db_query(f"SELECT COUNT(*) FROM measurements_result WHERE id={measurement_result_id}")[0][0]>0: measurement_result_id+=1
            for chan in range (8):
                if len(self.measured_data_channel[chan][0]):
                    sensor_sn = eval(f'{self.table_name}.item(chan,{int(1+1*(self.measuring_mode == "auto"))}).text()')
                    sensor_id = int(self.db_query('SELECT id FROM sensors WHERE sn="'+sensor_sn+'"')[0][0])
                    for i in range(len(self.measured_data_channel[chan][0])):
                        measurement_data_id += 1
                        data_tuple = (  measurement_data_id,
                                        self.counter_of_measurements,
                                        sensor_id,
                                        self.measured_data_channel[chan][1][i].toString("dd-MM-yyyy HH:mm:ss"),
                                        self.measured_data_channel[chan][0][i],         # measurement_data_t
                                        self.measured_data_channel[chan][3][i],         # measurement_data_r_or_e
                                        self.measured_data_channel[chan][2][i])         # ustavka number
                        exec(f'self.textEdit_log_{self.measuring_mode}.append("{data_tuple}")')
                        self.db_query("INSERT INTO measurements_data (id, measurement_id, sensor_id, datetime, t, r_or_e, ustavka_number) VALUES (?, ?, ?, ?, ?, ?, ?);", "write", data_tuple)
                    if self.measuring_mode == 'auto':
                        r_isolation = eval(f'{self.table_name}.item(chan,9).text()')
                        dielectric_strenght = 'n/a'
                        hk = 'n/a'  # компенсация холодных концов
                        r_pins = 0
                        ext_inspection_notes = eval(f'{self.table_name}.item(chan, 8).text()')
                        match self.auto_measuring_mode:
                            case 'verification_TS' | 'graduation_TS':
                                r_pins = eval(f'{self.table_name}.item(chan, 7).text()')
                            case 'verification_TP' | 'calibration_TP':
                                hk = eval(f'{self.table_name}.item(chan, 7).text()')
                                dielectric_strenght = eval(f'{self.table_name}.item(chan, 9).text()')
                        data_tuple = ( self.counter_of_measurements,
                                       sensor_id,
                                       ext_inspection_notes,
                                       dielectric_strenght,
                                       r_isolation,
                                       r_pins,
                                       hk
                        )
                        self.db_query("INSERT INTO measurements_sensors_data (measurement_id, sensor_id, ext_inspection_notes, dielectric_strenght, r_isolation, r_pins, hk) VALUES (?, ?, ?, ?, ?, ?, ?);", "write", data_tuple)
                        for measurement in range(eval('self.tableWidget_ustavka_' + self.measuring_mode + '.rowCount()')-1):
                            measurement_result_id += 1
                            match self.auto_measuring_mode:
                                case 'verification_TS':
                                    r_pins = eval(f'{self.table_name}.item(chan, 7).text()')
                                    result = self.result[chan][measurement][5]
                                case 'graduation_TS':
                                    r_pins = eval(f'{self.table_name}.item(chan, 7).text()')
                                    result = ''
                                case 'verification_TP':
                                    hk = eval(f'{self.table_name}.item(chan, 7).text()')
                                    dielectric_strenght = eval(f'{self.table_name}.item(chan, 9).text()')
                                    result = self.result[chan][measurement][5]
                                case 'calibration_TP':
                                    hk = eval(f'{self.table_name}.item(chan, 7).text()')
                                    dielectric_strenght = eval(f'{self.table_name}.item(chan, 9).text()')
                                    result = ''
                            data_tuple = (  measurement_result_id,
                                            self.counter_of_measurements,                                                               # measurement_id
                                            sensor_id,                                                                                  # sensor_id
                                            self.result[chan][measurement][0],                                                          # t_mid
                                            self.result[chan][measurement][1],                                                          # r_or_e_mid
                                            self.result[chan][measurement][2],                                                          # delta_t_etalon
                                            self.result[chan][measurement][3],                                                          # delta_t_dop
                                            self.result[chan][measurement][4],                                                          # u_itm
                                            int(eval(f'self.tableWidget_ustavka_{self.measuring_mode}.item({measurement},0).text()')),  # t_ustavka
                                            int(measurement+1),                                                                         # ustavka_number
                                            result                                                                                      # result
                                            )
                            self.db_query("INSERT INTO measurements_result (id, measurement_id, sensor_id, t_mid, r_or_e_mid, delta_t_etalon, delta_t_dop, u_itm, t_ustavka, ustavka_number, result) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);", "write", data_tuple)
            self.db_query(f"UPDATE options SET value='{str(measurement_data_id)}' WHERE option='counter_of_measurements_data'", "write")
            self.textEdit_log_manual.append(f'Измерение {self.counter_of_measurements} сохранено.')
            self.db_query(f"UPDATE options SET value='{str(measurement_result_id)}' WHERE option='counter_of_measurements_result'", "write")
            self.db_query(f"UPDATE options SET value='{str(self.counter_of_measurements)}' WHERE option='counter_of_measurements'", "write")
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setText("Обновить")')
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(96, 121, 255);")')
        exec(f'self.pushButton_save_data_{self.measuring_mode}.setEnabled(True)')
        print('Saved')

    def make_protocol(self):
        self.counter_of_measurements = int(self.db_query("SELECT value FROM options WHERE option='counter_of_measurements'")[0][0])
        template = f'Protocol_{self.auto_measuring_mode}.docx'
        doc = Document(f'{os.path.dirname(__file__)}{self.template_path}{template}')
        # doc = Document(f'g:\\Мой диск\\work\\Etalon\\Calibrator\\KS1200\\Calibrator_KS1200\\templates\\{template}')
        doc.tables[0].style = doc.tables[1].style = doc.tables[2].style = 'Table Grid'
        measurement = self.db_query(f'SELECT * FROM measurements WHERE id="{self.counter_of_measurements}"')[0]
        print (measurement)
        measurement_result = self.db_query(f'SELECT * FROM measurements_result WHERE measurement_id="{self.counter_of_measurements}"')
        print (measurement_result)
        sensors_id = set()
        for meas in measurement_result:
            sensors_id.add(str(meas[2]))
        sensors = []
        ish_set = self.db_query('SELECT value FROM options WHERE option="sh_set_ish"')
        sensor_etalon = None
        for sensor_id in sensors_id:
            sensor = list(self.db_query(f'SELECT * FROM sensors WHERE id="{sensor_id}"')[0])
            sensor_meas_data = self.db_query(f'SELECT * FROM measurements_sensors_data WHERE measurement_id="{self.counter_of_measurements}" AND sensor_id="{sensor[0]}"')[0]
            sensor.extend(sensor_meas_data[3:8])
            if ish_set[0][0].find(sensor[4])>=0:
                sensor_etalon = sensor
            else:
                sensors.append(sensor)
        print(sensors)
        device_nubmer = self.db_query('SELECT value FROM options WHERE option="device_nubmer"')[0][0]
        replace_data = [['measurementId', str(measurement[0])],
                        ['measurementDate', str(measurement[2])[0:10]],
                        ['deviceNubmer', str(device_nubmer)],
                        ['megaohmmeterModel', str(measurement[5])],
                        ['megaohmmeterSn', str(measurement[6])],
                        ['temperature', str(measurement[7])],
                        ['hydro', str(measurement[9])],
                        ['pressure', str(measurement[8])],
                        ['measurementOperator', str(measurement[4])]
        ]
        if sensor_etalon:
            replace_data.extend([['sensorClass', str(sensor_etalon[3])],
                        ['sensorTRange', str(sensor_etalon[5])],
                        ['sensorName', str(sensor_etalon[1])],
                        ['sensorSn', str(sensor_etalon[2])],
                        ['sensorYearOfIssue', str(sensor_etalon[6])]])
        for rd in replace_data:
            for para in doc.paragraphs:
                for run in para.runs:
                    if rd[0] in run.text:
                        run.text = run.text.replace(rd[0], rd[1])

        doc.tables[2].autofit = False
        doc.tables[2].allow_autofit = False
        for i, sensor in enumerate(sensors):
            table0_columns_data = [sensor[1], '-', sensor[2], sensor[6][0:4], measurement[3]]
            table1_columns_data = [sensor[1], sensor[2], sensor[7], sensor[9]]
            table2_columns_data = [sensor[1], sensor[2]]
            sensor_meas_result = self.db_query(f'SELECT * FROM measurements_result WHERE measurement_id="{self.counter_of_measurements}" AND sensor_id="{sensor[0]}"')
            match measurement[1]:
                case 'verification_TS':
                    table0_columns_data.extend([sensor[3], sensor[5], 'Поверка'])
                    columns_data = [sensor[1], sensor[2], sensor[4], sensor[10]]
                    for result in sensor_meas_result:
                        columns_data.extend([result[3], result[4], result[7]])
                    columns_data.extend([sensor[3]])
                    self.make_table_rows(doc.tables[2], columns_data)
                case 'verification_TP' | 'calibration_TP':
                    meas_result = 1
                    table0_columns_data.extend([sensor[3], sensor[5]])
                    if measurement[1] == 'verification_TP': table0_columns_data.extend(['Поверка'])
                    else: table0_columns_data.extend(['Калибровка'])
                    table1_columns_data.insert(3, sensor[8])
                    rows_data = [['Температура, оС'], ['Т, оС'], ['Δ, оС'], ['Δдоп, оС'], ['']]
                    for meas in sensor_meas_result:
                        table2_columns_data.extend([meas[8], meas[3], meas[5], meas[6], ''])
                        meas_result = meas_result and meas[10]
                        if i == 0:
                            for row_data in rows_data:
                                self.make_table_rows(doc.tables[2], row_data)
                    if measurement[1] == 'verification_TP' and i == 0:
                        self.make_table_rows(doc.tables[2], ['Результат'])
                        if meas_result: table2_columns_data.extend(['Годен'])
                        else: table2_columns_data.extend(['Не годен'])
                    self.make_table_columns(doc.tables[2], table2_columns_data)
                case 'graduation_TS':
                    table0_columns_data.extend([sensor[4], sensor[3], sensor[5]])
                    table2_columns_data.extend([sensor[10]])
                    for j, meas in enumerate(sensor_meas_result):
                        rows_data = [[f'Т{j+1}, оС'], [f'R{j+1}, оС'], ['']]
                        table2_columns_data.extend([meas[3], meas[4], ''])
                        if i == 0:
                            for row_data in rows_data:
                                self.make_table_rows(doc.tables[2], row_data)
                    self.make_table_columns(doc.tables[2], table2_columns_data)

            self.make_table_rows(doc.tables[0], table0_columns_data)
            self.make_table_rows(doc.tables[1], table1_columns_data)

        # doc.save(f'g:\\Мой диск\\work\\Etalon\\Calibrator\\KS1200\\Calibrator_KS1200\\{template}')
        doc.save(f'{os.path.dirname(__file__)}\\{template}')
        doc = Document()
        print('Done')


    def make_table_rows(self, table, columns_data):
        cells = table.add_row().cells
        for i, column_data in enumerate(columns_data):
            cells[i].text = str(column_data)
            paragraphs = cells[i].paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(9)

    def make_table_columns(self, table, rows_data):
        cells = table.add_column(Cm(5)).cells
        for i, row_data in enumerate(rows_data):
            cells[i].text = str(row_data)
            paragraphs = cells[i].paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(9)

    def coef_edit(self):
        button = QApplication.instance().sender()
        button_name = "self."+button.objectName()
        par = button_name.split("_")
        coef = par[2]
        ish_data_sensor = int(par[3])
        coef_val=self.db_query(f"SELECT {coef} FROM ish_data WHERE sensor_id={str(ish_data_sensor)}")[0][0]
        if coef_val is None:
            coef_val = 0
        else:
            coef_val = float(coef_val)

        coef_val, ok = QInputDialog.getDouble(button, coef, 'Введите новое значение:', coef_val, -2147483647, 2147483647,8)

        if ok:
            self.db_query(f"UPDATE ish_data SET {coef}={str(coef_val)} WHERE sensor_id={str(ish_data_sensor)}", "write")
            if coef_val is not None and coef_val != 0:
                exec(button_name+".setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(38, 0, 255);')")
            else:
                exec(button_name+".setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0);')")

    def change_tab(self,stack, tab):
        self.tab_path.append([stack, stack.currentIndex()])
        stack.setCurrentIndex(tab)

    def back_tab(self):
        tab = self.tab_path.pop()
        tab[0].setCurrentIndex(tab[1])

    def db_query(self, query, action='read', data_tuple = None):
        try:
            sqlite_connection = sqlite3.connect('ks1200.db')
            cursor = sqlite_connection.cursor()
            # print("База данных создана и успешно подключена к SQLite")
            # print("результат запроса:")
            # print (query+":")
            if data_tuple: cursor.execute(query, data_tuple)
            else: cursor.execute(query)
            if action =="write": sqlite_connection.commit()
            record = cursor.fetchall()
            # print(record)
            cursor.close()
        except sqlite3.Error as error:
            print("Класс исключения: ", error.__class__)
            print("Исключение", error.args)
            print("Печать подробноcтей исключения SQLite: ")
            exc_type, exc_value, exc_tb = sys.exc_info()
            print(traceback.format_exception(exc_type, exc_value, exc_tb))
        finally:
            if (sqlite_connection):
                sqlite_connection.close()
            return record

    def device_number_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='device_nubmer'", "write")

    def device_name_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='device_name'", "write")

    def device_veryfing_date_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='device_veryfing_date'", "write")

    def device_next_veryfing_date_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='device_next_veryfing_date'", "write")

    def device_produce_date_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='device_produce_date'", "write")

    def comport_bu7_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='com_port_bu7_{platform}'", "write")
        self.portname_bu7 = new_data

    def comport_itm_changed(self, new_data):
        self.db_query(f"UPDATE options SET value='{new_data}' WHERE option='com_port_itm_{platform}'", "write")
        self.portname_itm = new_data

    def open_side_menu(self):
        if self.widget_side_menu.isVisible() == False:
                self.toolButton_side_menu.setGeometry(QtCore.QRect(800, 9, 24, 582))
                self.widget_side_menu.show()
        else:
                self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
                self.widget_side_menu.close()

    def open_settings(self):
        self.widget_side_menu.close()
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        self.stackedWidget_main.setCurrentIndex(self.tab_settings)

    def open_settings_ISH(self):
        self.widget_side_menu.close()
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        self.change_tab(self.stackedWidget_main, self.tab_settings_ish)
    
    def open_settings_device(self):
        self.stackedWidget_main.setCurrentIndex(self.tab_settings_ish)

    def archive_delete_record(self):
        selected_items = self.tableWidget_archive.selectedItems()
        measurement_id = self.tableWidget_archive.verticalHeaderItem(selected_items[0].row()).text()
        self.db_query(f'DELETE FROM measurements WHERE id = {measurement_id}','write')
        self.db_query(f'DELETE FROM measurements_data WHERE measurement_id = {measurement_id}','write')
        self.db_query(f'DELETE FROM measurements_result WHERE measurement_id = {measurement_id}','write')
        self.db_query(f'DELETE FROM measurements_sensors_data WHERE measurement_id = {measurement_id}','write')
        self.tableWidget_archive.removeRow(selected_items[0].row())
        
    def archive_view_record(self):
        pass

    def open_mode_manual(self):
        self.widget_side_menu.close()
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        self.measuring_mode = 'manual'
        self.etalon_channel = -1
        self.table_name = 'self.tableWidget_manual_channels_settings'
        self.textEdit_log_manual.clear()
        self.timer_measuring.setInterval(5000)
        self.clear_tabs_and_itm()
        for channel in range(8):
            if eval(f"self.pushButton_{channel+1}.isChecked()"):
                exec(f"self.pushButton_{channel+1}.setEnabled(False)")
                exec(f"self.pushButton_{channel+1}.setChecked(False)")
                exec(f"self.label_temp_of_chan_{channel+1}.setText('-.-')")
                self.graphWidget.removeItem(self.sensor_graph[channel])
                if len(self.data_for_graph[channel][0]): self.data_for_graph[channel][0].clear()
                if len(self.data_for_graph[channel][1]): self.data_for_graph[channel][1].clear()
                self.itm_send_command(f"W,sen{str(channel+1)}=non", f"W,sen{str(channel+1)}-OK", f'Для канала {channel+1} установлен тип датчика non', f'Невозможно установить тип датчика non  для канала {channel+1}')
        if len(self.data_for_graph[8][0]): self.data_for_graph[8][0].clear()
        if len(self.data_for_graph[8][1]): self.data_for_graph[8][1].clear()
        self.itm_send_command("W,stopmeas=1", "W,stopmeas-OK", "Останов измерений в ИТМ", "Невозможно остановить измерения в ИТМ")
        self.sensor_graph[8].setData(self.data_for_graph[8][1], self.data_for_graph[8][0])
        self.stackedWidget_main.setCurrentIndex(self.tab_manual)
 
    def open_mode_auto(self):
        self.widget_side_menu.close()
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        self.auto_measuring_step_buttons(False)
        self.pushButton_verification_TP.setChecked(False)
        self.pushButton_verification_TS.setChecked(False)
        self.pushButton_graduation_TS.setChecked(False)
        self.pushButton_calibration_TP.setChecked(False)
        self.stackedWidget_auto_measuring_tabs.setCurrentIndex(0)
        self.measuring_mode = 'auto'
        self.textEdit_log_auto.clear()
        self.timer_measuring.setInterval(7000)
        self.stackedWidget_main.setCurrentIndex(self.tab_auto)

    def open_sensors(self):
        self.widget_side_menu.close() 
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        self.change_tab(self.stackedWidget_main, self.tab_sensors)

    def open_archive(self):
        self.widget_side_menu.close()
        self.toolButton_side_menu.setGeometry(QtCore.QRect(991, 9, 24, 582))
        self.change_tab(self.stackedWidget_main, self.tab_archive)

    def auto_measuring_step_buttons(self, status = True):
        self.pushButton_measuring_type.setEnabled(status)
        self.pushButton_tpoints_sensors.setEnabled(status)
        self.pushButton_measuring_settings.setEnabled(status)
        self.pushButton_start_stop_tab.setEnabled(status)
        self.pushButton_protocol.setEnabled(status)

    # Выбор режима автоматического измерения 
    def set_auto_verification_TS_mode(self):
        self.stackedWidget_measuring_types_tables.setCurrentIndex(1)
        self.pushButton_verification_TP.setChecked(False)
        self.pushButton_verification_TS.setChecked(True)
        self.pushButton_graduation_TS.setChecked(False)
        self.pushButton_calibration_TP.setChecked(False)
        self.auto_measuring_mode = 'verification_TS'
        self.table_name = 'self.tableWidget_auto_channels_settings_TS_verifications'
        self.auto_measuring_step_buttons()
        self.clear_tabs_and_itm()
        self.etalon_channel = -1
        for row in range(2):
            self.tableWidget_ustavka_auto.insertRow(row)
            item = QtWidgets.QTableWidgetItem()
            self.tableWidget_ustavka_auto.setItem(row, 0, item)
            self.tableWidget_ustavka_auto.item(row,0).setText(str(row*100))
        self.tableWidget_ustavka_auto.setEnabled(False)
        self.stackedWidget_auto_measuring_tabs.setCurrentIndex(1)

    def set_auto_graduation_TS_mode(self):
        self.stackedWidget_measuring_types_tables.setCurrentIndex(2)
        self.pushButton_verification_TP.setChecked(False)
        self.pushButton_verification_TS.setChecked(False)
        self.pushButton_graduation_TS.setChecked(True)
        self.pushButton_calibration_TP.setChecked(False)
        self.auto_measuring_mode = 'graduation_TS'
        self.table_name = 'self.tableWidget_auto_channels_settings_TS_gradiations'
        self.auto_measuring_step_buttons()
        self.clear_tabs_and_itm()
        self.etalon_channel = -1
        self.stackedWidget_auto_measuring_tabs.setCurrentIndex(1)

    def set_auto_calibration_TP_mode(self):
        self.stackedWidget_measuring_types_tables.setCurrentIndex(0)
        self.pushButton_verification_TP.setChecked(False)
        self.pushButton_verification_TS.setChecked(False)
        self.pushButton_graduation_TS.setChecked(False)
        self.pushButton_calibration_TP.setChecked(True)
        self.auto_measuring_mode = 'calibration_TP'
        self.table_name = 'self.tableWidget_auto_channels_settings_TP'
        self.auto_measuring_step_buttons()
        self.clear_tabs_and_itm()
        self.etalon_channel = -1
        self.stackedWidget_auto_measuring_tabs.setCurrentIndex(1)

    def set_auto_verification_TP_mode(self):
        self.stackedWidget_measuring_types_tables.setCurrentIndex(0)
        self.pushButton_verification_TP.setChecked(True)
        self.pushButton_verification_TS.setChecked(False)
        self.pushButton_graduation_TS.setChecked(False)
        self.pushButton_calibration_TP.setChecked(False)
        self.auto_measuring_mode = 'verification_TP'
        self.table_name = 'self.tableWidget_auto_channels_settings_TP'
        self.auto_measuring_step_buttons()
        self.clear_tabs_and_itm()
        self.etalon_channel = -1
        self.stackedWidget_auto_measuring_tabs.setCurrentIndex(1)

    def set_settings_basic_tab(self):
        self.stackedWidget_settings_tabs.setCurrentIndex(self.settings_basic_tab)
        self.pushButton_settings_basic.setChecked(True)
        self.pushButton_settings_connections.setChecked(False)
        self.pushButton_settings_verifying.setChecked(False)
        self.pushButton_settings_device.setChecked(False)
        self.pushButton_settings_system_update.setChecked(False)

    def set_settings_connections_tab(self):
        self.stackedWidget_settings_tabs.setCurrentIndex(self.settings_connections_tab)
        self.pushButton_settings_basic.setChecked(False)
        self.pushButton_settings_connections.setChecked(True)
        self.pushButton_settings_verifying.setChecked(False)
        self.pushButton_settings_device.setChecked(False)
        self.pushButton_settings_system_update.setChecked(False)
        
    def set_settings_verifying_tab(self):
        self.stackedWidget_settings_tabs.setCurrentIndex(self.settings_verifying_tab)
        self.pushButton_settings_basic.setChecked(False)
        self.pushButton_settings_connections.setChecked(False)
        self.pushButton_settings_verifying.setChecked(True)
        self.pushButton_settings_device.setChecked(False)
        self.pushButton_settings_system_update.setChecked(False)
        
    def set_settings_device_tab(self):
        self.stackedWidget_settings_tabs.setCurrentIndex(self.settings_device_tab)
        self.pushButton_settings_basic.setChecked(False)
        self.pushButton_settings_connections.setChecked(False)
        self.pushButton_settings_verifying.setChecked(False)
        self.pushButton_settings_device.setChecked(True)
        self.pushButton_settings_system_update.setChecked(False)
        
    def set_settings_system_update_tab(self):
        self.stackedWidget_settings_tabs.setCurrentIndex(self.settings_system_update_tab)
        self.pushButton_settings_basic.setChecked(False)
        self.pushButton_settings_connections.setChecked(False)
        self.pushButton_settings_verifying.setChecked(False)
        self.pushButton_settings_device.setChecked(False)
        self.pushButton_settings_system_update.setChecked(True)

    def clear_tabs_and_itm(self):
        for channel in range(8):
            self.itm_send_command(f"W,sen{str(channel+1)}=non", f"W,sen{str(channel+1)}-OK", f'Для канала {channel+1} установлен тип датчика non', f'Невозможно установить тип датчика non  для канала {channel+1}')
        tables = ['auto_channels_settings_TP', 'auto_channels_settings_TS_verifications', 'auto_channels_settings_TS_gradiations', 'manual_channels_settings']
        for table in tables:
            for row in range(eval(f'self.tableWidget_{table}.rowCount()')):
                for col in range(eval(f'self.tableWidget_{table}.columnCount()')):
                    if col == 0 and row == 0 and table == 'auto_channels_settings_TP':
                        self.itm_send_command('W,Rref1=R300', 'W,Rref1-OK', 'Для канала 1 установлено опорное сопротивление 300 Ом', 'Невозможно установить опорное сопротивление 300 Ом для канала 1')      # Выбираем опорный резистор для первого канала с заглушкой
                        self.itm_send_command('W,mc1=1.0', 'W,mc1-OK', 'Для канала 1 установлен ток измерения 1 мА', 'Невозможно установить ток измерения 1 мА для канала 1')   # Устанавливаем ток измерения для первого канала с заглушкой
                        val = "Перемычка"
                    elif col == 0:
                        val = '+'
                    else:
                        val = ""
                    exec(f'self.tableWidget_{table}.item({str(row)},{str(col)}).setText("{val}")')
        modes = ['auto', 'manual']
        for mode in modes:
            if eval(f'self.tableWidget_ustavka_{mode}.rowCount()') >1:
                for row in range(eval(f'self.tableWidget_ustavka_{mode}.rowCount()')-1):
                    exec(f'self.tableWidget_ustavka_{mode}.removeRow(0)')
            exec(f'self.tableWidget_ustavka_{mode}.setEnabled(True)')

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_F11:
            if MainWindow.windowState() == Qt.WindowState.WindowFullScreen:
                MainWindow.showNormal()
            else: MainWindow.showFullScreen()
        event.accept()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "KS1200"))
        self.label_device_nubmer.setText(_translate("MainWindow", "Заводской номер:"))
        self.lineEdit_device_number.setText(_translate("MainWindow", self.db_query("SELECT value FROM options WHERE option='device_nubmer'")[0][0]))
        self.label_device_name.setText(_translate("MainWindow", "Наименование прибора:"))
        self.lineEdit_device_name.setText(_translate("MainWindow", self.db_query("SELECT value FROM options WHERE option='device_name'")[0][0]))
        self.label_device_next_veryfing_date.setText(_translate("MainWindow", "Очередная поверка:"))
        self.lineEdit_device_next_veryfing_date.setText(_translate("MainWindow", self.db_query("SELECT value FROM options WHERE option='device_veryfing_date'")[0][0]))
        self.label_devic_veryfing_date.setText(_translate("MainWindow", "Дата поверки:"))
        self.lineEdit_device_produce_date.setText(_translate("MainWindow", self.db_query("SELECT value FROM options WHERE option='device_next_veryfing_date'")[0][0]))
        self.label_device_produce_date.setText(_translate("MainWindow", "Дата производства:"))
        self.lineEdit_device_veryfing_date.setText(_translate("MainWindow", self.db_query("SELECT value FROM options WHERE option='device_produce_date'")[0][0]))
        self.label_comport_bu7.setText(_translate("MainWindow", "Номер COM порта БУ7:"))
        self.label_comport_itm.setText(_translate("MainWindow", "Номер COM порта ИТМ:"))
        self.pushButton_settings_ISH.setText(_translate("MainWindow", "Настройки ИСХ"))
        self.pushButton_settings_basic.setText(_translate("MainWindow", "Основные настройки"))
        self.pushButton_settings_connections.setText(_translate("MainWindow", "Настройка подключений"))
        self.pushButton_settings_verifying.setText(_translate("MainWindow", "Поверка калибратора (пломбировка переключателя)"))
        self.pushButton_settings_device.setText(_translate("MainWindow", "Настройка оборудования"))
        self.pushButton_settings_system_update.setText(_translate("MainWindow", "Обновление системы"))
        item = self.tableWidget_settings_ISH.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Описание датчика"))
        item = self.tableWidget_settings_ISH.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "Тип"))
        item = self.tableWidget_settings_ISH.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "Коэффициенты"))
        __sortingEnabled = self.tableWidget_settings_ISH.isSortingEnabled()
        self.tableWidget_settings_ISH.setSortingEnabled(False)
        self.tableWidget_settings_ISH.setSortingEnabled(__sortingEnabled)
        items_text = ["Описание датчика", "Заводской номер", "Класс", "Тип", "Диапазон", "Дата выпуска"]
        for col in range(len(items_text)):
            item = self.tableWidget_sensors.horizontalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))
        items_text = ["Заказчик", "Тип измерения", "Дата и время"]
        for col in range(len(items_text)):
            item = self.tableWidget_archive.horizontalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))
        self.pushButton_ish_back.setText(_translate("MainWindow", "Назад"))
        self.pushButton_sensors_back.setText(_translate("MainWindow", "Назад"))
        self.pushButton_sensor_add.setText(_translate("MainWindow", "Добавить"))
        self.pushButton_start_stop_manual.setText(_translate("MainWindow", "Старт"))
        self.pushButton_manual_channels_settings.setText(_translate("MainWindow", "1. Настройка"))
        self.pushButton_manual_measuring.setText(_translate("MainWindow", "2. Измерение"))
        self.pushButton_manual_result.setText(_translate("MainWindow", "3. Результат"))
        item = self.tableWidget_ustavka_manual.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Уставка"))
        item = self.tableWidget_ustavka_manual.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "Время\n(чч:мм)"))
        __sortingEnabled = self.tableWidget_ustavka_manual.isSortingEnabled()
        self.tableWidget_ustavka_manual.setSortingEnabled(False)
        item = self.tableWidget_ustavka_manual.item(0, 0)
        item.setText(_translate("MainWindow", "+"))
        self.tableWidget_ustavka_manual.setSortingEnabled(__sortingEnabled)
        for row in range(8):
            item = self.tableWidget_manual_channels_settings.verticalHeaderItem(row)
            item.setText(_translate("MainWindow", str(row+1)))
            item = self.tableWidget_auto_channels_settings_TP.verticalHeaderItem(row)
            item.setText(_translate("MainWindow", str(row+1)))
            item = self.tableWidget_auto_channels_settings_TS_verifications.verticalHeaderItem(row)
            item.setText(_translate("MainWindow", str(row+1)))
            item = self.tableWidget_auto_channels_settings_TS_gradiations.verticalHeaderItem(row)
            item.setText(_translate("MainWindow", str(row+1)))
        items_text = ["Тип СИ", "Заводской\nномер", "Класс", "СХ", "Диапазон,\n°С", "Дата выпуска", "Комп.\nХК"]
        for col in range(len(items_text)):
            item = self.tableWidget_manual_channels_settings.horizontalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))
        self.groupBox_channels_buttons.setTitle(_translate("MainWindow", ""))
        for i in range(8):
            exec("self.pushButton_"+str(i+1)+".setText(_translate('MainWindow','"+str(i+1)+"'))")
            exec("self.label_temp_of_chan_"+str(i+1)+".setText(_translate('MainWindow','-.-'))")
        self.pushButton_9.setText(_translate('MainWindow',"К"))
        self.label_temp_of_chan_9.setText(_translate('MainWindow','-.-'))
        self.groupBox_navigate_buttons.setTitle(_translate("MainWindow", ""))
        self.pushButton_move_graph_left.setText(_translate("MainWindow", "←"))
        self.pushButton_move_graph_right.setText(_translate("MainWindow", "→"))
        self.pushButton_move_graph_up.setText(_translate("MainWindow", "↑"))
        self.pushButton_move_graph_down.setText(_translate("MainWindow", "↓"))
        self.pushButton_scale_vertical_up.setText(_translate("MainWindow", "UD+"))
        self.pushButton_scale_vertical_down.setText(_translate("MainWindow", "UD-"))
        self.pushButton_scale_horizontal_up.setText(_translate("MainWindow", "LR+"))
        self.pushButton_scale_horizontal_down.setText(_translate("MainWindow", "LR-"))
        self.pushButton_scale_auto.setText(_translate("MainWindow", "Auto"))
        self.pushButton_view_table.setText(_translate("MainWindow", "Table"))
        items_text = ["Измерительный канал:", "Номер датчика:", "Эталон:", "Статическая характеристика:", "Значение:", "Текущее отклонение:", "Дрейф на минуту:"]
        for col in range(len(items_text)):
            item = self.tableWidget_manual_measuring_result.verticalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))
        self.pushButton_view_graph.setText(_translate("MainWindow", "Graph"))
        self.groupBox_bottom_buttons.setTitle(_translate("MainWindow", "GroupBox"))
        self.pushButton_measuring_type.setText(_translate("MainWindow", "1. Тип измерения"))
        self.pushButton_tpoints_sensors.setText(_translate("MainWindow", "2. Температурные точки\nи данные датчиков"))
        self.pushButton_measuring_settings.setText(_translate("MainWindow", "3. Параметры измерения"))
        self.pushButton_start_stop_tab.setText(_translate("MainWindow", "4. Запуск/остановка"))
        self.pushButton_protocol.setText(_translate("MainWindow", "5. Протокол"))
        self.pushButton_verification_TS.setText(_translate("MainWindow", "Поверка ТС"))
        self.pushButton_graduation_TS.setText(_translate("MainWindow", "Градуировка ТС"))
        self.pushButton_calibration_TP.setText(_translate("MainWindow", "Калибровка ТП"))
        self.pushButton_verification_TP.setText(_translate("MainWindow", "Поверка ТП"))
        __sortingEnabled = self.tableWidget_auto_channels_settings_TS_verifications.isSortingEnabled()
        self.tableWidget_auto_channels_settings_TS_verifications.setSortingEnabled(False)
        self.tableWidget_auto_channels_settings_TS_verifications.setSortingEnabled(__sortingEnabled)
        __sortingEnabled = self.tableWidget_auto_channels_settings_TS_gradiations.isSortingEnabled()
        self.tableWidget_auto_channels_settings_TS_gradiations.setSortingEnabled(False)
        self.tableWidget_auto_channels_settings_TS_gradiations.setSortingEnabled(__sortingEnabled)
        __sortingEnabled = self.tableWidget_auto_channels_settings_TP.isSortingEnabled()
        self.tableWidget_auto_channels_settings_TP.setSortingEnabled(False)
        self.tableWidget_auto_channels_settings_TP.setSortingEnabled(__sortingEnabled)
        items_text = ["Тип СИ", "Рег.\nномер", "Заводской\nномер", "Год\nвыпуска", "Класс\nдопуска", "Диапазон,\n°С", "Вид испытаний", "Комп.\nХК", "Замечания по\nвнешнему осмотру", "Электрическая\nпрочность изоляции", "Электросопротивле-\nние изоляции"]
        for col in range(len(items_text)):
            item = self.tableWidget_auto_channels_settings_TP.horizontalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))
        items_text = ["Тип СИ", "Рег.\nномер", "Заводской\nномер", "Год\nвыпуска", "Класс\nдопуска", "Диапазон,\n°С", "Вид испытаний", "R выв.", "Замечания по\nвнешнему осмотру", "Электросопротивле-\nние изоляции"]
        for col in range(len(items_text)):
            item = self.tableWidget_auto_channels_settings_TS_verifications.horizontalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))            
        items_text = ["Тип СИ", "Рег.\nномер", "Заводской\nномер", "Год\nвыпуска", "НСХ", "Класс\nдопуска", "Диапазон,\n°С", "R выв.", "Замечания по\nвнешнему осмотру", "Электросопротивле-\nние изоляции"]
        for col in range(len(items_text)):
            item = self.tableWidget_auto_channels_settings_TS_gradiations.horizontalHeaderItem(col)
            item.setText(_translate("MainWindow", items_text[col]))  
        item = self.tableWidget_ustavka_auto.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Темп.\nточки"))
        __sortingEnabled = self.tableWidget_ustavka_auto.isSortingEnabled()
        self.tableWidget_ustavka_auto.setSortingEnabled(False)
        item = self.tableWidget_ustavka_auto.item(0, 0)
        item.setText(_translate("MainWindow", "+"))
        self.tableWidget_ustavka_auto.setSortingEnabled(__sortingEnabled)
        self.label_auto_number.setText(_translate("MainWindow", "номер"))
        self.label_auto_measurement_number_text.setText(_translate("MainWindow", "Измерение номер: "))
        self.label_auto_fio.setText(_translate("MainWindow", "Ф.И.О. оператора"))
        self.label_auto_model.setText(_translate("MainWindow", "модель"))
        self.label_auto_megaommetr.setText(_translate("MainWindow", "Данные о мегаомметре"))
        self.label_auto_customer.setText(_translate("MainWindow", "Данные для протокола"))
        self.label_auto_customer_2.setText(_translate("MainWindow", "Заказчик"))
        self.lineEdit_auto_customer.setText("Заказчик ")
        self.label_auto_t_atm.setText(_translate("MainWindow", "Температура окружающего воздуха. °С"))
        self.label_auto_p_atm.setText(_translate("MainWindow", "Атмосферное давлениеб кПа"))
        self.label_auto_hydro.setText(_translate("MainWindow", "Относительная влажность воздуха, %"))
        self.label_auto_article_name.setText(_translate("MainWindow", "Условия поверки"))
        self.label_manual_number.setText(_translate("MainWindow", "номер"))
        self.label_manual_measurement_number_text.setText(_translate("MainWindow", "Измерение номер: "))
        self.label_manual_fio.setText(_translate("MainWindow", "Ф.И.О. оператора"))
        self.label_manual_model.setText(_translate("MainWindow", "модель"))
        self.label_manual_megaommetr.setText(_translate("MainWindow", "Данные о мегаомметре"))
        self.label_manual_customer.setText(_translate("MainWindow", "Данные для протокола"))
        self.label_manual_customer_2.setText(_translate("MainWindow", "Заказчик"))
        self.lineEdit_manual_customer.setText("Заказчик")
        self.label_manual_t_atm.setText(_translate("MainWindow", "Температура окружающего воздуха. °С"))
        self.label_manual_p_atm.setText(_translate("MainWindow", "Атмосферное давлениеб кПа"))
        self.label_manual_hydro.setText(_translate("MainWindow", "Относительная влажность воздуха, %"))
        self.label_manual_article_name.setText(_translate("MainWindow", "Условия поверки"))
        self.pushButton_save_data_manual.setText(_translate("MainWindow", "Сохранить"))
        self.pushButton_start_stop_auto.setText(_translate("MainWindow", "Старт"))
        self.label_measuring_finished.setText(_translate("MainWindow", "Измерения завершены"))
        self.pushButton_print_auto.setText(_translate("MainWindow", "Печать"))
        self.pushButton_preview_auto.setText(_translate("MainWindow", "Просмотр"))
        self.pushButton_save_data_auto.setText(_translate("MainWindow", "Сохранить"))
        self.label_memory.setText(_translate("MainWindow", "Внутренняя память"))
        self.pushButton_archive_delete.setText(_translate("MainWindow", "Удалить"))
        self.pushButton_archive_export.setText(_translate("MainWindow", "Экспорт"))
        self.pushButton_archive_view.setText(_translate("MainWindow", "Просмотр"))
        self.pushButton_archive_back.setText(_translate("MainWindow", "Назад"))
        self.toolButton_side_menu.setText(_translate("MainWindow", "М"))
        self.pushButton_mode_mnual.setText(_translate("MainWindow", "Ручной режим"))
        self.pushButton_mode_auto.setText(_translate("MainWindow", "Авто-режим"))
        self.pushButton_sensors.setText(_translate("MainWindow", "Датчики"))
        self.pushButton_archive.setText(_translate("MainWindow", "Архив"))
        self.pushButton_settings.setText(_translate("MainWindow", "Настройка"))

class Change_ustavka_dialog(QDialog):
    def __init__(self, temperature, time, selected_row, rowCount, MainWindow):
        QDialog.__init__(self)
        self.setupUi(self)
        result = exchange_data_bu7(MainWindow, 'R,Tmin2,Tmax2', MainWindow.portname_bu7)
        min = extract_val_param(result, 'Tmin2')[:extract_val_param(result, 'Tmin2').find('.')]
        if min == 'errval':
            exec(f'MainWindow.textEdit_log_{MainWindow.measuring_mode}.append("Нет связи с БУ7")')
            min = 50
        max = extract_val_param(result, 'Tmax2')[:extract_val_param(result, 'Tmax2').find('.')]
        if max == 'errval':
            exec(f'MainWindow.textEdit_log_{MainWindow.measuring_mode}.append("Нет связи с БУ7")')
            max = 1300
        # Отслеживаем, что температуру можно задавать только на повышение.
        # if selected_row > 0:
            # min = MainWindow.tableWidget_ustavka_manual.item(selected_row-1,0).text()
        # if selected_row < rowCount - 2:
        #     max = MainWindow.tableWidget_ustavka_manual.item(selected_row+1,0).text()
        self.timeEdit.setMinimumTime(QtCore.QTime.fromString("00:01", "hh:mm"))
        if temperature == "+":
            self.pushButton_delete.setEnabled(False)
            time = "00:01"
            temperature = min
        self.spinBox.setRange(int(min),int(max))
        self.spinBox.setValue(int(temperature))
        self.timeEdit.setTime(QtCore.QTime.fromString(time, "hh:mm"))
        self.MainWindow = MainWindow
        self.buttonBox.accepted.connect(lambda: self.accept_data(temperature, time, selected_row, rowCount))
        self.buttonBox.rejected.connect(self.reject_data)
        self.pushButton_delete.clicked.connect(lambda: self.delete_data(selected_row))

    def accept_data(self, temperature, time, selected_row, rowCount):
        if selected_row+1 == rowCount:
            self.MainWindow.tableWidget_ustavka_manual.insertRow(selected_row)
            item = QtWidgets.QTableWidgetItem()
            self.MainWindow.tableWidget_ustavka_manual.setItem(selected_row, 0, item)
            item = QtWidgets.QTableWidgetItem()
            self.MainWindow.tableWidget_ustavka_manual.setItem(selected_row, 1, item)

        self.MainWindow.tableWidget_ustavka_manual.item(selected_row,0).setText(str(self.spinBox.value()))
        self.MainWindow.tableWidget_ustavka_manual.item(selected_row,1).setText(self.timeEdit.time().toString("hh:mm"))
        self.close()

    def reject_data(self):
        self.close()

    def delete_data(self, selected_row):
        self.MainWindow.tableWidget_ustavka_manual.removeRow(selected_row)
        self.MainWindow.tableWidget_ustavka_manual.selectionModel().clearCurrentIndex()
        self.close()

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.setModal(True)
        Dialog.resize(362, 106)
        self.verticalLayout = QtWidgets.QVBoxLayout(Dialog)
        self.verticalLayout.setObjectName("verticalLayout")
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.label_temperature = QtWidgets.QLabel(Dialog)
        self.label_temperature.setObjectName("label_temperature")
        self.horizontalLayout.addWidget(self.label_temperature)
        self.spinBox = QtWidgets.QSpinBox(Dialog)
        self.spinBox.setObjectName("spinBox")
        self.horizontalLayout.addWidget(self.spinBox)
        self.label_time = QtWidgets.QLabel(Dialog)
        self.label_time.setObjectName("label_time")
        self.horizontalLayout.addWidget(self.label_time)
        self.timeEdit = QtWidgets.QTimeEdit(Dialog)
        self.timeEdit.setObjectName("timeEdit")
        self.horizontalLayout.addWidget(self.timeEdit)
        self.verticalLayout.addLayout(self.horizontalLayout)
        self.groupBox = QtWidgets.QGroupBox(Dialog)
        self.groupBox.setObjectName("groupBox")
        self.buttonBox = QtWidgets.QDialogButtonBox(self.groupBox)
        self.buttonBox.setGeometry(QtCore.QRect(141, 15, 191, 31))
        self.buttonBox.setOrientation(QtCore.Qt.Orientation.Horizontal)
        self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.StandardButton.Cancel|QtWidgets.QDialogButtonBox.StandardButton.Ok)
        self.buttonBox.setCenterButtons(False)
        self.buttonBox.setObjectName("buttonBox")
        self.pushButton_delete = QtWidgets.QPushButton(self.groupBox)
        self.pushButton_delete.setGeometry(QtCore.QRect(10, 15, 101, 31))
        self.pushButton_delete.setObjectName("pushButton_delete")
        self.verticalLayout.addWidget(self.groupBox)

        self.retranslateUi(Dialog)
        self.buttonBox.accepted.connect(Dialog.accept) # type: ignore
        self.buttonBox.rejected.connect(Dialog.reject) # type: ignore
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Настройка уставки"))
        self.label_time.setText(_translate("Dialog", "Время (Ч:ММ):"))
        self.label_temperature.setText(_translate("Dialog", "Температура, °С:"))
        # self.groupBox.setTitle(_translate("Dialog", "GroupBox"))
        self.pushButton_delete.setText(_translate("Dialog", "Удалить"))

class Change_confidence_error_dialog(QDialog):  # окно изменения строки таблицы
    def __init__(self, selected_row, ParentWindow):
        QDialog.__init__(self)
        self.setupUi(self)
        self.ParentWindow = ParentWindow
        if self.ParentWindow.tableWidget_confidence_error.item(selected_row,0).text() !='+':
            self.spinbox_t.setValue(float(self.ParentWindow.tableWidget_confidence_error.item(selected_row,0).text()))
            self.spinbox_delta_t.setValue(float(self.ParentWindow.tableWidget_confidence_error.item(selected_row,1).text()))
            self.spinbox_delta_t_max.setValue(float(self.ParentWindow.tableWidget_confidence_error.item(selected_row,2).text()))
            self.pushButton_delete.setEnabled(True)
        self.pushButton_delete.clicked.connect(lambda: self.delete_row(selected_row))
        self.buttonBox.accepted.connect(lambda: self.accept_data(selected_row))
        self.buttonBox.rejected.connect(self.reject_data)

    def accept_data(self, selected_row):
        if self.ParentWindow.tableWidget_confidence_error.rowCount() < 4 and self.ParentWindow.tableWidget_confidence_error.item(selected_row,0).text() == '+':
            self.ParentWindow.tableWidget_confidence_error.insertRow(selected_row)
            for col in range(3):
                item = QtWidgets.QTableWidgetItem()
                item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                self.ParentWindow.tableWidget_confidence_error.setItem(selected_row,col, item)
        self.ParentWindow.tableWidget_confidence_error.item(selected_row,0).setText(str(self.spinbox_t.value()))
        self.ParentWindow.tableWidget_confidence_error.item(selected_row,1).setText(str(self.spinbox_delta_t.value()))
        self.ParentWindow.tableWidget_confidence_error.item(selected_row,2).setText(str(self.spinbox_delta_t_max.value()))
        rowCount = self.ParentWindow.tableWidget_confidence_error.rowCount()
        if self.ParentWindow.tableWidget_confidence_error.item(rowCount-1,0).text() == '+': rowCount -= 1
        data = []
        for row in range (rowCount):
            data.append(list())
            for col in range(3):
                data[row].append(float(self.ParentWindow.tableWidget_confidence_error.item(row,col).text()))
        data.sort()
        for row in range (rowCount):
            for col in range(3):
                self.ParentWindow.tableWidget_confidence_error.item(row,col).setText(str(data[row][col]))

        self.close()

    def reject_data(self):
        self.close()

    def delete_row(self, selected_row):
        self.ParentWindow.tableWidget_confidence_error.removeRow(selected_row)
        self.ParentWindow.tableWidget_confidence_error.selectionModel().clearCurrentIndex()
        rowCount = self.ParentWindow.tableWidget_confidence_error.rowCount()
        if rowCount < 4 and self.ParentWindow.tableWidget_confidence_error.item(rowCount-1,0).text() != '+':
            self.ParentWindow.tableWidget_confidence_error.insertRow(rowCount)
            for col in range(3):
                item = QtWidgets.QTableWidgetItem()
                item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                self.ParentWindow.tableWidget_confidence_error.setItem(rowCount, col, item)
            self.ParentWindow.tableWidget_confidence_error.item(rowCount,0).setText('+')
        self.close()
    
    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.setModal(True)
        Dialog.resize(475, 100)
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_t = QtWidgets.QLabel(Dialog)
        self.label_t.setGeometry(QtCore.QRect(15, 15, 50, 22))
        self.label_t.setFont(font)
        self.label_t.setObjectName("label_t")
        self.label_delta_t = QtWidgets.QLabel(Dialog)
        self.label_delta_t.setGeometry(QtCore.QRect(165, 15, 50, 22))
        self.label_delta_t.setFont(font)
        self.label_delta_t.setObjectName("label_delta_t")
        self.label_delta_t_max = QtWidgets.QLabel(Dialog)
        self.label_delta_t_max.setGeometry(QtCore.QRect(305, 15, 90, 22))
        self.label_delta_t_max.setFont(font)
        self.label_delta_t_max.setObjectName("label_delta_t_max")
        self.spinbox_t = QtWidgets.QDoubleSpinBox(Dialog)
        self.spinbox_t.setGeometry(QtCore.QRect(55, 15, 90, 22))
        self.spinbox_t.setFont(font)
        self.spinbox_t.setRange(-1000, 1000)
        self.spinbox_t.setDecimals(3)
        self.spinbox_t.setStepType(1)
        self.spinbox_t.setObjectName("spinbox_t")
        self.spinbox_delta_t = QtWidgets.QDoubleSpinBox(Dialog)
        self.spinbox_delta_t.setGeometry(QtCore.QRect(215, 15, 70, 22))
        self.spinbox_delta_t.setDecimals(3)
        self.spinbox_delta_t.setStepType(1)
        self.spinbox_delta_t.setFont(font)
        self.spinbox_delta_t.setObjectName("spinbox_delta_t")
        self.spinbox_delta_t_max = QtWidgets.QDoubleSpinBox(Dialog)
        self.spinbox_delta_t_max.setGeometry(QtCore.QRect(390, 15, 70, 22))
        self.spinbox_delta_t_max.setDecimals(3)
        self.spinbox_delta_t_max.setStepType(1)
        self.spinbox_delta_t_max.setFont(font)
        self.spinbox_delta_t_max.setObjectName("spinbox_delta_t_max")
        self.pushButton_delete = QtWidgets.QPushButton(Dialog)
        self.pushButton_delete.setGeometry(QtCore.QRect(15, 60, 101, 25))
        self.pushButton_delete.setObjectName("pushButton_delete")
        self.pushButton_delete.setEnabled(False)
        self.buttonBox = QtWidgets.QDialogButtonBox(Dialog)
        self.buttonBox.setGeometry(QtCore.QRect(290, 55, 170, 31))
        self.buttonBox.setOrientation(QtCore.Qt.Horizontal)
        self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.Cancel|QtWidgets.QDialogButtonBox.Ok)
        self.buttonBox.setObjectName("buttonBox")
        self.retranslateUi(Dialog)
        self.buttonBox.accepted.connect(Dialog.accept)
        self.buttonBox.rejected.connect(Dialog.reject)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Доверительная погрешность"))
        self.label_t.setText(_translate("Dialog_sensor_edit", "t, °C:"))
        self.label_delta_t.setText(_translate("Dialog_sensor_edit", "δt, °C:"))
        self.label_delta_t_max.setText(_translate("Dialog_sensor_edit", "max δt, °C:"))
        self.pushButton_delete.setText(_translate("Dialog", "Удалить"))

class Confidence_error_dialog(QDialog):
    def __init__(self, MainWindow):
        self.MainWindow = MainWindow
        button = QApplication.instance().sender()
        button_name = "self."+button.objectName()
        self.ish_data_sensor = int(button_name.split("_")[3])
        QDialog.__init__(self)
        self.setupUi(self)
        self.buttonBox.accepted.connect(lambda: self.accept_data())
        self.buttonBox.rejected.connect(self.reject_data)
        self.pushButton_clear.clicked.connect(lambda: self.clear_data())

    def accept_data(self):
        self.MainWindow.db_query(f'DELETE FROM confidence_error WHERE sensor_id="{str(self.ish_data_sensor)}";', 'write')
        rowCount = self.tableWidget_confidence_error.rowCount()
        if self.tableWidget_confidence_error.item(rowCount-1,0).text() == '+': rowCount -= 1
        for row in range (rowCount):
            data = []
            data.append(int(self.ish_data_sensor))
            for col in range(3): data.append(float(self.tableWidget_confidence_error.item(row,col).text()))
            self.MainWindow.db_query(f'INSERT INTO confidence_error (sensor_id, t, delta_ets, instability_ets) VALUES (?, ?, ?, ?);', 'write', data)
        self.close()

    def reject_data(self):
        self.close()

    def clear_data(self):
        self.tableWidget_confidence_error.setRowCount(0)
        self.tableWidget_confidence_error.selectionModel().clearCurrentIndex()
        self.tableWidget_confidence_error.insertRow(0)
        for col in range(3):
            item = QtWidgets.QTableWidgetItem()
            item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
            item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
            item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.tableWidget_confidence_error.setItem(0, col, item)
        self.tableWidget_confidence_error.item(0,0).setText('+')

    def Change_confidence_error(self, selected_row):
        Change_ustavka_dialog_inst = Change_confidence_error_dialog(selected_row, self)
        Change_ustavka_dialog_inst.show()
        Change_ustavka_dialog_inst.exec()

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.setModal(True)
        Dialog.resize(350, 230)
        font = QtGui.QFont()
        font.setPointSize(12)
        self.verticalLayout = QtWidgets.QVBoxLayout(Dialog)
        self.verticalLayout.setObjectName("verticalLayout")
        self.tableWidget_confidence_error = QtWidgets.QTableWidget(Dialog)
        # self.tableWidget_confidence_error.setGeometry(QtCore.QRect(0, 0, 100, 50))
        self.tableWidget_confidence_error.setObjectName("tableWidget_confidence_error")
        columnWidths = [50, 80, 100]
        self.tableWidget_confidence_error.setColumnCount(len(columnWidths))
        for column in range(len(columnWidths)):
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            self.tableWidget_confidence_error.setHorizontalHeaderItem(column, item)
            self.tableWidget_confidence_error.setColumnWidth(column, columnWidths[column])
        self.tableWidget_confidence_error.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        rowCount = self.MainWindow.db_query(f'SELECT COUNT(*) FROM confidence_error WHERE sensor_id="{str(self.ish_data_sensor)}"')[0][0]
        self.tableWidget_confidence_error.setRowCount(rowCount)
        if rowCount > 0:
            delta_data = self.MainWindow.db_query(f'SELECT * FROM confidence_error WHERE sensor_id="{str(self.ish_data_sensor)}"')
            for row in range(rowCount):
                for col in range(3):
                    item = QtWidgets.QTableWidgetItem(str(delta_data[row][col+2]))
                    item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                    item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                    self.tableWidget_confidence_error.setItem(row, col, item)
        if rowCount < 4:
            self.tableWidget_confidence_error.insertRow(rowCount)
            for col in range(3):
                item = QtWidgets.QTableWidgetItem()
                item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignHCenter)
                item.setFlags(Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEnabled)
                item.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                self.tableWidget_confidence_error.setItem(rowCount, col, item)
            self.tableWidget_confidence_error.item(rowCount,0).setText('+')
        self.tableWidget_confidence_error.cellClicked.connect(self.Change_confidence_error)
        self.verticalLayout.addWidget(self.tableWidget_confidence_error)
        self.groupBox = QtWidgets.QGroupBox(Dialog)
        self.groupBox.setObjectName("groupBox")
        self.groupBox.setMinimumHeight(51)
        self.buttonBox = QtWidgets.QDialogButtonBox(self.groupBox)
        self.buttonBox.setGeometry(QtCore.QRect(133, 10, 191, 31))
        self.buttonBox.setOrientation(QtCore.Qt.Orientation.Horizontal)
        self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.StandardButton.Cancel|QtWidgets.QDialogButtonBox.StandardButton.Ok)
        self.buttonBox.setCenterButtons(False)
        self.buttonBox.setObjectName("buttonBox")
        self.pushButton_clear = QtWidgets.QPushButton(self.groupBox)
        self.pushButton_clear.setGeometry(QtCore.QRect(5, 13, 100, 25))
        self.pushButton_clear.setObjectName("pushButton_clear")
        self.verticalLayout.addWidget(self.groupBox)

        self.retranslateUi(Dialog)
        self.buttonBox.accepted.connect(Dialog.accept) # type: ignore
        self.buttonBox.rejected.connect(Dialog.reject) # type: ignore
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Доверительная погрешность"))
        item = self.tableWidget_confidence_error.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "t, °C"))
        item = self.tableWidget_confidence_error.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "δt, °C"))
        item = self.tableWidget_confidence_error.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "max δt, °C"))
        self.pushButton_clear.setText(_translate("Dialog", "Очистить"))


class Change_temp_dialog(QDialog):
    def __init__(self, temperature, selected_row, rowCount, MainWindow):
        QDialog.__init__(self)
        self.setupUi(self)
        result = exchange_data_bu7(MainWindow, 'R,Tmin2,Tmax2', MainWindow.portname_bu7)
        min = extract_val_param(result, 'Tmin2')[:extract_val_param(result, 'Tmin2').find('.')]
        if min == 'errval':
            exec(f'MainWindow.textEdit_log_{MainWindow.measuring_mode}.append("Нет связи с БУ7")')
            min = 50
        max = extract_val_param(result, 'Tmax2')[:extract_val_param(result, 'Tmax2').find('.')]
        if max == 'errval':
            exec(f'MainWindow.textEdit_log_{MainWindow.measuring_mode}.append("Нет связи с БУ7")')
            max = 1300
        if temperature == "+":
            self.pushButton_delete.setEnabled(False)
            temperature = min
        self.spinBox.setRange(int(min),int(max))
        self.spinBox.setValue(int(temperature))
        self.MainWindow = MainWindow
        self.buttonBox.accepted.connect(lambda: self.accept_data(temperature, selected_row, rowCount))
        self.buttonBox.rejected.connect(self.reject_data)
        self.pushButton_delete.clicked.connect(lambda: self.delete_data(selected_row))

    def accept_data(self, temperature, selected_row, rowCount):
        if selected_row+1 == rowCount:
            self.MainWindow.tableWidget_ustavka_auto.insertRow(selected_row)
            item = QtWidgets.QTableWidgetItem()
            self.MainWindow.tableWidget_ustavka_auto.setItem(selected_row, 0, item)
        self.MainWindow.tableWidget_ustavka_auto.item(selected_row,0).setText(str(self.spinBox.value()))
        self.close()

    def reject_data(self):
        self.close()

    def delete_data(self, selected_row):
        self.MainWindow.tableWidget_ustavka_auto.removeRow(selected_row)
        self.MainWindow.tableWidget_ustavka_auto.selectionModel().clearCurrentIndex()
        self.close()

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.setModal(True)
        Dialog.resize(362, 106)
        self.verticalLayout = QtWidgets.QVBoxLayout(Dialog)
        self.verticalLayout.setObjectName("verticalLayout")
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.label_temperature = QtWidgets.QLabel(Dialog)
        self.label_temperature.setObjectName("label_temperature")
        self.horizontalLayout.addWidget(self.label_temperature)
        self.spinBox = QtWidgets.QSpinBox(Dialog)
        self.spinBox.setObjectName("spinBox_temperature")
        self.horizontalLayout.addWidget(self.spinBox)
        self.verticalLayout.addLayout(self.horizontalLayout)
        self.groupBox = QtWidgets.QGroupBox(Dialog)
        self.groupBox.setObjectName("groupBox")
        self.buttonBox = QtWidgets.QDialogButtonBox(self.groupBox)
        self.buttonBox.setGeometry(QtCore.QRect(141, 15, 191, 31))
        self.buttonBox.setOrientation(QtCore.Qt.Orientation.Horizontal)
        self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.StandardButton.Cancel|QtWidgets.QDialogButtonBox.StandardButton.Ok)
        self.buttonBox.setCenterButtons(False)
        self.buttonBox.setObjectName("buttonBox")
        self.pushButton_delete = QtWidgets.QPushButton(self.groupBox)
        self.pushButton_delete.setGeometry(QtCore.QRect(10, 15, 101, 31))
        self.pushButton_delete.setObjectName("pushButton_delete")
        self.verticalLayout.addWidget(self.groupBox)

        self.retranslateUi(Dialog)
        self.buttonBox.accepted.connect(Dialog.accept) # type: ignore
        self.buttonBox.rejected.connect(Dialog.reject) # type: ignore
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Установка температуры"))
        self.label_temperature.setText(_translate("Dialog", "Температура, °С:"))
        # self.groupBox.setTitle(_translate("Dialog", "GroupBox"))
        self.pushButton_delete.setText(_translate("Dialog", "Удалить"))

class Sensor_edit_dialog(QDialog):
    def __init__(self, MainWindow, selected_row):
        QDialog.__init__(self)
        self.setupUi(self)
        self.MainWindow = MainWindow
        self.count_of_types = []
        all_types = ['sh_set_ish','sh_set_nsh_tp','sh_set_nsh_tr']
        types = []
        for sh in all_types:
            self.count_of_types.append(len(MainWindow.db_query(f"SELECT value FROM options WHERE option='{sh}'")[0][0].split(',')))
            for i in MainWindow.db_query(f"SELECT value FROM options WHERE option='{sh}'")[0][0].split(','):
                types.append(i)
        self.comboBox_sensor_type.addItems(types)
        if selected_row >= 0:
            self.pushButton_sensor_delete.setEnabled(True)
            sensor_sn = MainWindow.tableWidget_sensors.item(selected_row,1).text()
            self.sensor = MainWindow.db_query("SELECT * FROM sensors WHERE sn='"+sensor_sn+"'")
            self.lineEdit_sensor_name.setText(str(self.sensor[0][1]))
            self.lineEdit_sensor_sn.setText(str(self.sensor[0][2]))
            self.dateEdit_sensor_year_of_issue.setDate(QDate.fromString(self.sensor[0][6],'yyyy-MM-dd'))
            self.comboBox_sensor_type.setCurrentIndex(types.index(self.sensor[0][4]))
            if self.sensor[0][4] == 'ЭТС': self.change_sensor_classes()     # т.к. ЭТС первый в списке и событие изменения не генерится
            sensor_id = int(self.sensor[0][0])
        else:
            self.pushButton_sensor_delete.setEnabled(False)
            sensor_id = int(MainWindow.db_query("SELECT value FROM options WHERE option='counter_of_sensors'")[0][0])+1
        self.buttonBox.accepted.connect(lambda: self.accept_data(selected_row, sensor_id))
        self.buttonBox.rejected.connect(self.reject_data)
        self.pushButton_sensor_delete.clicked.connect(lambda: self.delete_data(selected_row, sensor_id))
        self.MainWindow.tableWidget_sensors.clearSelection()

    def change_sensor_classes(self):
        sensor_type = self.comboBox_sensor_type.currentText()
        classes = []
        for i in MainWindow.db_query(f"SELECT classes FROM gost_6616 WHERE sensor_type='{sensor_type}'")[0][0].split(','): classes.append(i)
        self.comboBox_sensor_class.clear()
        self.comboBox_sensor_class.addItems(classes)
        if 'sensor' in self.__dict__: self.comboBox_sensor_class.setCurrentIndex(classes.index(self.sensor[0][3]))

    def chanhe_sensor_range(self):
        sensor_gost_id = MainWindow.db_query(f"SELECT id FROM gost_6616 WHERE sensor_type='{self.comboBox_sensor_type.currentText()}'")[0][0]
        sensor_class = self.comboBox_sensor_class.currentText()
        if not sensor_class: return
        # print(sensor_gost_id)
        # print(f"SELECT t_min FROM gost_6616_data WHERE gost_6616_id='{sensor_gost_id}' AND class='{sensor_class}'")
        class_gost_t_min = MainWindow.db_query(f"SELECT t_min FROM gost_6616_data WHERE gost_6616_id='{sensor_gost_id}' AND class='{sensor_class}'")[0][0]
        class_gost_t_max = MainWindow.db_query(f"SELECT t_max FROM gost_6616_data WHERE gost_6616_id='{sensor_gost_id}' AND class='{sensor_class}'")[0][0]
        # print(class_gost_t_min)
        # print(class_gost_t_max)
        self.spinbox_sensor_t_min.setRange(int(class_gost_t_min), int(class_gost_t_max))
        self.spinbox_sensor_t_max.setRange(int(class_gost_t_min), int(class_gost_t_max))
        self.comboBox_sensor_type.currentText()
        self.comboBox_sensor_class.currentText()
        if 'sensor' in self.__dict__:
            self.spinbox_sensor_t_max.setValue(int(self.sensor[0][5].split('...')[1]))
            self.spinbox_sensor_t_min.setValue(int(self.sensor[0][5].split('...')[0]))

    def accept_data(self, selected_row, sensor_id):
        data_tuple = (  sensor_id,
                        self.lineEdit_sensor_name.text(),
                        self.lineEdit_sensor_sn.text(),
                        self.comboBox_sensor_class.currentText(),
                        self.comboBox_sensor_type.currentText(),
                        str(self.spinbox_sensor_t_min.value()) + '...' + str(self.spinbox_sensor_t_max.value()),
                        self.dateEdit_sensor_year_of_issue.date().toString('yyyy-MM-dd'),
                        )   
        if selected_row >= 0:
            MainWindow.db_query("UPDATE sensors SET id = ?, name = ?, sn = ?, class = ?, type = ?, t_range = ?, year_of_issue = ? WHERE id ='" + str(sensor_id)+"';", "write", data_tuple)
        else:
            MainWindow.db_query("INSERT INTO sensors (id, name, sn, class, type, t_range, year_of_issue) VALUES (?, ?, ?, ?, ?, ?, ?);", "write", data_tuple)
            MainWindow.db_query(f'UPDATE options SET value = {str(sensor_id)} WHERE option = "counter_of_sensors"', 'write')
            # Если добавлен датчик ИСХ, то добаввляем строку в таблицу коэффициентов.
            if self.comboBox_sensor_type.currentText() == "ЭТС" or self.comboBox_sensor_type.currentText() == "ППО" or self.comboBox_sensor_type.currentText() == "ПРО":
                ish_data_id = int(MainWindow.db_query("SELECT value FROM options WHERE option='counter_of_sensors_ish'")[0][0])+1
                MainWindow.db_query(f'UPDATE options SET value = {str(ish_data_id)} WHERE option = "counter_of_sensors_ish"', 'write')
                data_tuple_ish = ( ish_data_id, sensor_id, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0)
                MainWindow.db_query("INSERT INTO ish_data (id, sensor, A, B, C, D, W, Rttb, M, tZn, tAl, tCu, eZn, eAl, eCu, tPd, tPt, ePd, ePt) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);", "write", data_tuple_ish)
            selected_row = MainWindow.tableWidget_sensors.rowCount()
            MainWindow.tableWidget_sensors.insertRow(selected_row)
            font = QtGui.QFont()
            font.setPointSize(14)
            item = QtWidgets.QTableWidgetItem()
            item.setFont(font)
            item.setText(str(selected_row+1))
            MainWindow.tableWidget_sensors.setVerticalHeaderItem(selected_row, item)
            for col in range(6):
                item = QtWidgets.QTableWidgetItem()
                item.setFont(font)
                item.setFlags(Qt.ItemFlag.ItemIsSelectable)
                MainWindow.tableWidget_sensors.setItem(selected_row, col, item)
        for col in range(len(data_tuple)-1):
            MainWindow.tableWidget_sensors.item(selected_row,col).setText(data_tuple[col+1])
        self.close()

    def reject_data(self):
        self.close()

    def delete_data(self, selected_row, sensor_id):
        MainWindow.tableWidget_sensors.removeRow(selected_row)
        MainWindow.tableWidget_sensors.selectionModel().clearCurrentIndex()
        MainWindow.db_query(f'DELETE from sensors WHERE id ="{str(sensor_id)}";','write')
        if self.comboBox_sensor_type.currentText() == "ЭТС" or self.comboBox_sensor_type.currentText() == "ППО" or self.comboBox_sensor_type.currentText() == "ПРО":
            MainWindow.db_query(f"DELETE from ish_data WHERE sensor_id ='{str(sensor_id)}';","write")
        self.close()

    def setupUi(self, Dialog_sensor_edit):
        Dialog_sensor_edit.setObjectName("Dialog_sensor_edit")
        Dialog_sensor_edit.setModal(True)
        Dialog_sensor_edit.resize(531, 162)
        font = QtGui.QFont()
        font.setPointSize(12)
        Dialog_sensor_edit.setFont(font)
        self.label_sensor_name = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_name.setGeometry(QtCore.QRect(10, 20, 70, 22))
        self.label_sensor_name.setFont(font)
        self.label_sensor_name.setObjectName("label_sensor_name")
        self.label_sensor_sn = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_sn.setGeometry(QtCore.QRect(270, 20, 110, 22))
        self.label_sensor_sn.setFont(font)
        self.label_sensor_sn.setObjectName("label_sensor_sn")
        self.label_sensor_type = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_type.setGeometry(QtCore.QRect(10, 50, 41, 22))
        self.label_sensor_type.setFont(font)
        self.label_sensor_type.setObjectName("label_sensor_type")
        self.label_sensor_class = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_class.setGeometry(QtCore.QRect(185, 50, 51, 22))
        self.label_sensor_class.setFont(font)
        self.label_sensor_class.setObjectName("label_sensor_class")
        self.label_sensor_t_range = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_t_range.setGeometry(QtCore.QRect(10, 80, 250, 22))
        self.label_sensor_t_range.setFont(font)
        self.label_sensor_t_range.setObjectName("label_sensor_t_range")
        self.label_sensor_t_max = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_t_max.setGeometry(QtCore.QRect(360, 80, 50, 22))
        self.label_sensor_t_max.setFont(font)
        self.label_sensor_t_max.setObjectName("label_sensor_t_max")
        self.lineEdit_sensor_name = QtWidgets.QLineEdit(Dialog_sensor_edit)
        self.lineEdit_sensor_name.setGeometry(QtCore.QRect(75, 20, 180, 22))
        self.lineEdit_sensor_name.setFont(font)
        self.lineEdit_sensor_name.setText("")
        self.lineEdit_sensor_name.setObjectName("lineEdit_sensor_name")
        self.lineEdit_sensor_sn = QtWidgets.QLineEdit(Dialog_sensor_edit)
        self.lineEdit_sensor_sn.setGeometry(QtCore.QRect(380, 20, 140, 22))
        self.lineEdit_sensor_sn.setFont(font)
        self.lineEdit_sensor_sn.setText("")
        self.lineEdit_sensor_sn.setObjectName("lineEdit_sensor_sn")
        self.comboBox_sensor_type = QtWidgets.QComboBox(Dialog_sensor_edit)
        self.comboBox_sensor_type.setGeometry(QtCore.QRect(50, 50, 120, 22))
        self.comboBox_sensor_type.setObjectName("comboBox_sensor_type")
        self.comboBox_sensor_type.currentIndexChanged.connect(self.change_sensor_classes)
        self.comboBox_sensor_class = QtWidgets.QComboBox(Dialog_sensor_edit)
        self.comboBox_sensor_class.setGeometry(QtCore.QRect(240, 50, 50, 22))
        self.comboBox_sensor_class.setObjectName("comboBox_sensor_class")
        self.comboBox_sensor_class.currentIndexChanged.connect(self.chanhe_sensor_range)
        self.spinbox_sensor_t_min = QtWidgets.QSpinBox(Dialog_sensor_edit)
        self.spinbox_sensor_t_min.setGeometry(QtCore.QRect(270, 80, 70, 22))
        self.spinbox_sensor_t_min.setFont(font)
        self.spinbox_sensor_t_min.setObjectName("spinbox_sensor_t_min")
        self.spinbox_sensor_t_max = QtWidgets.QSpinBox(Dialog_sensor_edit)
        self.spinbox_sensor_t_max.setGeometry(QtCore.QRect(410, 80, 70, 22))
        self.spinbox_sensor_t_max.setFont(font)
        self.spinbox_sensor_t_max.setObjectName("spinbox_sensor_t_max")
        self.label_sensor_year_of_issue = QtWidgets.QLabel(Dialog_sensor_edit)
        self.label_sensor_year_of_issue.setGeometry(QtCore.QRect(300, 50, 111, 22))
        self.label_sensor_year_of_issue.setFont(font)
        self.label_sensor_year_of_issue.setObjectName("label_sensor_year_of_issue")
        self.dateEdit_sensor_year_of_issue = QtWidgets.QDateEdit(Dialog_sensor_edit)
        self.dateEdit_sensor_year_of_issue.setGeometry(QtCore.QRect(410, 50, 110, 22))
        self.dateEdit_sensor_year_of_issue.setFont(font)
        self.dateEdit_sensor_year_of_issue.setObjectName("dateEdit_sensor_year_of_issue")
        self.buttonBox = QtWidgets.QDialogButtonBox(Dialog_sensor_edit)
        self.buttonBox.setGeometry(QtCore.QRect(340, 120, 171, 31))
        self.buttonBox.setOrientation(QtCore.Qt.Horizontal)
        self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.Cancel|QtWidgets.QDialogButtonBox.Ok)
        self.buttonBox.setObjectName("buttonBox")
        self.pushButton_sensor_delete = QtWidgets.QPushButton(Dialog_sensor_edit)
        self.pushButton_sensor_delete.setGeometry(QtCore.QRect(20, 120, 81, 31))
        self.pushButton_sensor_delete.setObjectName("pushButton_sensor_delete")

        self.retranslateUi(Dialog_sensor_edit)
        self.buttonBox.accepted.connect(Dialog_sensor_edit.accept) # type: ignore
        self.buttonBox.rejected.connect(Dialog_sensor_edit.reject) # type: ignore
        QtCore.QMetaObject.connectSlotsByName(Dialog_sensor_edit)

    def retranslateUi(self, Dialog_sensor_edit):
        _translate = QtCore.QCoreApplication.translate
        Dialog_sensor_edit.setWindowTitle(_translate("Dialog_sensor_edit", "Редактирование датчика"))
        self.label_sensor_name.setText(_translate("Dialog_sensor_edit", "Тип СИ:"))
        self.label_sensor_sn.setText(_translate("Dialog_sensor_edit", "Заводской №:"))
        self.label_sensor_class.setText(_translate("Dialog_sensor_edit", "Класс:"))
        self.label_sensor_type.setText(_translate("Dialog_sensor_edit", "Тип:"))
        self.label_sensor_t_range.setText(_translate("Dialog_sensor_edit", "Диапазон температур:         t min"))
        self.label_sensor_t_max.setText(_translate("Dialog_sensor_edit", "t max"))
        self.label_sensor_year_of_issue.setText(_translate("Dialog_sensor_edit", "Дата выпуска:"))
        self.pushButton_sensor_delete.setText(_translate("Dialog_sensor_edit", "Удалить"))


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = MainWindow()
    MainWindow.show()
    sys.exit(app.exec())
