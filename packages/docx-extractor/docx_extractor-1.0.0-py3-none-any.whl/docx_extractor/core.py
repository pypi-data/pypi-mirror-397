import os
import tempfile
import shutil
import glob
from pathlib import Path

from docx import Document
import olefile
import zipfile

from .processors import process_excel, process_text
from .toon_formatter import format_toon


def extract_text_and_tables(docx_path):
    """Extract non-empty paragraphs and tables from DOCX."""
    doc = Document(docx_path)
    
    # Filter empty paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Tables: filter empty cells too
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:  # only non-empty cells
                    row_data.append(cell_text)
            if row_data:  # only non-empty rows
                table_data.append(row_data)
        if table_data:
            tables.append(table_data)
    
    return {"paragraphs": paragraphs, "tables": tables}


def extract_embedded_files(docx_path, output_dir):
    """Extract embedded files from DOCX."""
    os.makedirs(output_dir, exist_ok=True)
    
    with tempfile.TemporaryDirectory() as temp_dir:
        with zipfile.ZipFile(docx_path) as zf:
            zf.extractall(temp_dir)
        
        embeddings_dir = os.path.join(temp_dir, "word", "embeddings")
        if not os.path.exists(embeddings_dir):
            return []
        
        extracted = []
        for emb_path in glob.glob(os.path.join(embeddings_dir, "*")):
            filename = os.path.basename(emb_path)
            try:
                if olefile.isOleFile(emb_path):
                    with olefile.OleFileIO(emb_path) as ole:
                        for stream in ("Package", "package"):
                            if ole.exists(stream):
                                data = ole.openstream(stream).read()
                                out_name = filename.replace(".bin", "")
                                out_path = os.path.join(output_dir, out_name)
                                with open(out_path, "wb") as f:
                                    f.write(data)
                                extracted.append(out_path)
                                break
                        else:
                            shutil.copy(emb_path, output_dir)
                            extracted.append(os.path.join(output_dir, filename))
                else:
                    shutil.copy(emb_path, output_dir)
                    extracted.append(os.path.join(output_dir, filename))
            except Exception:
                pass
        return extracted


def process_file_recursive(path, output_base, level=0, max_depth=5):
    """Recursive processor for all file types."""
    if level > max_depth:
        return None
    
    path = str(path)
    ext = Path(path).suffix.lower()
    
    result = {
        "file": os.path.basename(path),
        "path": path,
        "type": ext,
        "level": level,
    }
    
    child_dir = os.path.join(output_base, f"level_{level}")
    os.makedirs(child_dir, exist_ok=True)
    
    if ext == ".docx":
        # DOCX content
        content = extract_text_and_tables(path)
        result["content"] = content
        
        # Embedded files
        emb_dir = os.path.join(child_dir, "embedded")
        embedded = extract_embedded_files(path, emb_dir)
        result["embedded_count"] = len(embedded)
        
        # Recurse
        children = []
        for emb in embedded:
            child_result = process_file_recursive(emb, output_base, level + 1, max_depth)
            if child_result:
                children.append(child_result)
        result["children"] = children
    
    elif ext in (".xlsx", ".xls"):
        excel_data = process_excel(path)
        if excel_data:
            result["excel"] = excel_data
    
    elif ext in (".txt", ".sql"):
        text = process_text(path)
        if text.strip():
            result["text"] = text
    
    return result


def extract_docx_recursive(docx_path, output_dir="output", max_depth=5):
    """Main entry point: process DOCX recursively."""
    os.makedirs(output_dir, exist_ok=True)
    result = process_file_recursive(docx_path, output_dir, max_depth=max_depth)
    return result
