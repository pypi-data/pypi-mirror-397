from typing import Any, Optional
from pydantic import BaseModel
from agentmail import AgentMail

import io
import filetype
import pymupdf
import docx


type Kwargs = dict[str, Any]


class Attachment(BaseModel):
    text: Optional[str] = None
    error: Optional[str] = None
    file_type: Optional[str] = None


def list_inboxes(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.list(**kwargs)


def get_inbox(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.get(**kwargs)


def create_inbox(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.create(**kwargs)


def delete_inbox(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.delete(**kwargs)


def list_threads(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.threads.list(**kwargs)


def get_thread(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.threads.get(**kwargs)


def get_attachment(client: AgentMail, kwargs: Kwargs):
    it = client.threads.get_attachment(
        thread_id=kwargs["thread_id"], attachment_id=kwargs["attachment_id"]
    )
    file_bytes = b"".join(it)

    file_kind = filetype.guess(file_bytes)
    file_type = file_kind.mime if file_kind else None

    text = ""
    if file_type == "application/pdf":
        for page in pymupdf.Document(stream=file_bytes):
            text += page.get_text() + "\n"
    elif (
        file_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        for paragraph in docx.Document(io.BytesIO(file_bytes)).paragraphs:
            text += paragraph.text + "\n"
    else:
        return Attachment(
            error=f"Unsupported file type: {file_type or 'unknown'}",
            file_type=file_type,
        )

    return Attachment(text=text, file_type=file_type)


def send_message(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.messages.send(**kwargs)


def reply_to_message(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.messages.reply(**kwargs)


def update_message(client: AgentMail, kwargs: Kwargs):
    return client.inboxes.messages.update(**kwargs)
