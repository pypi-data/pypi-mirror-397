from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config
import re

class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = self.format_markdown(markdown_content)
        # print(self.markdown_content)
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        # print(self.html_content)
        # 逐行处理HTML代码，在<ol>、<li>、<ul>标签前添加换行
        self.html_content = self._process_html_line_breaks(self.html_content)
        print('='*50)
        print(self.html_content)
        FileUtils.create_paths(Config.BASE_PATH)

    def format_markdown(self, markdown_text):
        """
        简化Markdown格式处理，优化多层列表嵌套支持
        
        Args:
            markdown_text (str): 原始Markdown文本
            
        Returns:
            str: 格式化后的Markdown文本
        """
        lines = markdown_text.split('\n')
        processed_lines = []
        
        # 跟踪列表嵌套级别
        list_levels = []
        
        for i, line in enumerate(lines):
            line = line.rstrip()  # 移除行尾空白
            
            # 跳过空行
            if not line.strip():
                processed_lines.append('')
                continue
            
            # 处理标题行：去掉前面的空格，确保#号后面有空格
            if line.strip().startswith('#'):
                line = re.sub(r'^\s*(#+)\s*', r'\1 ', line)
            
            # 处理表格和标题之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and line.strip().startswith('#'):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理表格和列表之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and re.match(r'^\s*(\d+\.|\-)', line.strip()):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理冒号后的换行 - 确保嵌套列表有正确的换行
            if '：' in line and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # 如果下一行是列表项，在当前行后添加空行
                if re.match(r'^\s*(\d+\.|\-)', next_line):
                    line = line + '\n'
            
            # 特殊处理：对于以数字开头且后面是粗体的行，进行转义
            # 避免被Markdown库误识别为有序列表项
            # if re.match(r'^\s*\d+\.\s*\*\*', line):
            #     # 转义数字点号，使其不被识别为列表项
            #     line = re.sub(r'^(\s*)(\d+)\.(\s*\*\*)', r'\1\2\\.\3', line)
                
            #     # 确保转义后的行后面有适当的换行，避免HTML结构问题
            #     if i + 1 < len(lines) and re.match(r'^\s*\-\s*\*\*', lines[i + 1]):
            #         line = line + '\n\n'  # 添加两个换行，确保正确的HTML结构
            
            # 计算当前行的缩进级别
            line_stripped = line.lstrip()
            indent_count = len(line) - len(line_stripped)
            # print(f"line: {line}，indent_count: {indent_count}")
            if indent_count > 0 :
                indent_count = indent_count * 2
            # 保持原有缩进级别，不要翻倍，避免过度缩进
                tab_indent = ' ' * indent_count
                a = tab_indent + line_stripped
                processed_lines.append(a)
            else:
                processed_lines.append(line)
        return '\n'.join(processed_lines)

    def to_html(self):
        return self.html_content

    def _process_html_line_breaks(self, html_content):
        """
        处理HTML代码，去除<li>、<ul>、<ol>标签内的<p>标签，并处理换行结构
        
        Args:
            html_content (str): 原始HTML内容
            
        Returns:
            str: 处理后的HTML内容
        """
        # 逐个处理<li>标签内的<p>和</p>标签
        # 使用更精确的方法，避免统一正则表达式的问题
        
        # 方法1：处理完整的<li><p>...</p></li>结构
        def process_li_tags(match):
            li_content = match.group(1)
            
            # 首先移除所有的<p>和</p>标签
            li_content = re.sub(r'<p>|</p>', '', li_content)
            
            # 处理嵌套列表：确保<ul>或<ol>在单独的一行开始
            if '<ul>' in li_content or '<ol>' in li_content:
                # 将内容分为嵌套列表前和嵌套列表部分
                # 找到第一个<ul>或<ol>的位置
                list_start_match = re.search(r'(<ul>|<ol>)', li_content)
                if list_start_match:
                    list_start_pos = list_start_match.start()
                    # 嵌套列表前的内容
                    before_list = li_content[:list_start_pos].strip()
                    # 嵌套列表部分
                    list_part = li_content[list_start_pos:]
                    
                    # 压缩嵌套列表前内容的空白字符到一行
                    before_list = re.sub(r'\s+', ' ', before_list).strip()
                    
                    # 格式化嵌套列表部分，确保每个<li>在一行
                    list_part = re.sub(r'>\s*<', '><', list_part)  # 移除标签间的空白
                    list_part = re.sub(r'<li>\s*', '<li>', list_part)  # 移除<li>后的空白
                    list_part = re.sub(r'\s*</li>', '</li>', list_part)  # 移除</li>前的空白
                    
                    # 如果嵌套列表前有内容，确保嵌套列表在新的一行开始
                    if before_list:
                        li_content = f'{before_list}\n{list_part}'
                    else:
                        li_content = list_part
                # 确保<li>标签的开始和结束在一行，但保留嵌套列表的换行
                # 先压缩所有空白字符
                li_content = re.sub(r'\s+', ' ', li_content)
                # 然后确保<ul>或<ol>单独一行，后面的内容放到下一行
                #li_content = re.sub(r' (<ul>|<ol>)', r'\n\1\n', li_content)
                li_content = li_content.strip()
            else:
                # 对于不包含嵌套列表的情况，压缩所有空白字符到一行
                li_content = re.sub(r'\s+', ' ', li_content).strip()
            
            return f'<li>{li_content}</li>'
        
        # 使用更精确的正则匹配每个<li>标签
        html_content = re.sub(r'<li>(.*?)</li>', process_li_tags, html_content, flags=re.DOTALL)
        return html_content

    def _process_list(self, doc, list_tag, list_style, level=0, start_number=1):
        """处理列表标签，支持嵌套列表"""
        # 对于有序列表，手动设置起始序号
        if list_style == 'ListNumber':
            # 为每个有序列表项手动编号
            for i, li in enumerate(list_tag.find_all('li', recursive=False), start=start_number):
                # 创建列表项段落
                paragraph = doc.add_paragraph()
                
                # 设置缩进级别：按层级递增缩进
                if level > 0:  # 第二层及以上都缩进
                    paragraph.paragraph_format.left_indent = Inches(0.2 * level)
                # 第一层（level=0）不设置缩进
                
                # 手动添加序号
                run = paragraph.add_run(f"{i}. ")
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
                
                # 处理li的所有内容，包括文本和嵌套列表
                self._process_li_word_content(paragraph, li, list_style, level, doc)
        else:
            # 无序列表保持原有逻辑
            for li in list_tag.find_all('li', recursive=False):
                # 创建列表项段落
                paragraph = doc.add_paragraph(style=list_style)
                
                # 设置缩进级别：按层级递增缩进
                if level > 0:  # 第二层及以上都缩进
                    paragraph.paragraph_format.left_indent = Inches(0.2 * level)
                # 第一层（level=0）不设置缩进
                
                # 处理li的所有内容，包括文本和嵌套列表
                self._process_li_word_content(paragraph, li, list_style, level, doc)

    def _process_paragraph_content(self, paragraph, p_tag):
        """处理p标签内容（Word文档生成）"""
        # 处理p标签的所有子元素
        for child in p_tag.children:
            if hasattr(child, 'name'):
                if child.name == 'strong':
                    # 处理粗体文本
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                else:
                    # 处理其他标签
                    text = child.get_text().strip()
                    if text:
                        paragraph.add_run(text)
            elif isinstance(child, NavigableString) and child.strip():
                # 处理纯文本，去除前后空白
                text = child.strip()
                if text:
                    paragraph.add_run(text)

    def _process_li_word_content(self, paragraph, li_tag, parent_style, level=0, doc=None):
        """处理li标签内容（Word文档生成）"""
        # 处理li标签的所有子元素
        for child in li_tag.children:
            if hasattr(child, 'name'):
                if child.name == 'strong':
                    # 处理粗体文本
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['ul', 'ol']:
                    # 处理嵌套列表 - 创建新的段落并设置缩进
                    # 对于嵌套的有序列表，确保从1开始编号
                    nested_style = 'ListBullet' if child.name == 'ul' else 'ListNumber'
                    self._process_nested_list_separate(child, nested_style, level + 1, doc)
                else:
                    # 处理其他标签
                    text = child.get_text().strip()
                    if text:
                        paragraph.add_run(text)
            elif isinstance(child, NavigableString) and child.strip():
                # 处理纯文本，去除前后空白
                text = child.strip()
                if text:
                    paragraph.add_run(text)

    def _process_nested_list_separate(self, list_tag, parent_style, level, doc):
        """处理嵌套列表（单独段落）"""
        # 为嵌套列表创建新的段落
        # 对于嵌套的有序列表，确保从1开始编号
        if parent_style == 'ListNumber':
            # 手动处理有序列表
            for i, li in enumerate(list_tag.find_all('li', recursive=False), start=1):
                # 创建新的段落
                paragraph = doc.add_paragraph()
                
                # 设置悬挂缩进，确保列表符号和文字都能正确显示
                paragraph.paragraph_format.first_line_indent = Inches(-0.2)
                paragraph.paragraph_format.left_indent = Inches(0.2 * level + 0.2)
                
                # 手动添加序号
                run = paragraph.add_run(f"{i}. ")
                run.bold = False
                
                # 处理列表项内容（不包含嵌套列表，因为嵌套列表会在_process_li_word_content中处理）
                self._process_li_word_content(paragraph, li, parent_style, level, doc)
        else:
            # 无序列表保持原有逻辑
            for li in list_tag.find_all('li', recursive=False):
                # 创建新的段落
                paragraph = doc.add_paragraph(style=parent_style)
                
                # 设置悬挂缩进，确保列表符号和文字都能正确显示
                paragraph.paragraph_format.first_line_indent = Inches(-0.2)
                paragraph.paragraph_format.left_indent = Inches(0.2 * level + 0.2)
                
                # 处理列表项内容（不包含嵌套列表，因为嵌套列表会在_process_li_word_content中处理）
                self._process_li_word_content(paragraph, li, parent_style, level, doc)

    def to_word(self,file_name):
        """
        将markdown文本转换成word，
        :param file_name
        :return: 返回文件名
        """
        if file_name is None or file_name =='':
            file_name = CommonUtils.generate_uuid() + ".docx"
        elif not file_name.endswith(".docx"):
            file_name += ".docx"
        doc = Document()
        # 设置文档默认字体为宋体（同时设置中文字体和西文字体）
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Normal'].font.size = Pt(12)  # 宋体小四
        doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
        doc.styles['Normal'].font.italic = False  # 去除斜体
        
        # 设置标题样式也使用宋体
        for i in range(1, 7):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = '宋体'
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
            heading_style.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
            heading_style.font.italic = False  # 去除斜体
        soup = BeautifulSoup(self.html_content, 'html.parser')
        # 智能处理标签：对列表标签使用recursive=False避免重复，其他标签正常处理
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img','ul','ol']):
            if tag.name in ['ul', 'ol']:
                # 跳过嵌套的列表标签，它们会在_process_li_word_content中处理
                if tag.find_parent(['ul', 'ol']):
                    continue
                else:
                    # 处理顶层列表标签
                    if tag.name == 'ul':
                        self._process_list(doc, tag, 'ListBullet', 0)
                    elif tag.name == 'ol':
                        self._process_list(doc, tag, 'ListNumber', 0)
                    continue
            
            # 处理非列表标签
            if tag.name == 'h1':
                heading = doc.add_paragraph()
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)  # 宋体小二
                run.font.name = '宋体'  # 设置为宋体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
            elif tag.name == 'h2':
                heading = doc.add_paragraph()
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)  # 宋体三号
                run.font.name = '宋体'  # 设置为宋体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
            elif tag.name == 'h3':
                heading = doc.add_paragraph()
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)  # 宋体四号
                run.font.name = '宋体'  # 设置为宋体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
            elif tag.name == 'h4':
                heading = doc.add_paragraph()
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)  # 宋体小四
                run.font.name = '宋体'  # 设置为宋体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
            elif tag.name == 'h5':
                heading = doc.add_paragraph()
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)  # 宋体小四
                run.font.name = '宋体'  # 设置为宋体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
            elif tag.name == 'h6':
                heading = doc.add_paragraph()
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)  # 宋体小四
                run.font.name = '宋体'  # 设置为宋体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                run.font.italic = False  # 去除斜体
            elif tag.name == 'p':
                paragraph = doc.add_paragraph()
                # 处理p标签内的所有子元素，支持strong等内联标签
                self._process_paragraph_content(paragraph, tag)
                # 设置段落字体大小和颜色
                for run in paragraph.runs:
                    run.font.size = Pt(12)  # 宋体小四
                    run.font.name = '宋体'  # 设置为宋体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                    run.font.italic = False  # 去除斜体
            elif tag.name == 'img':
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        fixed_width = Inches(6)

                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'ul':
                self._process_list(doc, tag, 'ListBullet', 0)
            elif tag.name == 'ol':
                self._process_list(doc, tag, 'ListNumber', 0)
            elif tag.name == 'table':
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for cell in table.rows[0].cells:
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = '宋体'  # 设置为宋体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
                    run.font.italic = False  # 去除斜体
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1

                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)
                
                # 在表格后添加空段落，增加与下一行内容的距离
                doc.add_paragraph()

        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,file_name,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param file_name:
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        if file_name is None or file_name=='':
            file_name = CommonUtils.generate_uuid()+ ".pdf";
        elif not file_name.endswith(".pdf"):
            file_name += ".pdf"
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name
